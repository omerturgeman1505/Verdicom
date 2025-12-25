import re
import os
import requests
from bs4 import BeautifulSoup
from flask import Flask, request, render_template, redirect, url_for, jsonify
from dotenv import load_dotenv
import certifi
import ssl
import urllib3
import json
from datetime import datetime
from collections import defaultdict
from functools import lru_cache
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse
import sqlite3
import tempfile
import traceback
import hashlib
import base64
import zipfile
import io
import xml.etree.ElementTree as ET
import shutil
from werkzeug.utils import secure_filename
import subprocess
import whois
import dns.resolver
# Import Msg-Reader wrapper for enhanced attachment extraction
try:
    from msg_reader_wrapper import extract_email_attachments
    MSG_READER_AVAILABLE = True
    print("[App] ✓ Msg-Reader wrapper available")
except ImportError as e:
    MSG_READER_AVAILABLE = False
    print(f"[App] ⚠️ Msg-Reader wrapper not available: {e}")

# Import DKIM check functions from Magic-Spoofing
try:
    import sys
    magic_spoofing_path = os.path.join(os.path.dirname(__file__), 'Magic-Spoofing')
    if magic_spoofing_path not in sys.path:
        sys.path.insert(0, magic_spoofing_path)
    from dns_checks import check_dkim, check_spf, check_dmarc
    from utils import setup_colors
    DKIM_CHECK_AVAILABLE = True
    SPF_CHECK_AVAILABLE = True
    DMARC_CHECK_AVAILABLE = True
    print("[App] ✓ Magic-Spoofing DKIM, SPF, and DMARC checks available")
except ImportError as e:
    DKIM_CHECK_AVAILABLE = False
    SPF_CHECK_AVAILABLE = False
    DMARC_CHECK_AVAILABLE = False
    print(f"[App] ⚠️ Magic-Spoofing DKIM/SPF/DMARC checks not available: {e}")
try:
    import extract_msg
    EXTRACT_MSG_AVAILABLE = True
except ImportError:
    EXTRACT_MSG_AVAILABLE = False
try:
    from msg_parser import MsOxMessage
    MSG_PARSER_AVAILABLE = True
except ImportError:
    MSG_PARSER_AVAILABLE = False
try:
    import filetype
    FILETYPE_AVAILABLE = True
except ImportError:
    FILETYPE_AVAILABLE = False
try:
    from pdfminer.high_level import extract_text
    from pdfminer.layout import LAParams
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False
    print("[Warning] pdfminer.six not installed. PDF text extraction will be disabled.")
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    print("[Warning] olefile not installed. OLE structure validation will be disabled.")
# LinksReader dependencies for enhanced link extraction
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("[Warning] PyPDF2 not installed. Enhanced PDF link extraction will be disabled.")
try:
    from docx import Document
    DOCX_LIB_AVAILABLE = True
except ImportError:
    DOCX_LIB_AVAILABLE = False
    print("[Warning] python-docx not installed. Enhanced DOCX link extraction will be disabled.")

# Active Recon (OSINT) dependencies
try:
    import whois
    WHOIS_AVAILABLE = True
except ImportError:
    WHOIS_AVAILABLE = False
    print("[Warning] python-whois not installed. Active recon WHOIS lookups will be disabled.")
try:
    import dns.resolver
    DNS_AVAILABLE = True
except ImportError:
    DNS_AVAILABLE = False
    print("[Warning] dnspython not installed. Active recon DNS lookups will be disabled.")

load_dotenv()
app = Flask(__name__)

# Template filter for converting Unix timestamp to readable date
@app.template_filter('timestamp_to_date')
def timestamp_to_date_filter(timestamp):
    """Convert Unix timestamp to readable date string."""
    try:
        from datetime import datetime
        dt = datetime.fromtimestamp(int(timestamp))
        return dt.strftime('%Y-%m-%d %H:%M:%S')
    except:
        return 'Invalid date'

# ======= DATABASE SETUP =======
DB_PATH = "threat_intel_cache.db"

def init_db():
    """Initialize SQLite database for caching threat intelligence data."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS cache (
            id TEXT NOT NULL,
            data TEXT NOT NULL,
            timestamp REAL NOT NULL,
            type TEXT NOT NULL,
            PRIMARY KEY (id, type)
        )
    ''')
    # Create index for faster history queries
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_timestamp ON cache(timestamp DESC)
    ''')
    
    # Create AI training data table for lifelong learning
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS ai_training_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            domain TEXT NOT NULL,
            features TEXT NOT NULL,
            verdict TEXT,
            user_feedback INTEGER,
            created_at REAL NOT NULL,
            updated_at REAL
        )
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_domain ON ai_training_data(domain)
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_created_at ON ai_training_data(created_at DESC)
    ''')
    
    conn.commit()
    conn.close()
    print(f"[DB] Database initialized at {DB_PATH}")

# Initialize database on startup
init_db()

# ======= ATTACHMENT CACHE (for HTML viewing) =======
# Store all attachments temporarily for viewing
# Format: {hash: {'content': bytes, 'timestamp': float, 'filename': str, 'mime_type': str}}
_attachment_cache = {}
_ATTACHMENT_CACHE_TTL = 3600  # 1 hour
_ATTACHMENT_CACHE_CLEANUP_INTERVAL = 300  # Clean up every 5 minutes
_last_attachment_cache_cleanup = time.time()

def _cleanup_attachment_cache():
    """Remove expired attachments from cache."""
    global _last_attachment_cache_cleanup
    current_time = time.time()
    
    # Only cleanup every N seconds to avoid overhead
    if current_time - _last_attachment_cache_cleanup < _ATTACHMENT_CACHE_CLEANUP_INTERVAL:
        return
    
    _last_attachment_cache_cleanup = current_time
    expired_hashes = []
    
    for hash_val, data in _attachment_cache.items():
        if current_time - data['timestamp'] > _ATTACHMENT_CACHE_TTL:
            expired_hashes.append(hash_val)
    
    for hash_val in expired_hashes:
        del _attachment_cache[hash_val]
    
    if expired_hashes:
        print(f"[Attachment Cache] Cleaned up {len(expired_hashes)} expired attachments")

def cache_attachment(file_hash: str, file_content: bytes, filename: str, mime_type: str = None):
    """Cache an attachment for later viewing."""
    _cleanup_attachment_cache()
    
    if file_hash and file_content:
        # Normalize hash to lowercase and strip whitespace for consistency
        file_hash = str(file_hash).lower().strip()
        
        # Verify hash is valid SHA256 (64 hex chars)
        if len(file_hash) != 64 or not all(c in '0123456789abcdef' for c in file_hash):
            print(f"[Attachment Cache] WARNING: Invalid hash format: {file_hash[:32]}... (length: {len(file_hash)})")
            # Recalculate hash if invalid
            file_hash = hashlib.sha256(file_content).hexdigest()
            print(f"[Attachment Cache] Recalculated hash: {file_hash[:32]}...")
        
        _attachment_cache[file_hash] = {
            'content': file_content,
            'timestamp': time.time(),
            'filename': filename,
            'mime_type': mime_type
        }
        print(f"[Attachment Cache] ✓ Cached attachment: {filename}")
        print(f"[Attachment Cache]   Hash: {file_hash[:32]}... (length: {len(file_hash)})")
        print(f"[Attachment Cache]   Size: {len(file_content)} bytes, MIME: {mime_type or 'unknown'}")
        print(f"[Attachment Cache]   Total cached: {len(_attachment_cache)} attachment(s)")

def get_cached_attachment(file_hash: str):
    """Retrieve a cached attachment by hash."""
    _cleanup_attachment_cache()
    
    # Normalize hash to lowercase for consistency
    file_hash = file_hash.lower().strip()
    
    if file_hash in _attachment_cache:
        data = _attachment_cache[file_hash]
        # Check if expired
        if time.time() - data['timestamp'] > _ATTACHMENT_CACHE_TTL:
            del _attachment_cache[file_hash]
            print(f"[Attachment Cache] Attachment expired: {file_hash[:16]}...")
            return None
        print(f"[Attachment Cache] ✓ Found attachment: {data['filename']} (hash: {file_hash[:16]}...)")
        return data
    
    print(f"[Attachment Cache] ✗ Attachment not found: {file_hash[:16]}...")
    return None

# Legacy functions for PDF compatibility
def cache_pdf_attachment(pdf_hash: str, pdf_content: bytes, filename: str):
    """Cache a PDF attachment for later HTML conversion."""
    cache_attachment(pdf_hash, pdf_content, filename, 'application/pdf')

def get_cached_pdf(pdf_hash: str):
    """Retrieve a cached PDF attachment by hash."""
    return get_cached_attachment(pdf_hash)

def get_cached_data(indicator: str, fetch_function, cache_type: str, max_age_hours: int = 24):
    """
    Smart caching wrapper for threat intelligence lookups.
    
    Args:
        indicator: The indicator to lookup (IP, hash, domain, etc.)
        fetch_function: Function to call if cache miss
        cache_type: Type of cache ('vt' or 'abuseipdb')
        max_age_hours: Maximum age of cached data in hours (default 24)
    
    Returns:
        Tuple of (data, error) - same format as fetch_function
    """
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Check cache
    cursor.execute('''
        SELECT data, timestamp FROM cache 
        WHERE id = ? AND type = ?
    ''', (indicator, cache_type))
    
    row = cursor.fetchone()
    current_time = time.time()
    
    if row:
        cached_data, cached_timestamp = row
        age_hours = (current_time - cached_timestamp) / 3600
        
        if age_hours < max_age_hours:
            # Cache hit - return cached data
            conn.close()
            try:
                data = json.loads(cached_data)
                return data, None
            except json.JSONDecodeError:
                # Corrupted cache, fall through to fetch
                pass
    
    # Cache miss or expired - fetch fresh data
    data, error = fetch_function(indicator)
    
    # Save to cache if successful (even if error, save empty data to track search history)
    # Always save to maintain persistent history
    try:
        cache_data = data if data else {}
        cursor.execute('''
            INSERT OR REPLACE INTO cache (id, data, timestamp, type)
            VALUES (?, ?, ?, ?)
        ''', (indicator, json.dumps(cache_data), current_time, cache_type))
        conn.commit()
        print(f"[Cache] Saved {cache_type} lookup for {indicator} (persistent)")
    except Exception as e:
        print(f"[Cache] Error saving to cache: {e}")
    
    conn.close()
    return data, error

def get_recent_history(limit: int = 10):
    """Get recent search history from cache - persistent across sessions."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Get unique indicators (one per indicator, regardless of type) 
    # ordered by most recent timestamp - this ensures history persists
    cursor.execute('''
        SELECT id, type, MAX(timestamp) as latest_timestamp 
        FROM cache
        GROUP BY id
        ORDER BY latest_timestamp DESC
        LIMIT ?
    ''', (limit,))
    
    rows = cursor.fetchall()
    conn.close()
    
    history = []
    for row in rows:
        indicator, cache_type, timestamp = row
        history.append({
            'indicator': indicator,
            'type': cache_type,
            'timestamp': timestamp,
            'age_hours': (time.time() - timestamp) / 3600
        })
    
    print(f"[History] Retrieved {len(history)} items from persistent cache (DB: {DB_PATH})")
    return history

def get_cached_lookup_data(query: str):
    """Get lookup data from cache only (no API calls). Returns None if not cached."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Get VT data
    cursor.execute('''
        SELECT data, timestamp FROM cache 
        WHERE id = ? AND type = 'vt'
    ''', (query,))
    vt_row = cursor.fetchone()
    
    vt_data = None
    if vt_row:
        try:
            vt_data = json.loads(vt_row[0])
        except json.JSONDecodeError:
            pass
    
    # Get AbuseIPDB data (if IP)
    ipdb_data = None
    if guess_type(query) == "ip":
        cursor.execute('''
            SELECT data, timestamp FROM cache 
            WHERE id = ? AND type = 'abuseipdb'
        ''', (query,))
        ipdb_row = cursor.fetchone()
        if ipdb_row:
            try:
                ipdb_data = json.loads(ipdb_row[0])
            except json.JSONDecodeError:
                pass
    
    conn.close()
    
    return vt_data, ipdb_data

# ======= LOAD KEYS SAFELY =======
VT_API_KEY = os.environ.get("VT_API_KEY")
ABUSEIPDB_API_KEY = os.environ.get("ABUSEIPDB_API_KEY")
RAPIDAPI_KEY = os.environ.get("RAPIDAPI_KEY")

# ======= CONFIGURABLE DOMAIN PATTERNS =======
# Allow customizing email domain patterns via environment variable
# Default: mobileye.com (can be changed to any domain or multiple domains)
CUSTOM_EMAIL_DOMAIN = os.environ.get("CUSTOM_EMAIL_DOMAIN", "mobileye.com")
# Support multiple domains separated by comma
EMAIL_DOMAINS = [d.strip().lower() for d in CUSTOM_EMAIL_DOMAIN.split(",") if d.strip()]
# Build regex pattern for all configured domains
EMAIL_DOMAIN_PATTERN = "|".join([re.escape(d) for d in EMAIL_DOMAINS])
_MOBILEYE_EMAIL_REGEX_PATTERN = re.compile(rf"\b([A-Za-z0-9._-]+)@({EMAIL_DOMAIN_PATTERN})\b", re.IGNORECASE)
# ================================

# ======= SSL VERIFICATION CONFIGURATION =======
# Allow disabling SSL verification via environment variable (for corporate proxies)
# WARNING: Only use this if you understand the security implications
# To disable SSL verification, add DISABLE_SSL_VERIFY=true to your .env file
DISABLE_SSL_VERIFY = os.environ.get("DISABLE_SSL_VERIFY", "false").lower() == "true"
CUSTOM_CA_BUNDLE = os.environ.get("REQUESTS_CA_BUNDLE") or os.environ.get("CURL_CA_BUNDLE")

# Configure SSL verification with smart fallback
def get_ssl_verify():
    """
    Returns SSL verification setting based on environment configuration.
    Handles corporate proxies and self-signed certificates gracefully.
    """
    if DISABLE_SSL_VERIFY:
        print("[WARNING] SSL verification is DISABLED. This is not recommended for production!")
        try:
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        except:
            pass
        return False
    
    # Use custom CA bundle if specified
    if CUSTOM_CA_BUNDLE and os.path.exists(CUSTOM_CA_BUNDLE):
        return CUSTOM_CA_BUNDLE
    
    # Try system default certificates first (includes Windows certificate store)
    # This works better in corporate environments where CAs are installed system-wide
    return True  # requests will use system default certificate store

def get_requests_session():
    """
    Creates and returns a requests.Session with proper SSL/TLS configuration.
    This session can be reused across multiple requests for better performance and SSL handling.
    """
    session = requests.Session()
    ssl_verify = get_ssl_verify()
    
    # Configure session with SSL verification setting
    session.verify = ssl_verify
    
    # Set default timeout for all requests in this session
    session.timeout = 10
    
    # Add default headers for better compatibility
    session.headers.update({
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    })
    
    return session

# Configure urllib3 to handle SSL/TLS better
def configure_ssl_for_requests():
    """
    Configure urllib3 and requests to handle SSL/TLS negotiation better.
    This helps prevent ALPN negotiation failures.
    """
    try:
        # Create a custom SSL context that's more flexible
        import ssl as ssl_module
        
        # Disable SSL warnings if verification is disabled
        if DISABLE_SSL_VERIFY:
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        # Set default SSL context options for better compatibility
        try:
            # Try to use system default SSL context
            default_context = ssl_module.create_default_context()
            # Allow more protocol versions for compatibility
            default_context.options &= ~ssl_module.OP_NO_SSLv2
            default_context.options &= ~ssl_module.OP_NO_SSLv3
        except:
            pass
        
        print("[SSL] SSL/TLS configuration initialized")
    except Exception as e:
        print(f"[SSL] Warning: Could not configure SSL context: {e}")

# Initialize SSL configuration on module load
configure_ssl_for_requests()
# ================================

# --- VirusTotal Comments API ---
def get_vt_comments(limit=10, filter_tag=None, cursor=None):
    """
    Fetch latest comments from VirusTotal API.
    
    Args:
        limit: Number of comments to retrieve (default: 10)
        filter_tag: Filter comments by tag (e.g., 'malware', 'phishing')
        cursor: Continuation cursor for pagination
    
    Returns:
        Tuple of (comments_data, error)
    """
    if not VT_API_KEY:
        return None, "VT_API_KEY is not configured."
    
    url = "https://www.virustotal.com/api/v3/comments"
    headers = {"x-apikey": VT_API_KEY, "Accept": "application/json"}
    params = {"limit": limit}
    
    if filter_tag:
        params["filter"] = f"tag:{filter_tag}"
    
    if cursor:
        params["cursor"] = cursor
    
    ssl_verify = get_ssl_verify()
    session = get_requests_session()
    
    try:
        # Use session for better SSL/TLS handling and retry logic
        r = session.get(url, headers=headers, params=params, timeout=10, verify=ssl_verify)
        print(f"[VT Comments] GET {url} -> {r.status_code}")
        
        if r.status_code != 200:
            error_snippet = r.text[:400].replace("\n", " ")
            return None, f"HTTP {r.status_code}: {error_snippet}"
        
        try:
            js = r.json()
            return js, None
        except Exception as json_err:
            return None, f"Invalid JSON from VirusTotal: {json_err}"
            
    except requests.exceptions.SSLError as e:
        error_msg = str(e)
        if not ssl_verify:
            return None, f"SSL Error (verification disabled): {e}"
        
        # Handle ALPN negotiation failures and other SSL errors
        if ("certificate verify failed" in error_msg or "self-signed certificate" in error_msg or 
            "ALPN" in error_msg or "negotiation" in error_msg.lower()):
            print(f"[VT Comments] SSL/TLS error detected, retrying without verification")
            try:
                urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
                fallback_session = requests.Session()
                fallback_session.verify = False
                r = fallback_session.get(url, headers=headers, params=params, timeout=10)
                if r.status_code == 200:
                    try:
                        js = r.json()
                        return js, None
                    except Exception as json_err:
                        return None, f"Invalid JSON from VirusTotal: {json_err}"
                else:
                    return None, f"HTTP {r.status_code} (SSL verification disabled)"
            except Exception as retry_err:
                return None, f"SSL Error (retry failed): {retry_err}"
        return None, f"SSL Error: {e}"
        
    except requests.exceptions.Timeout:
        return None, "Request timeout (10 seconds)"
    except Exception as e:
        return None, f"Error fetching comments: {e}"

def get_vt_comments_cached(limit=10, filter_tag=None, cursor=None):
    """Get VirusTotal comments with caching."""
    cache_key = f"vt_comments_{limit}_{filter_tag or 'all'}_{cursor or 'none'}"
    return get_cached_data(cache_key, lambda k: get_vt_comments(limit, filter_tag, cursor), 'vt_comments', max_age_hours=1)

# --- 2. Updated News Function (Web Scraping) ---
def get_cyber_news():
    """Fetches top 5 cyber news headlines by scraping The Hacker News."""
    URL = "https://thehackernews.com/"
    news_items = []
    
    user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36"
    headers = {"User-Agent": user_agent}
    
    # First try with SSL verification
    ssl_verify = get_ssl_verify()
    try:
        print(f"[News] Scraping news from {URL}")
        r = requests.get(URL, headers=headers, timeout=3, verify=ssl_verify)
        
        if r.status_code != 200:
            print(f"[News] Error: Got status code {r.status_code}")
            return []

        soup = BeautifulSoup(r.text, 'html.parser')
        headlines = soup.find_all('h3', class_='home-title', limit=5)
        
        for item in headlines:
            link_tag = item.find_parent('a')
            if link_tag and link_tag.has_attr('href'):
                news_items.append({
                    "title": item.text.strip(),
                    "link": link_tag['href']
                })
                
        print(f"[News] Found {len(news_items)} items via scraping.")
    except requests.exceptions.SSLError as e:
        error_msg = str(e)
        # If SSL verification is already disabled, don't retry
        if not ssl_verify:
            print(f"Error scraping cyber news (SSL verification disabled): {e}")
            return []
        
        # If SSL error occurs and we're using verification, try without verification as fallback
        if "certificate verify failed" in error_msg or "self-signed certificate" in error_msg or "CERTIFICATE_VERIFY_FAILED" in error_msg:
            print(f"[News] SSL verification failed, retrying without verification (corporate proxy detected)")
            try:
                urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
                r = requests.get(URL, headers=headers, timeout=3, verify=False)
                print(f"[News] Got status code {r.status_code} (SSL verification disabled)")
                
                if r.status_code != 200:
                    print(f"[News] Error: Got status code {r.status_code}")
                    return []

                soup = BeautifulSoup(r.text, 'html.parser')
                headlines = soup.find_all('h3', class_='home-title', limit=5)
                
                for item in headlines:
                    link_tag = item.find_parent('a')
                    if link_tag and link_tag.has_attr('href'):
                        news_items.append({
                            "title": item.text.strip(),
                            "link": link_tag['href']
                        })
                        
                print(f"[News] Found {len(news_items)} items via scraping.")
            except Exception as retry_e:
                print(f"Error scraping cyber news (retry failed): {retry_e}")
                return []
        return []
    except requests.exceptions.Timeout:
        print(f"[News] Timeout after 3 seconds, returning empty list")
        return []
    except Exception as e:
        print(f"Error scraping cyber news: {e}")
        return []
    
    return news_items
# --- End Update ---


# ---------- helpers ----------
_ipv4 = re.compile(r"^((25[0-5]|2[0-4]\d|1?\d?\d)(\.|$)){4}$")
_domain = re.compile(r"^(?=.{1,253}$)(?!-)(?:[A-Za-z0-9-]{1,63}\.)+[A-Za-z]{2,63}$")
_md5 = re.compile(r"^[A-Fa-f0-9]{32}$")
_sha1 = re.compile(r"^[A-Fa-f0-9]{40}$")
_sha224 = re.compile(r"^[A-Fa-f0-9]{56}$")
_sha256 = re.compile(r"^[A-Fa-f0-9]{64}$")
_sha384 = re.compile(r"^[A-Fa-f0-9]{96}$")
_sha512 = re.compile(r"^[A-Fa-f0-9]{128}$")
_sha3_224 = re.compile(r"^[A-Fa-f0-9]{56}$")  # Same length as SHA-224
_sha3_256 = re.compile(r"^[A-Fa-f0-9]{64}$")  # Same length as SHA-256
_sha3_384 = re.compile(r"^[A-Fa-f0-9]{96}$")  # Same length as SHA-384
_sha3_512 = re.compile(r"^[A-Fa-f0-9]{128}$")  # Same length as SHA-512

# ======= COMPILED REGEX PATTERNS FOR SIEM LOG PARSING (Module-level for performance) =======
# Compiled once at module load to avoid recompiling on every function call
_IP_REGEX_PATTERN = re.compile(r"(?:(25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(25[0-5]|2[0-4]\d|1?\d?\d)")
_IP_REGEX_PATTERN_BYTES = re.compile(rb'(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)')
_URL_REGEX_PATTERN_BYTES = re.compile(rb'https?://[a-zA-Z0-9./?=_~-]+')
_ASCII_STRINGS = re.compile(rb'[\x20-\x7E]{5,}')  # 5+ printable ASCII characters
_PORT_REGEX_PATTERN = re.compile(r":(\d{1,5})\b")
_DOMAIN_REGEX_PATTERN = re.compile(r"(?:https?://)?([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)+)")
_URL_REGEX_PATTERN = re.compile(r"https?://[^\s<>\"{}|\\^`\[\]]+")
_USER_AGENT_REGEX_PATTERN = re.compile(r"User-Agent['\":\s]+([^\n\"']+)", re.IGNORECASE)
_HTTP_METHOD_REGEX_PATTERN = re.compile(r"\b(GET|POST|PUT|DELETE|PATCH|HEAD|OPTIONS|CONNECT|TRACE)\s+")
_HTTP_STATUS_REGEX_PATTERN = re.compile(r"\s(\d{3})\s")
_EMAIL_REGEX_PATTERN = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9_.-]+")
# Note: _MOBILEYE_EMAIL_REGEX_PATTERN is now defined above with configurable domains
_USERNAME_PATTERNS_COMPILED = [
    re.compile(r"(?:user(?:name)?|account|acct|login|principal|subject|actor)[\s_\-]*(?:name|id)?\s*[:=]\s*['\"]?([a-zA-Z0-9_\-\.@\\/]+)", re.IGNORECASE),
    re.compile(r"(?:Account Name|TargetUserName|SubjectUserName|TargetUserSid|SubjectUserSid)\s*[:=]\s*['\"]?([a-zA-Z0-9_\-\.@\\/]+)", re.IGNORECASE),
    re.compile(r"User\s*'([^']+)'", re.IGNORECASE),
    re.compile(r'User\s*"([^"]+)"', re.IGNORECASE)
]
_FILE_PATH_REGEX_PATTERN = re.compile(r"[C-Z]:\\[^\s<>\"|]+\.[a-zA-Z0-9]{1,5}|/(?:[^/\s]+/)+[^/\s]+\.[a-zA-Z0-9]{1,5}")
_PROCESS_REGEX_PATTERN = re.compile(r"(?:process|exe|program)[\s:=]+([a-zA-Z0-9_\-\.]+\.(exe|dll|bat|sh|py|js))", re.IGNORECASE)
_HASH_REGEX_PATTERN = re.compile(r"\b[A-Fa-f0-9]{32,128}\b")
_SRC_IP_CAMEL_PATTERN = re.compile(r"(?:src|source)[Ii][Pp]\s*[:=]\s*['\"]?((?:(25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(25[0-5]|2[0-4]\d|1?\d?\d))['\"]?", re.IGNORECASE)
_DST_IP_CAMEL_PATTERN = re.compile(r"(?:dst|dest|destination)[Ii][Pp]\s*[:=]\s*['\"]?((?:(25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(25[0-5]|2[0-4]\d|1?\d?\d))['\"]?", re.IGNORECASE)
_SRC_IP_PATTERN = re.compile(r"(?:src|source)(?:\s+ip|\s+_ip|\s+_address|\s+address)?\s*[:=]\s*['\"]?((?:(25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(25[0-5]|2[0-4]\d|1?\d?\d))['\"]?", re.IGNORECASE)
_DST_IP_PATTERN = re.compile(r"(?:dst|dest|destination)(?:\s+ip|\s+_ip|\s+_address|\s+address)?\s*[:=]\s*['\"]?((?:(25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(25[0-5]|2[0-4]\d|1?\d?\d))['\"]?", re.IGNORECASE)
_SRC_PORT_PATTERN = re.compile(r"src[^:]*:(\d{1,5})\b", re.IGNORECASE)
_DST_PORT_PATTERN = re.compile(r"dst[^:]*:(\d{1,5})\b", re.IGNORECASE)
_PROTOCOL_PATTERN = re.compile(r"\b(TCP|UDP|ICMP|HTTP|HTTPS|FTP|SSH|TLS|SSL|SMTP|DNS)\b", re.IGNORECASE)
_TIMESTAMP_PATTERNS = [
    re.compile(r"(\d{4}-\d{2}-\d{2}[T\s]\d{2}:\d{2}:\d{2}(?:\.\d+)?(?:Z|[+-]\d{2}:\d{2})?)"),
    re.compile(r"(\d{2}/\d{2}/\d{4}\s+\d{2}:\d{2}:\d{2})"),
    re.compile(r"(\w+\s+\d{1,2}\s+\d{2}:\d{2}:\d{2})"),
    re.compile(r"(\d{10,13})")  # Unix timestamp
]
_EVENT_KEYWORDS_PATTERN = re.compile(r"\b(login|logout|access|denied|allowed|blocked|failed|success|error|alert|attack|malware|virus|firewall|intrusion)\b", re.IGNORECASE)
_RISK_KEYWORDS_HIGH_PATTERN = re.compile(r"\b(denied|blocked|failed|attack|malware|virus|intrusion|exploit|breach|unauthorized|suspicious)\b", re.IGNORECASE)
_RISK_KEYWORDS_MEDIUM_PATTERN = re.compile(r"\b(error|warning|alert|unusual|anomaly|abnormal)\b", re.IGNORECASE)
_WINDOWS_FILE_PATTERN = re.compile(r"[C-Z]:\\[^\s<>\"|\\]+\.[a-zA-Z0-9]{1,10}", re.IGNORECASE)
_UNIX_FILE_PATTERN = re.compile(r"(?:/|\.\.?/)[^\s<>\"|\\]+\.[a-zA-Z0-9]{1,10}", re.IGNORECASE)

# ======= PHISHING HUNTER MODULE - KEYWORDS =======
PHISHING_KEYWORDS = ['urgent', 'immediate', 'verify', 'account', 'suspended', 'expire', 'bank', 'password', 'credential', 'confirm', 'action required']
_QUOTED_FILE_PATTERN = re.compile(r"['\"]([^\s<>\"']+\.[a-zA-Z0-9]{1,10})['\"]", re.IGNORECASE)
_URL_FILE_PATTERN = re.compile(r"https?://[^\s<>\"']+/([^\s<>\"'/?]+\.(?:[a-zA-Z0-9]{1,10}))", re.IGNORECASE)
_KEYWORD_FILE_PATTERN = re.compile(r"(?:file|path|filename|target|source|destination)[\s:=]+['\"]?([^\s<>\"'|]+\.[a-zA-Z0-9]{1,10})", re.IGNORECASE)
_EXFIL_KEYWORDS_HIGH_PATTERN = re.compile(r"\b(export|backup|dump|sync|archive|exfil)\b", re.IGNORECASE)
_EXFIL_KEYWORDS_MEDIUM_PATTERN = re.compile(r"\b(upload|transfer|copy|move|send)\b", re.IGNORECASE)

# ======= UTILITY FUNCTIONS =======
def is_private_ip(ip: str) -> bool:
    """Check if an IP address is in private/internal ranges. Optimized version."""
    if not ip:
        return False
    parts = ip.split('.')
    if len(parts) != 4:
        return False
    try:
        first = int(parts[0])
        second = int(parts[1])
        if first == 10: return True  # 10.0.0.0/8
        if first == 172 and 16 <= second <= 31: return True  # 172.16.0.0/12
        if first == 192 and second == 168: return True  # 192.168.0.0/16
        if first == 127: return True  # 127.0.0.0/8
        if first == 169 and second == 254: return True  # 169.254.0.0/16
    except (ValueError, IndexError):
        return False
    return False

def extract_urls_from_docx(docx_data: bytes) -> list:
    """Extract URLs from DOCX file content using LinksReader approach (python-docx library)."""
    urls = []
    
    # Try using python-docx library (LinksReader approach) - more reliable
    if DOCX_LIB_AVAILABLE:
        try:
            doc = Document(io.BytesIO(docx_data))
            
            # Extract hyperlinks from relationships (like LinksReader does)
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel._target
                    if url:
                        urls.append(url)
            
            # Extract URLs from text using regex (like LinksReader does)
            for para in doc.paragraphs:
                url_matches = _URL_REGEX_PATTERN.findall(para.text)
                urls.extend(url_matches)
            
            # Also check tables for URLs
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            url_matches = _URL_REGEX_PATTERN.findall(para.text)
                            urls.extend(url_matches)
            
            print(f"[DOCX Extraction] Found {len(set(urls))} unique URLs using python-docx")
            return list(set(urls))  # Remove duplicates
        except Exception as e:
            print(f"[DOCX Extraction] Error using python-docx library: {e}, falling back to XML parsing...")
    
    # Fallback: Original XML parsing approach
    try:
        # DOCX is a ZIP archive
        with zipfile.ZipFile(io.BytesIO(docx_data), 'r') as zip_file:
            # Extract text from word/document.xml
            try:
                doc_xml = zip_file.read('word/document.xml')
                # Parse XML and extract URLs
                root = ET.fromstring(doc_xml)
                # Search for hyperlinks in the document
                # DOCX hyperlinks are in <w:hyperlink> elements with r:id attributes
                # We also search for plain URLs in text
                for elem in root.iter():
                    # Check text content for URLs
                    if elem.text:
                        url_matches = _URL_REGEX_PATTERN.findall(elem.text)
                        urls.extend(url_matches)
                    # Check tail content
                    if elem.tail:
                        url_matches = _URL_REGEX_PATTERN.findall(elem.tail)
                        urls.extend(url_matches)
            except Exception as xml_err:
                print(f"[DOCX Extraction] Error parsing document.xml: {xml_err}")
                # Fallback: search raw XML as text
                try:
                    doc_xml_str = doc_xml.decode('utf-8', errors='ignore')
                    url_matches = _URL_REGEX_PATTERN.findall(doc_xml_str)
                    urls.extend(url_matches)
                except:
                    pass
    except Exception as e:
        print(f"[DOCX Extraction] Error extracting URLs from DOCX: {e}")
    return list(set(urls))  # Remove duplicates

def extract_text_from_docx(docx_data: bytes) -> str:
    """Extract text content from DOCX file. DOCX is a ZIP archive containing XML files. Robust version."""
    text_parts = []
    
    # Validate that data looks like a ZIP file
    if not docx_data or len(docx_data) < 4:
        print("[DOCX Text Extraction] Invalid DOCX data (too short)")
        return ""
    
    # Check for ZIP signature (PK)
    if not docx_data.startswith(b'PK'):
        print(f"[DOCX Text Extraction] Invalid DOCX data (not a ZIP file, starts with: {docx_data[:4].hex()})")
        return ""
    
    try:
        # DOCX is a ZIP archive
        with zipfile.ZipFile(io.BytesIO(docx_data), 'r') as zip_file:
            # Verify it's a valid DOCX by checking for required files
            file_list = zip_file.namelist()
            if 'word/document.xml' not in file_list:
                print(f"[DOCX Text Extraction] Not a valid DOCX (missing document.xml). Files: {file_list[:5]}")
                return ""
            
            # Extract text from word/document.xml
            try:
                doc_xml = zip_file.read('word/document.xml')
                # Parse XML and extract text
                root = ET.fromstring(doc_xml)
                # Define namespace for Word documents
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Find all text elements
                for t in root.findall('.//w:t', ns):
                    if t.text:
                        text_parts.append(t.text)
                
                # If no text found with namespace, try without namespace
                if not text_parts:
                    for elem in root.iter():
                        if elem.text and elem.text.strip():
                            text_parts.append(elem.text)
                        if elem.tail and elem.tail.strip():
                            text_parts.append(elem.tail)
                            
                print(f"[DOCX Text Extraction] Extracted {len(text_parts)} text elements")
            except Exception as xml_err:
                print(f"[DOCX Text Extraction] Error parsing document.xml: {xml_err}")
                # Fallback: search raw XML as text (basic extraction)
                try:
                    doc_xml_str = doc_xml.decode('utf-8', errors='ignore')
                    # Simple text extraction - remove XML tags
                    text_content = re.sub(r'<[^>]+>', ' ', doc_xml_str)
                    text_content = ' '.join(text_content.split())  # Normalize whitespace
                    if text_content and len(text_content) > 50:  # Only use if we got meaningful content
                        text_parts.append(text_content)
                        print(f"[DOCX Text Extraction] Fallback extraction: {len(text_content)} characters")
                except Exception as fallback_err:
                    print(f"[DOCX Text Extraction] Fallback extraction also failed: {fallback_err}")
    except zipfile.BadZipFile as zip_err:
        print(f"[DOCX Text Extraction] Invalid ZIP structure: {zip_err}")
        return ""
    except Exception as e:
        print(f"[DOCX Text Extraction] Error extracting text from DOCX: {e}")
        return ""
    
    final_text = ' '.join(text_parts).strip()
    print(f"[DOCX Text Extraction] Final extracted text: {len(final_text)} characters")
    return final_text

# ======= NEWS CACHING =======
_news_cache = {"data": None, "timestamp": 0}
_NEWS_CACHE_TTL = 300  # 5 minutes

def get_cyber_news_cached():
    """Get cyber news with caching to avoid repeated HTTP requests."""
    current_time = time.time()
    if _news_cache["data"] and (current_time - _news_cache["timestamp"]) < _NEWS_CACHE_TTL:
        return _news_cache["data"]
    
    try:
        news_data = get_cyber_news()
        _news_cache["data"] = news_data
        _news_cache["timestamp"] = current_time
        return news_data
    except Exception as e:
        print(f"[News] Error fetching news: {e}")
        # Return empty list immediately on any failure to avoid blocking
        return []

def guess_type(q: str) -> str:
    q = (q or "").strip()
    if _ipv4.match(q): return "ip"
    # Check for URLs (http:// or https://)
    if q.startswith(("http://", "https://")):
        return "url"
    if _domain.match(q): return "domain"
    if (_sha512.match(q) or _sha384.match(q) or _sha256.match(q) or 
        _sha224.match(q) or _sha1.match(q) or _md5.match(q)): return "hash"
    return "domain" if "." in q else "unknown"

def get_vt_community_link(indicator: str, indicator_type: str = None) -> str:
    """
    Generate VirusTotal community tab URL for an indicator.
    
    Args:
        indicator: The indicator (URL, IP, domain, or hash)
        indicator_type: Optional type hint (url, ip, domain, hash)
    
    Returns:
        VirusTotal community tab URL
    """
    if not indicator:
        return None
    
    if indicator_type is None:
        indicator_type = guess_type(indicator)
    
    if indicator_type == "url":
        # For URLs, we need to base64 encode without padding
        import base64
        url_bytes = indicator.encode('utf-8')
        url_b64 = base64.urlsafe_b64encode(url_bytes).decode('utf-8').rstrip('=')
        return f"https://www.virustotal.com/gui/url/{url_b64}/community"
    elif indicator_type == "ip":
        return f"https://www.virustotal.com/gui/ip-address/{indicator}/community"
    elif indicator_type == "domain":
        return f"https://www.virustotal.com/gui/domain/{indicator}/community"
    elif indicator_type == "hash":
        return f"https://www.virustotal.com/gui/file/{indicator}/community"
    else:
        return None

def vt_url_for(q: str) -> str:
    t = guess_type(q)
    if t == "ip": return f"https://www.virustotal.com/api/v3/ip_addresses/{q}"
    if t == "url":
        # VirusTotal requires URLs to be Base64 URL-safe encoded without padding
        # Trim whitespace from URL before encoding
        url_trimmed = q.strip()
        # Generate URL Identifier for VT (Base64 without padding)
        url_id = base64.urlsafe_b64encode(url_trimmed.encode()).decode().strip("=")
        return f"https://www.virustotal.com/api/v3/urls/{url_id}"
    if t == "domain": return f"https://www.virustotal.com/api/v3/domains/{q}"
    if t == "hash": return f"https://www.virustotal.com/api/v3/files/{q}"
    return f"https://www.virustotal.com/api/v3/search?query={q}"

def safe_text(s, limit=400):
    s = "" if s is None else str(s).strip()
    return (s[:limit] + "…") if len(s) > limit else s

def summarize_vt(payload: dict) -> dict:
    """Summarizes a VT payload, handling both direct lookups and search results."""
    d = (payload or {}).get("data") or {}

    if isinstance(d, list):
        if not d:
            return {"name": "Not Found", "malicious": 0, "suspicious": 0, "harmless": 0, "undetected": 0, "timeout": 0, "community_score": 0, "kind": "—", "detections": []}
        d = d[0]

    attr = d.get("attributes") or {}
    stats = attr.get("last_analysis_stats") or {}
    votes = attr.get("total_votes", {})
    # For URLs, use the URL itself or the id
    name = attr.get("meaningful_name") or attr.get("url") or d.get("id") or "—"
    kind = d.get("type") or "—"
    
    detections_list = []
    if kind in ('file', 'ip_address', 'domain', 'url'):
        last_analysis_results = attr.get("last_analysis_results", {})
        for engine, details in last_analysis_results.items():
            detections_list.append({
                "engine": engine,
                "result": details.get('result'),
                "category": details.get('category', 'unknown')
            })
        
        category_order = {
            "malicious": 1, "suspicious": 2, "harmless": 3,
            "undetected": 4, "timeout": 5, "type-unsupported": 6, "unknown": 7
        }
        
        detections_list.sort(key=lambda x: (category_order.get(x['category'], 99), x['engine']))
    
    return {
        "kind": kind, "name": name,
        "malicious": stats.get("malicious", 0),
        "suspicious": stats.get("suspicious", 0),
        "harmless": stats.get("harmless", 0),
        "undetected": stats.get("undetected", 0),
        "timeout": stats.get("timeout", 0),
        "community_score": (votes.get("harmless", 0) - votes.get("malicious", 0)),
        "detections": detections_list
    }

def _vt_fetch_internal(query: str):
    """Internal VT fetch function (without caching)."""
    if not VT_API_KEY:
        return None, "VT_API_KEY is not configured."
    # Trim query to handle trailing spaces
    query = query.strip()
    url = vt_url_for(query)
    headers = {"x-apikey": VT_API_KEY, "Accept": "application/json"}
    
    # First try with SSL verification using shared session
    ssl_verify = get_ssl_verify()
    session = get_requests_session()
    try:
        # Use session for better SSL/TLS handling and retry logic
        r = session.get(url, headers=headers, timeout=10, verify=ssl_verify)
        print(f"[VT] {url} -> {r.status_code}")
        
        # If URL lookup returns 404, fallback to domain check
        if r.status_code == 404 and "/api/v3/urls/" in url:
            print(f"[VT] URL not found, falling back to Domain check for: {query}")
            # Extract domain
            try:
                parsed = urlparse(query)
                domain = parsed.netloc
                if ':' in domain:
                    domain = domain.split(':')[0]
                
                if domain:
                    fallback_url = f"https://www.virustotal.com/api/v3/domains/{domain}"
                    r = requests.get(fallback_url, headers=headers, timeout=3, verify=ssl_verify)
                    print(f"[VT] Fallback domain lookup: {fallback_url} -> {r.status_code}")
            except Exception as fallback_err:
                print(f"[VT] Fallback domain extraction failed: {fallback_err}")
                pass  # If fallback fails, keep original error
        
        if r.status_code != 200:
            snippet = r.text[:400].replace("\n", " ")
            return None, f"HTTP {r.status_code}: {snippet}"
        try:
            js = r.json()
        except Exception:
            return None, "Invalid JSON from VirusTotal"
        print(f"[VT] sample: {str(js)[:200]}")
        return summarize_vt(js), None
    except requests.exceptions.SSLError as e:
        error_msg = str(e)
        # If SSL verification is already disabled, don't retry
        if not ssl_verify:
            return None, f"SSL Error (verification disabled): {e}"
        
        # If SSL error occurs and we're using verification, try without verification as fallback
        if "certificate verify failed" in error_msg or "self-signed certificate" in error_msg or "CERTIFICATE_VERIFY_FAILED" in error_msg:
            print(f"[VT] SSL verification failed, retrying without verification (corporate proxy detected)")
            try:
                urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
                r = requests.get(url, headers=headers, timeout=3, verify=False)
                print(f"[VT] {url} -> {r.status_code} (SSL verification disabled)")
                
                # If URL lookup returns 404, fallback to domain check (even in SSL retry)
                if r.status_code == 404 and "/api/v3/urls/" in url:
                    print(f"[VT] URL not found, falling back to Domain check for: {query}")
                    # Extract domain
                    try:
                        parsed = urlparse(query)
                        domain = parsed.netloc
                        if ':' in domain:
                            domain = domain.split(':')[0]
                        
                        if domain:
                            fallback_url = f"https://www.virustotal.com/api/v3/domains/{domain}"
                            r = requests.get(fallback_url, headers=headers, timeout=10, verify=False)
                            print(f"[VT] Fallback domain lookup: {fallback_url} -> {r.status_code}")
                    except Exception as fallback_err:
                        print(f"[VT] Fallback domain extraction failed: {fallback_err}")
                        pass  # If fallback fails, keep original error
                
                if r.status_code != 200:
                    snippet = r.text[:400].replace("\n", " ")
                    return None, f"HTTP {r.status_code}: {snippet}"
                try:
                    js = r.json()
                except Exception:
                    return None, "Invalid JSON from VirusTotal"
                print(f"[VT] sample: {str(js)[:200]}")
                return summarize_vt(js), None
            except Exception as retry_e:
                return None, f"SSL Error (retry failed): {retry_e}"
        return None, f"SSL Error: {e}"
    except Exception as e:
        return None, f"Network error: {e}"

def vt_fetch(query: str):
    """VT fetch with smart caching."""
    return get_cached_data(query, _vt_fetch_internal, 'vt', max_age_hours=24)

def _vt_fetch_relationships_internal(indicator: str, indicator_type: str):
    """
    Internal function to fetch related entities from VirusTotal.
    Returns relationships dict with categories and related entities.
    """
    if not VT_API_KEY:
        return None, "VT_API_KEY is not configured."
    
    relationships = {
        "communicating_files": [],
        "referrer_files": [],
        "subdomains": [],
        "contacted_urls": [],
        "contacted_ips": []
    }
    
    headers = {"x-apikey": VT_API_KEY, "Accept": "application/json"}
    ssl_verify = get_ssl_verify()
    
    try:
        if indicator_type == "ip":
            # Fetch communicating files
            url1 = f"https://www.virustotal.com/api/v3/ip_addresses/{indicator}/communicating_files"
            try:
                r1 = requests.get(url1, headers=headers, timeout=3, verify=ssl_verify)
                if r1.status_code == 200:
                    data1 = r1.json()
                    files = data1.get("data", [])
                    # Sort by malicious count, limit to top 5
                    files_sorted = sorted(files, key=lambda x: x.get("attributes", {}).get("last_analysis_stats", {}).get("malicious", 0), reverse=True)[:5]
                    relationships["communicating_files"] = files_sorted
            except Exception as e:
                print(f"[VT Relationships] Error fetching communicating_files for {indicator}: {e}")
            
            # Fetch referrer files
            url2 = f"https://www.virustotal.com/api/v3/ip_addresses/{indicator}/referrer_files"
            try:
                r2 = requests.get(url2, headers=headers, timeout=3, verify=ssl_verify)
                if r2.status_code == 200:
                    data2 = r2.json()
                    files = data2.get("data", [])
                    files_sorted = sorted(files, key=lambda x: x.get("attributes", {}).get("last_analysis_stats", {}).get("malicious", 0), reverse=True)[:5]
                    relationships["referrer_files"] = files_sorted
            except Exception as e:
                print(f"[VT Relationships] Error fetching referrer_files for {indicator}: {e}")
        
        elif indicator_type == "domain":
            # Fetch communicating files
            url1 = f"https://www.virustotal.com/api/v3/domains/{indicator}/communicating_files"
            try:
                r1 = requests.get(url1, headers=headers, timeout=3, verify=ssl_verify)
                if r1.status_code == 200:
                    data1 = r1.json()
                    files = data1.get("data", [])
                    files_sorted = sorted(files, key=lambda x: x.get("attributes", {}).get("last_analysis_stats", {}).get("malicious", 0), reverse=True)[:5]
                    relationships["communicating_files"] = files_sorted
            except Exception as e:
                print(f"[VT Relationships] Error fetching communicating_files for {indicator}: {e}")
            
            # Fetch subdomains
            url2 = f"https://www.virustotal.com/api/v3/domains/{indicator}/subdomains"
            try:
                r2 = requests.get(url2, headers=headers, timeout=3, verify=ssl_verify)
                if r2.status_code == 200:
                    data2 = r2.json()
                    subdomains = data2.get("data", [])
                    # Limit to top 5 subdomains
                    relationships["subdomains"] = subdomains[:5]
            except Exception as e:
                print(f"[VT Relationships] Error fetching subdomains for {indicator}: {e}")
        
        elif indicator_type == "hash":
            # Fetch contacted URLs
            url1 = f"https://www.virustotal.com/api/v3/files/{indicator}/contacted_urls"
            try:
                r1 = requests.get(url1, headers=headers, timeout=3, verify=ssl_verify)
                if r1.status_code == 200:
                    data1 = r1.json()
                    urls = data1.get("data", [])
                    # Sort by malicious count if available, limit to top 5
                    relationships["contacted_urls"] = urls[:5]
            except Exception as e:
                print(f"[VT Relationships] Error fetching contacted_urls for {indicator}: {e}")
            
            # Fetch contacted IPs
            url2 = f"https://www.virustotal.com/api/v3/files/{indicator}/contacted_ips"
            try:
                r2 = requests.get(url2, headers=headers, timeout=3, verify=ssl_verify)
                if r2.status_code == 200:
                    data2 = r2.json()
                    ips = data2.get("data", [])
                    relationships["contacted_ips"] = ips[:5]
            except Exception as e:
                print(f"[VT Relationships] Error fetching contacted_ips for {indicator}: {e}")
        
        return relationships, None
        
    except Exception as e:
        return None, f"Error fetching relationships: {e}"

def vt_fetch_relationships(indicator: str, indicator_type: str):
    """
    Fetch related entities from VirusTotal with caching.
    Returns relationships dict and error.
    """
    # Use cache key that includes type
    cache_key = f"{indicator}|||relationships|||{indicator_type}"
    # Create a wrapper function that matches get_cached_data signature
    def fetch_wrapper(key):
        return _vt_fetch_relationships_internal(indicator, indicator_type)
    
    return get_cached_data(cache_key, fetch_wrapper, 'vt_relationships', max_age_hours=24)

def _get_abuseipdb_info_internal(ip: str):
    """Internal AbuseIPDB fetch function (without caching)."""
    if not ABUSEIPDB_API_KEY:
        return None, "ABUSEIPDB_API_KEY is not configured."
    url = "https://api.abuseipdb.com/api/v2/check"
    headers = {"Key": ABUSEIPDB_API_KEY, "Accept": "application/json"}
    params = {"ipAddress": ip, "maxAgeInDays": "90"}
    
    # First try with SSL verification
    ssl_verify = get_ssl_verify()
    try:
        r = requests.get(url, headers=headers, params=params, timeout=3, verify=ssl_verify)
        print(f"[IPDB] {ip} -> {r.status_code}")
        if r.status_code != 200:
            snippet = r.text[:400].replace("\n", " ")
            return None, f"HTTP {r.status_code}: {snippet}"
        try:
            d = r.json().get("data", {}) or {}
        except Exception:
            return None, "Invalid JSON from AbuseIPDB"
        print(f"[IPDB] sample: {str(d)[:200]}")
        return {
            "abuseScore": d.get("abuseConfidenceScore", "None"),
            "totalReports": d.get("totalReports", "None"),
            "isp": d.get("isp", "None"),
            "country": d.get("countryCode", "None"),
            "domain": d.get("domain", "None"),
            "flag_url": (f"https://flagcdn.com/w40/{(d.get('countryCode') or '').lower()}.png" if d.get("countryCode") else None),
            "reports": (d.get("reports", []) or [])[:5]
        }, None
    except requests.exceptions.SSLError as e:
        error_msg = str(e)
        # If SSL verification is already disabled, don't retry
        if not ssl_verify:
            return None, f"SSL Error (verification disabled): {e}"
        
        # If SSL error occurs and we're using verification, try without verification as fallback
        if "certificate verify failed" in error_msg or "self-signed certificate" in error_msg or "CERTIFICATE_VERIFY_FAILED" in error_msg:
            print(f"[IPDB] SSL verification failed, retrying without verification (corporate proxy detected)")
            try:
                urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
                r = requests.get(url, headers=headers, params=params, timeout=3, verify=False)
                print(f"[IPDB] {ip} -> {r.status_code} (SSL verification disabled)")
                if r.status_code != 200:
                    snippet = r.text[:400].replace("\n", " ")
                    return None, f"HTTP {r.status_code}: {snippet}"
                try:
                    d = r.json().get("data", {}) or {}
                except Exception:
                    return None, "Invalid JSON from AbuseIPDB"
                print(f"[IPDB] sample: {str(d)[:200]}")
                return {
                    "abuseScore": d.get("abuseConfidenceScore", "None"),
                    "totalReports": d.get("totalReports", "None"),
                    "isp": d.get("isp", "None"),
                    "country": d.get("countryCode", "None"),
                    "domain": d.get("domain", "None"),
                    "flag_url": (f"https://flagcdn.com/w40/{(d.get('countryCode') or '').lower()}.png" if d.get("countryCode") else None),
                    "reports": (d.get("reports", []) or [])[:5]
                }, None
            except Exception as retry_e:
                return None, f"SSL Error (retry failed): {retry_e}"
        return None, f"SSL Error: {e}"
    except Exception as e:
        return None, f"Network error: {e}"

def get_abuseipdb_info(ip: str):
    """AbuseIPDB fetch with smart caching."""
    return get_cached_data(ip, _get_abuseipdb_info_internal, 'abuseipdb', max_age_hours=24)

# ---------- routes ----------
@app.route("/", methods=["GET"])
def home():
    news_data = get_cyber_news_cached()
    # Get latest VirusTotal comments (cached for 1 hour)
    vt_comments, comments_error = get_vt_comments_cached(limit=10)
    recent_history = get_recent_history(limit=10)
    missing = []
    if not VT_API_KEY: missing.append("VT_API_KEY")
    if not ABUSEIPDB_API_KEY: missing.append("ABUSEIPDB_API_KEY")
    
    result = ("Missing API keys. Set them in your .env file: " + ", ".join(missing)) if missing else None
    return render_template("index.html", result=result, query="", vt=None, ipdb=None,
                           malicious=0, suspicious=0, harmless=0, undetected=0, timeout=0,
                           news=news_data, recent_history=recent_history,
                           vt_comments=vt_comments, comments_error=comments_error)

def active_recon_domain(domain: str) -> dict:
    """
    Perform active reconnaissance on a domain using WHOIS and DNS.
    
    Args:
        domain: Domain name to investigate
        
    Returns:
        Dictionary containing:
        - domain: The queried domain
        - whois_data: Parsed WHOIS information
        - domain_age_days: Age of domain in days (None if unavailable)
        - mx_records: List of MX records
        - risk_flags: List of risk indicators
        - is_high_risk: Boolean indicating if domain is high risk
        - error: Error message if any
    """
    result = {
        "domain": domain,
        "whois_data": {},
        "domain_age_days": None,
        "mx_records": [],
        "risk_flags": [],
        "is_high_risk": False,
        "error": None
    }
    
    try:
        # Clean domain (remove protocol, path, etc.)
        clean_domain = domain.lower().strip()
        if '://' in clean_domain:
            from urllib.parse import urlparse
            parsed = urlparse(clean_domain)
            clean_domain = parsed.netloc or parsed.path.split('/')[0]
        # Remove port if present
        if ':' in clean_domain:
            clean_domain = clean_domain.split(':')[0]
        # Remove www. prefix
        if clean_domain.startswith('www.'):
            clean_domain = clean_domain[4:]
        
        result["domain"] = clean_domain
        
        # WHOIS Lookup
        if WHOIS_AVAILABLE:
            try:
                print(f"[Active Recon] Fetching WHOIS data for {clean_domain}")
                whois_result = whois.whois(clean_domain)
                
                if whois_result:
                    # Extract key information
                    whois_info = {}
                    
                    # Creation date (most important for domain age)
                    creation_date = None
                    if hasattr(whois_result, 'creation_date'):
                        creation_date = whois_result.creation_date
                    elif 'creation_date' in whois_result:
                        creation_date = whois_result['creation_date']
                    
                    if creation_date:
                        # Handle list of dates (some WHOIS returns multiple)
                        if isinstance(creation_date, list):
                            creation_date = creation_date[0]
                        
                        if isinstance(creation_date, str):
                            # Try to parse string date
                            try:
                                # Try common date formats
                                date_formats = [
                                    '%Y-%m-%d',
                                    '%Y-%m-%d %H:%M:%S',
                                    '%d-%m-%Y',
                                    '%m/%d/%Y',
                                    '%Y.%m.%d'
                                ]
                                parsed = False
                                for fmt in date_formats:
                                    try:
                                        creation_date = datetime.strptime(creation_date, fmt)
                                        parsed = True
                                        break
                                    except:
                                        continue
                                if not parsed:
                                    # Try dateutil parser as fallback
                                    try:
                                        from dateutil import parser
                                        creation_date = parser.parse(creation_date)
                                    except:
                                        pass
                            except:
                                pass
                        
                        if isinstance(creation_date, datetime):
                            whois_info['creation_date'] = creation_date.isoformat()
                            # Calculate domain age
                            # Handle timezone-aware vs timezone-naive datetime
                            now = datetime.now()
                            
                            # If creation_date is timezone-aware, make now timezone-aware too
                            if creation_date.tzinfo is not None:
                                # creation_date has timezone info
                                try:
                                    from datetime import timezone
                                    # Make now timezone-aware (UTC)
                                    now = datetime.now(timezone.utc)
                                except:
                                    # Fallback: remove timezone from creation_date
                                    creation_date = creation_date.replace(tzinfo=None)
                            else:
                                # creation_date is naive, ensure now is also naive
                                if now.tzinfo is not None:
                                    now = now.replace(tzinfo=None)
                            
                            try:
                                age_delta = now - creation_date
                                domain_age_days = age_delta.days
                                result["domain_age_days"] = domain_age_days
                                
                                # Check if domain is less than 30 days old (HIGH RISK)
                                if domain_age_days < 30:
                                    result["risk_flags"].append(f"Domain is only {domain_age_days} days old (< 30 days)")
                                    result["is_high_risk"] = True
                            except Exception as age_error:
                                print(f"[Active Recon] ⚠️ Error calculating domain age: {age_error}")
                                # Still store the creation date even if age calculation failed
                                result["error"] = f"Could not calculate domain age: {str(age_error)}"
                        else:
                            whois_info['creation_date'] = str(creation_date) if creation_date else None
                    
                    # Expiration date
                    exp_date = None
                    if hasattr(whois_result, 'expiration_date'):
                        exp_date = whois_result.expiration_date
                    elif 'expiration_date' in whois_result:
                        exp_date = whois_result['expiration_date']
                    
                    if exp_date:
                        if isinstance(exp_date, list):
                            exp_date = exp_date[0]
                        whois_info['expiration_date'] = str(exp_date) if exp_date else None
                    
                    # Registrar
                    registrar = None
                    if hasattr(whois_result, 'registrar'):
                        registrar = whois_result.registrar
                    elif 'registrar' in whois_result:
                        registrar = whois_result['registrar']
                    whois_info['registrar'] = registrar
                    
                    # Name servers
                    name_servers = []
                    if hasattr(whois_result, 'name_servers'):
                        ns = whois_result.name_servers
                    elif 'name_servers' in whois_result:
                        ns = whois_result['name_servers']
                    else:
                        ns = None
                    
                    if ns:
                        if isinstance(ns, list):
                            name_servers = [str(n) for n in ns]
                        else:
                            name_servers = [str(ns)]
                    whois_info['name_servers'] = name_servers
                    
                    # Organization/Registrant
                    org = None
                    if hasattr(whois_result, 'org'):
                        org = whois_result.org
                    elif 'org' in whois_result:
                        org = whois_result['org']
                    elif hasattr(whois_result, 'registrant'):
                        org = whois_result.registrant
                    whois_info['organization'] = org
                    
                    result["whois_data"] = whois_info
                    print(f"[Active Recon] ✓ WHOIS data retrieved for {clean_domain}")
                else:
                    result["error"] = "No WHOIS data found"
            except Exception as whois_error:
                print(f"[Active Recon] ⚠️ WHOIS error for {clean_domain}: {whois_error}")
                result["error"] = f"WHOIS lookup failed: {str(whois_error)}"
        else:
            result["error"] = "WHOIS library not available"
        
        # DNS MX Records Lookup
        if DNS_AVAILABLE:
            try:
                print(f"[Active Recon] Fetching MX records for {clean_domain}")
                mx_records = []
                try:
                    answers = dns.resolver.resolve(clean_domain, 'MX')
                    for rdata in answers:
                        mx_records.append({
                            "priority": rdata.preference,
                            "exchange": str(rdata.exchange).rstrip('.')
                        })
                    # Sort by priority
                    mx_records.sort(key=lambda x: x["priority"])
                    result["mx_records"] = mx_records
                    print(f"[Active Recon] ✓ Found {len(mx_records)} MX record(s) for {clean_domain}")
                except dns.resolver.NXDOMAIN:
                    result["risk_flags"].append("Domain does not exist (NXDOMAIN)")
                    result["is_high_risk"] = True
                except dns.resolver.NoAnswer:
                    result["risk_flags"].append("No MX records found - Suspicious configuration")
                    result["is_high_risk"] = True
                except dns.resolver.Timeout:
                    result["risk_flags"].append("DNS query timeout")
                except Exception as dns_error:
                    result["error"] = f"DNS lookup failed: {str(dns_error)}"
            except Exception as dns_error:
                print(f"[Active Recon] ⚠️ DNS error for {clean_domain}: {dns_error}")
                if not result["error"]:
                    result["error"] = f"DNS lookup failed: {str(dns_error)}"
        else:
            if not result["error"]:
                result["error"] = "DNS library not available"
            else:
                result["error"] += "; DNS library not available"
        
    except Exception as e:
        print(f"[Active Recon] ⚠️ Error in active_recon_domain for {domain}: {e}")
        traceback.print_exc()
        result["error"] = str(e)
    
    return result

@app.route("/api/osint/dossier", methods=["POST"])
def get_dossier_api():
    """
    API endpoint to fetch OSINT dossier (Active Recon) data for a domain.
    
    POST body:
    {
        "domain": "example.com"
    }
    
    Returns:
    {
        "domain": "example.com",
        "whois_data": {...},
        "domain_age_days": 1234,
        "mx_records": [...],
        "risk_flags": [...],
        "is_high_risk": false,
        "error": null
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        domain = data.get("domain")
        if not domain:
            return jsonify({"error": "Domain parameter is required"}), 400
        
        # Perform active reconnaissance
        dossier = active_recon_domain(domain)
        
        # Get SSL relations count for recommendation engine
        ssl_relations = get_ssl_relations(domain)
        ssl_count = len(ssl_relations) if ssl_relations else 0
        
        # Determine if SSL pattern is suspicious (3+ unrelated base domains)
        ssl_suspicious = False
        if ssl_count > 0:
            main_domain_parts = domain.lower().split('.')
            main_base = main_domain_parts[-2:] if len(main_domain_parts) >= 2 else []
            main_base_name = '.'.join(main_base) if main_base else ''
            
            related_base_names = set()
            for rel_domain in ssl_relations:
                parts = rel_domain.lower().split('.')
                if len(parts) >= 2:
                    base_name = '.'.join(parts[-2:])
                    related_base_names.add(base_name)
            
            unrelated_count = sum(1 for base in related_base_names 
                                if base != main_base_name and base != '')
            ssl_suspicious = unrelated_count >= 3
        
        # Add SSL data to dossier for recommendations
        dossier["ssl_relations_count"] = ssl_count
        dossier["ssl_suspicious_pattern"] = ssl_suspicious
        
        # Fetch VirusTotal data for recommendations
        vt_data = None
        try:
            vt_summary, vt_err = vt_fetch(domain)
            if vt_summary and not vt_err:
                # Extract relevant VT data for recommendations
                vt_data = {
                    "malicious": vt_summary.get("malicious", 0),
                    "suspicious": vt_summary.get("suspicious", 0),
                    "harmless": vt_summary.get("harmless", 0),
                    "community_score": vt_summary.get("community_score", 0),
                    "kind": vt_summary.get("kind", "—")
                }
                dossier["vt_data"] = vt_data
                print(f"[Dossier] VirusTotal data: {vt_data.get('malicious')} malicious, {vt_data.get('suspicious')} suspicious")
        except Exception as vt_err:
            print(f"[Dossier] Error fetching VirusTotal data: {vt_err}")
            # Continue without VT data
        
        # Fetch AbuseIPDB data if domain is actually an IP
        # Note: AbuseIPDB is for IPs, but we check anyway in case user passes IP as domain
        abuseipdb_data = None
        try:
            # Check if domain is actually an IP address
            import re
            ip_pattern = re.compile(r'^(\d{1,3}\.){3}\d{1,3}$')
            if ip_pattern.match(domain):
                ipdb_view, ipdb_err = get_abuseipdb_info(domain)
                if ipdb_view and not ipdb_err:
                    abuseipdb_data = {
                        "abuseScore": ipdb_view.get("abuseScore", "None"),
                        "totalReports": ipdb_view.get("totalReports", "None"),
                        "isp": ipdb_view.get("isp", "None"),
                        "country": ipdb_view.get("country", "None")
                    }
                    dossier["abuseipdb_data"] = abuseipdb_data
                    print(f"[Dossier] AbuseIPDB data: Score {abuseipdb_data.get('abuseScore')}, Reports {abuseipdb_data.get('totalReports')}")
        except Exception as ipdb_err:
            print(f"[Dossier] Error fetching AbuseIPDB data: {ipdb_err}")
            # Continue without AbuseIPDB data
        
        # Generate AI recommendations (now includes VT and AbuseIPDB data)
        recommendations = generate_security_recommendations(dossier)
        dossier["recommendations"] = recommendations
        
        # Extract features for training data
        features = {
            "domain_age_days": dossier.get("domain_age_days"),
            "is_high_risk": dossier.get("is_high_risk", False),
            "risk_flags": dossier.get("risk_flags", []),
            "mx_records_count": len(dossier.get("mx_records", [])),
            "ssl_relations_count": ssl_count,
            "ssl_suspicious_pattern": ssl_suspicious,
            "has_whois_data": bool(dossier.get("whois_data")),
            "vt_malicious": vt_data.get("malicious", 0) if vt_data else 0,
            "vt_suspicious": vt_data.get("suspicious", 0) if vt_data else 0,
            "vt_community_score": vt_data.get("community_score", 0) if vt_data else 0,
            "abuse_score": abuseipdb_data.get("abuseScore", "None") if abuseipdb_data else "None",
            "total_reports": abuseipdb_data.get("totalReports", "None") if abuseipdb_data else "None"
        }
        
        # Determine verdict from recommendations
        verdict = "info"
        if any(r.get("severity") == "high" for r in recommendations):
            verdict = "high"
        elif any(r.get("severity") == "medium" for r in recommendations):
            verdict = "medium"
        elif any(r.get("severity") == "low" for r in recommendations):
            verdict = "low"
        
        # Save analysis vector to AI training data table (for future ML training)
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            current_time = time.time()
            
            cursor.execute('''
                INSERT INTO ai_training_data (domain, features, verdict, created_at)
                VALUES (?, ?, ?, ?)
            ''', (domain, json.dumps(features), verdict, current_time))
            
            conn.commit()
            conn.close()
            print(f"[AI Training] Saved analysis features for domain: {domain}")
        except Exception as db_err:
            print(f"[AI Training] Error saving to database: {db_err}")
            # Don't fail the request if DB save fails
        
        return jsonify(dossier), 200
        
    except Exception as e:
        print(f"[Dossier API] Error: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


def generate_security_recommendations(analysis_data: dict) -> list:
    """
    Generate security recommendations based on analysis data.
    Uses rule-based logic to provide actionable advice.
    
    Args:
        analysis_data: Dictionary containing domain analysis results including:
            - domain_age_days: Age of domain in days
            - is_high_risk: Boolean indicating high risk
            - risk_flags: List of risk indicators
            - mx_records: List of MX records
            - ssl_relations_count: Number of SSL-related domains (optional)
            - ssl_suspicious_pattern: Boolean indicating suspicious SSL pattern (optional)
            - vt_data: VirusTotal data (optional) with malicious, suspicious, community_score
            - abuseipdb_data: AbuseIPDB data (optional) with abuseScore, totalReports (for IPs only)
    
    Returns:
        List of recommendation dictionaries with severity, action, and reason
    """
    recommendations = []
    
    domain = analysis_data.get("domain", "unknown")
    domain_age = analysis_data.get("domain_age_days")
    is_high_risk = analysis_data.get("is_high_risk", False)
    risk_flags = analysis_data.get("risk_flags", [])
    ssl_relations_count = analysis_data.get("ssl_relations_count", 0)
    ssl_suspicious_pattern = analysis_data.get("ssl_suspicious_pattern", False)
    
    # VirusTotal data
    vt_data = analysis_data.get("vt_data", {})
    vt_malicious = vt_data.get("malicious", 0) if isinstance(vt_data, dict) else 0
    vt_suspicious = vt_data.get("suspicious", 0) if isinstance(vt_data, dict) else 0
    vt_community_score = vt_data.get("community_score", 0) if isinstance(vt_data, dict) else 0
    
    # AbuseIPDB data (only for IPs, but we check anyway)
    abuseipdb_data = analysis_data.get("abuseipdb_data", {})
    abuse_score = abuseipdb_data.get("abuseScore", "None") if isinstance(abuseipdb_data, dict) else "None"
    total_reports = abuseipdb_data.get("totalReports", "None") if isinstance(abuseipdb_data, dict) else "None"
    
    # High Risk - Immediate Action
    if is_high_risk:
        recommendations.append({
            "severity": "high",
            "action": "Immediate Action: Block domain on Firewall & EDR.",
            "reason": "Domain flagged as high risk based on multiple indicators."
        })
    
    # Domain Age < 30 days - Suspicious
    if domain_age is not None and domain_age < 30:
        recommendations.append({
            "severity": "high",
            "action": "Investigate: Domain is less than 30 days old. Common in phishing campaigns.",
            "reason": f"Domain age: {domain_age} days"
        })
    elif domain_age is not None and domain_age < 365:
        recommendations.append({
            "severity": "medium",
            "action": "Monitor: Domain is less than 1 year old. Keep under observation.",
            "reason": f"Domain age: {domain_age} days"
        })
    
    # Missing MX records - Suspicious
    mx_records = analysis_data.get("mx_records", [])
    if not mx_records or len(mx_records) == 0:
        recommendations.append({
            "severity": "medium",
            "action": "Alert: No MX records found. Domain may not be used for email or is misconfigured.",
            "reason": "Missing mail server configuration"
        })
    
    # SSL Shared Infrastructure - Hunt
    if ssl_relations_count > 0:
        if ssl_suspicious_pattern:
            recommendations.append({
                "severity": "high",
                "action": "Hunt: Search logs for other domains sharing this certificate. Suspicious pattern detected.",
                "reason": f"{ssl_relations_count} unrelated domains share SSL certificate"
            })
        else:
            recommendations.append({
                "severity": "medium",
                "action": "Hunt: Search logs for other domains on this certificate.",
                "reason": f"{ssl_relations_count} related domains share SSL certificate"
            })
    
    # Risk Flags
    if "young_domain" in risk_flags:
        recommendations.append({
            "severity": "medium",
            "action": "Verify: Check domain registration details and historical usage.",
            "reason": "Domain age risk flag"
        })
    
    if "no_mx_records" in risk_flags:
        recommendations.append({
            "severity": "low",
            "action": "Note: Domain does not accept email. Verify if this is expected.",
            "reason": "Missing MX records"
        })
    
    # VirusTotal Recommendations
    if vt_malicious > 0:
        if vt_malicious >= 5:
            recommendations.append({
                "severity": "high",
                "action": "CRITICAL: Domain flagged as malicious by multiple security engines. Block immediately.",
                "reason": f"VirusTotal: {vt_malicious} security engines detected malicious activity"
            })
        elif vt_malicious >= 2:
            recommendations.append({
                "severity": "high",
                "action": "URGENT: Domain flagged as malicious. Review and consider blocking.",
                "reason": f"VirusTotal: {vt_malicious} security engines detected malicious activity"
            })
        else:
            recommendations.append({
                "severity": "medium",
                "action": "Alert: Domain flagged by at least one security engine. Investigate further.",
                "reason": f"VirusTotal: {vt_malicious} security engine(s) detected malicious activity"
            })
    
    if vt_suspicious > 0:
        recommendations.append({
            "severity": "medium",
            "action": "Monitor: Domain shows suspicious indicators. Keep under observation.",
            "reason": f"VirusTotal: {vt_suspicious} security engine(s) flagged as suspicious"
        })
    
    if vt_community_score < -5:
        recommendations.append({
            "severity": "high",
            "action": "Warning: Strong negative community sentiment. Domain likely malicious.",
            "reason": f"VirusTotal community score: {vt_community_score} (negative indicates malicious)"
        })
    elif vt_community_score < 0:
        recommendations.append({
            "severity": "medium",
            "action": "Caution: Negative community sentiment detected. Review domain reputation.",
            "reason": f"VirusTotal community score: {vt_community_score}"
        })
    
    # AbuseIPDB Recommendations (only if data exists - typically for IPs)
    if abuse_score != "None" and isinstance(abuse_score, (int, float)):
        try:
            abuse_score_num = int(abuse_score) if isinstance(abuse_score, str) else abuse_score
            if abuse_score_num >= 75:
                recommendations.append({
                    "severity": "high",
                    "action": "BLOCK: IP has high abuse confidence score. Immediate blocking recommended.",
                    "reason": f"AbuseIPDB: Abuse confidence {abuse_score_num}% (threshold: 75%)"
                })
            elif abuse_score_num >= 50:
                recommendations.append({
                    "severity": "medium",
                    "action": "Alert: IP shows moderate abuse indicators. Review and consider blocking.",
                    "reason": f"AbuseIPDB: Abuse confidence {abuse_score_num}%"
                })
        except (ValueError, TypeError):
            pass
    
    if total_reports != "None" and isinstance(total_reports, (int, str)):
        try:
            reports_num = int(total_reports) if isinstance(total_reports, str) else total_reports
            if reports_num >= 10:
                recommendations.append({
                    "severity": "high",
                    "action": "URGENT: IP has multiple abuse reports. Block on firewall immediately.",
                    "reason": f"AbuseIPDB: {reports_num} abuse reports in last 90 days"
                })
            elif reports_num >= 3:
                recommendations.append({
                    "severity": "medium",
                    "action": "Investigate: IP has several abuse reports. Review activity logs.",
                    "reason": f"AbuseIPDB: {reports_num} abuse reports in last 90 days"
                })
        except (ValueError, TypeError):
            pass
    
    # If no specific risks, provide monitoring recommendation
    if not recommendations:
        recommendations.append({
            "severity": "info",
            "action": "Monitor: No immediate action required, but log access for baseline analysis.",
            "reason": "Domain appears clean based on current analysis"
        })
    
    return recommendations


def get_ssl_relations(domain: str) -> list:
    """
    Get related domains from Certificate Transparency logs using crt.sh.
    
    This function queries crt.sh to find all SSL certificates that contain
    the given domain, then extracts all other domains from those certificates.
    This is useful for infrastructure hunting - finding related domains that
    share SSL certificates.
    
    Args:
        domain: Domain name to search for (e.g., "example.com")
        
    Returns:
        List of unique related domains found on the same SSL certificates.
        Returns empty list on error. Limited to top 30 results.
    """
    related_domains = []
    
    # Disable SSL warnings for corporate proxy/firewall environments
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    try:
        # Clean domain (remove protocol, path, etc.)
        clean_domain = domain.lower().strip()
        if '://' in clean_domain:
            parsed = urlparse(clean_domain)
            clean_domain = parsed.netloc or parsed.path.split('/')[0]
        # Remove port if present
        if ':' in clean_domain:
            clean_domain = clean_domain.split(':')[0]
        # Remove www. prefix
        if clean_domain.startswith('www.'):
            clean_domain = clean_domain[4:]
        
        # Build crt.sh API URL
        # The % wildcard searches for certificates containing the domain
        api_url = f"https://crt.sh/?q=%.{clean_domain}&output=json"
        
        print(f"[SSL Relations] Querying crt.sh for domain: {clean_domain}")
        
        # Headers with real browser User-Agent (required by crt.sh to bypass bot protection)
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'application/json',
            'Accept-Language': 'en-US,en;q=0.9'
        }
        
        # Make GET request with 30 second timeout (crt.sh API is often slow)
        # verify=False to bypass SSL verification for corporate proxy/firewall environments
        response = requests.get(api_url, headers=headers, timeout=30, verify=False)
        
        # Check if request was successful
        if response.status_code != 200:
            print(f"[SSL Relations] HTTP error for domain {domain}: Status code {response.status_code}")
            print(f"[SSL Relations] Response content (first 500 chars): {response.text[:500]}")
            return []
        
        # Parse JSON response
        try:
            certificates = response.json()
        except json.JSONDecodeError as json_err:
            print(f"[SSL Relations] JSON decode error for domain {domain}: {json_err}")
            print(f"[SSL Relations] Response status code: {response.status_code}")
            print(f"[SSL Relations] Response content (first 500 chars): {response.text[:500]}")
            return []
        
        if not isinstance(certificates, list):
            print(f"[SSL Relations] Unexpected response format: {type(certificates)}")
            print(f"[SSL Relations] Response type: {type(certificates)}, Value: {str(certificates)[:200]}")
            return []
        
        print(f"[SSL Relations] Found {len(certificates)} certificates")
        
        # Extract name_value fields and collect unique domains
        domain_set = set()
        
        for cert in certificates:
            if not isinstance(cert, dict):
                continue
                
            # Extract name_value field
            name_value = cert.get("name_value", "")
            if not name_value:
                continue
            
            # name_value can be a string with multiple domains (newline-separated)
            # or a single domain - split by newlines to extract all domains
            if isinstance(name_value, str):
                # Split by newlines and also by commas (some certificates use comma separation)
                # First split by newlines, then by commas
                domains = []
                for line in name_value.split('\n'):
                    # Also handle comma-separated values
                    for item in line.split(','):
                        domains.append(item.strip())
            elif isinstance(name_value, list):
                domains = [str(d).strip() for d in name_value]
            else:
                domains = [str(name_value).strip()]
            
            # Process each domain
            for domain_entry in domains:
                if not domain_entry:
                    continue
                
                # Remove wildcard prefixes (*.)
                domain_entry = domain_entry.replace('*.', '').strip()
                
                # Skip if it's just "*" or empty after wildcard removal
                if not domain_entry or domain_entry == '*':
                    continue
                
                # Skip if it's the input domain itself
                if domain_entry.lower() == clean_domain:
                    continue
                
                # Exclude emails (strings containing @)
                if '@' in domain_entry:
                    continue
                
                # Basic domain validation - skip if doesn't look like a domain
                # Should contain at least one dot and be alphanumeric with dots/hyphens
                if '.' not in domain_entry:
                    continue
                
                # Skip if contains invalid characters for a domain
                if not all(c.isalnum() or c in '.-' for c in domain_entry):
                    continue
                
                # Add to set (automatically deduplicates)
                domain_set.add(domain_entry.lower())
        
        # Convert set to sorted list and limit to top 30
        related_domains = sorted(list(domain_set))[:30]
        
        print(f"[SSL Relations] Found {len(domain_set)} unique related domains (returning top {len(related_domains)})")
        
    except requests.exceptions.Timeout as e:
        print(f"[SSL Relations] Request timeout (30s) for domain: {domain}")
        print(f"[SSL Relations] Timeout error details: {e}")
    except requests.exceptions.ConnectionError as e:
        print(f"[SSL Relations] Connection error for domain {domain}: {e}")
        print(f"[SSL Relations] Connection error type: {type(e).__name__}")
    except requests.exceptions.HTTPError as e:
        status_code = e.response.status_code if hasattr(e, 'response') and e.response is not None else 'N/A'
        print(f"[SSL Relations] HTTP error for domain {domain}: {e}")
        print(f"[SSL Relations] HTTP status code: {status_code}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"[SSL Relations] Response content (first 500 chars): {e.response.text[:500]}")
    except json.JSONDecodeError as e:
        print(f"[SSL Relations] JSON decode error for domain {domain}: {e}")
        print(f"[SSL Relations] JSON error at position: {e.pos if hasattr(e, 'pos') else 'N/A'}")
    except Exception as e:
        print(f"[SSL Relations] Unexpected error for domain {domain}: {e}")
        print(f"[SSL Relations] Error type: {type(e).__name__}")
        traceback.print_exc()
    
    return related_domains


@app.route("/api/ghost/ssl", methods=["POST"])
def get_ssl_relations_api():
    """
    API endpoint to fetch related domains from Certificate Transparency logs.
    
    This endpoint uses crt.sh to find domains that share SSL certificates
    with the given domain. Useful for infrastructure hunting.
    
    POST body:
    {
        "domain": "example.com"
    }
    
    Returns:
    {
        "domain": "example.com",
        "related_domains": ["sub1.example.com", "sub2.example.com", ...],
        "count": 5,
        "error": null
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        domain = data.get("domain")
        if not domain:
            return jsonify({"error": "Domain parameter is required"}), 400
        
        # Get related domains from SSL certificates
        related_domains = get_ssl_relations(domain)
        
        return jsonify({
            "domain": domain,
            "related_domains": related_domains,
            "count": len(related_domains),
            "error": None
        }), 200
        
    except Exception as e:
        print(f"[SSL Relations API] Error: {e}")
        traceback.print_exc()
        return jsonify({
            "domain": data.get("domain", ""),
            "related_domains": [],
            "count": 0,
            "error": str(e)
        }), 500


def prepare_safe_html(html_content: str, base_url: str) -> str:
    """
    Prepare safe HTML content by sanitizing and fixing links.
    
    Args:
        html_content: Raw HTML content from the response
        base_url: Base URL for fixing relative links (e.g., "https://example.com")
    
    Returns:
        Sanitized HTML string safe to display in iframe
    """
    if not html_content:
        return ""
    
    try:
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove all <script> tags and their contents
        for script in soup.find_all('script'):
            script.decompose()
        
        # Remove all event handler attributes (onclick, onload, etc.)
        for tag in soup.find_all(True):
            # Remove all attributes that start with 'on'
            attrs_to_remove = [attr for attr in tag.attrs.keys() if attr.lower().startswith('on')]
            for attr in attrs_to_remove:
                del tag.attrs[attr]
        
        # Remove Content-Security-Policy meta tags to allow images in iframe
        if soup.head:
            for meta in soup.head.find_all('meta'):
                http_equiv = meta.get('http-equiv', '').lower()
                if http_equiv == 'content-security-policy':
                    meta.decompose()
            
            # Remove existing base tag if any
            existing_base = soup.head.find('base')
            if existing_base:
                existing_base.decompose()
            
            # Create and inject new base tag at the very top of <head>
            base_tag = soup.new_tag('base', href=base_url)
            # Insert at position 0 (very top of head)
            soup.head.insert(0, base_tag)
        else:
            # If no head tag, create one
            head_tag = soup.new_tag('head')
            base_tag = soup.new_tag('base', href=base_url)
            head_tag.insert(0, base_tag)
            if soup.html:
                # Insert head at the beginning of html
                soup.html.insert(0, head_tag)
            else:
                # If no html tag, wrap everything
                html_tag = soup.new_tag('html')
                html_tag.insert(0, head_tag)
                if soup.body:
                    html_tag.append(soup.body)
                else:
                    html_tag.append(soup)
                soup = BeautifulSoup(str(html_tag), 'html.parser')
        
        return str(soup)
    except Exception as e:
        print(f"[Safe HTML] Error sanitizing HTML: {e}")
        # Fallback: basic sanitization with regex
        import re
        # Remove script tags
        safe_html = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        # Remove Content-Security-Policy meta tags
        safe_html = re.sub(r'<meta[^>]*http-equiv\s*=\s*["\']?content-security-policy["\']?[^>]*>', '', safe_html, flags=re.IGNORECASE)
        # Remove on* attributes
        safe_html = re.sub(r'\s+on\w+\s*=\s*["\'][^"\']*["\']', '', safe_html, flags=re.IGNORECASE)
        # Inject base tag at the very top of head (or create head if needed)
        base_injection = f'<base href="{base_url}">'
        if '<head>' in safe_html:
            # Insert base tag right after <head> tag (very top)
            safe_html = safe_html.replace('<head>', f'<head>{base_injection}', 1)
        elif '<html>' in safe_html:
            safe_html = safe_html.replace('<html>', f'<html><head>{base_injection}</head>', 1)
        else:
            safe_html = f'<head>{base_injection}</head>{safe_html}'
        return safe_html


def detect_cloaking(domain: str) -> dict:
    """
    Detect if a domain returns different content based on User-Agent (cloaking detection).
    This is a common phishing evasion technique where sites show different content
    to mobile users vs security bots.
    
    Args:
        domain: Domain name or full URL to test (e.g., "example.com" or "https://example.com/path")
    
    Returns:
        Dictionary with cloaking detection results including:
        - mobile_result: {status_code, content_size, error, safe_html}
        - desktop_result: {status_code, content_size, error, safe_html}
        - bot_result: {status_code, content_size, error, safe_html}
        - deviation_score: Calculated deviation between mobile/bot responses
        - verdict: "High Probability of Cloaking" or "No Cloaking Detected"
        - is_cloaking: Boolean indicating if cloaking was detected
    """
    results = {
        "domain": domain,
        "mobile_result": {"status_code": None, "content_size": 0, "error": None, "safe_html": ""},
        "desktop_result": {"status_code": None, "content_size": 0, "error": None, "safe_html": ""},
        "bot_result": {"status_code": None, "content_size": 0, "error": None, "safe_html": ""},
        "deviation_score": 0,
        "verdict": "Unknown",
        "is_cloaking": False
    }
    
    # Disable SSL warnings for corporate proxy/firewall environments
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    # Determine if input is a full URL or just a domain
    input_lower = domain.strip()
    
    # Check if it's already a full URL (starts with http:// or https://)
    if input_lower.startswith('http://') or input_lower.startswith('https://'):
        # Use the URL as-is (keep the full path)
        url = input_lower
        print(f"[Cloaking Detection] Using provided URL as-is: {url}")
    else:
        # It's just a domain - construct URL (try https:// first)
        clean_domain = input_lower.lower()
        # Remove port if present
        if ':' in clean_domain and not clean_domain.startswith('http'):
            clean_domain = clean_domain.split(':')[0]
        # Try https first
        url = f"https://{clean_domain}"
        print(f"[Cloaking Detection] Constructed URL from domain: {url}")
    
    # Define 3 distinct User-Agent headers
    user_agents = {
        "mobile": "Mozilla/5.0 (iPhone; CPU iPhone OS 16_6 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.6 Mobile/15E148 Safari/604.1",
        "desktop": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "bot": "Googlebot/2.1 (+http://www.google.com/bot.html)"
    }
    
    def fetch_with_ua(persona: str, ua: str):
        """Helper function to fetch with specific User-Agent using streaming for speed."""
        try:
            headers = {
                "User-Agent": ua,
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.9",
                "Connection": "close"  # Prevent pool blocking behind proxies
            }
            # Use streaming=True for partial download, timeout=30 for reliability
            response = requests.get(url, headers=headers, timeout=30, verify=False, allow_redirects=True, stream=True)
            
            # Download only first 150KB for visual analysis (optimized for proxy)
            content_chunks = []
            total_bytes = 0
            max_bytes = 150000  # 150KB limit
            
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    content_chunks.append(chunk)
                    total_bytes += len(chunk)
                    if total_bytes >= max_bytes:
                        # Only take the first 150KB
                        if total_bytes > max_bytes:
                            # Trim the last chunk if we exceeded
                            excess = total_bytes - max_bytes
                            content_chunks[-1] = content_chunks[-1][:-excess]
                        break
            
            # Combine chunks and decode to text
            content_bytes = b''.join(content_chunks)
            content_text = content_bytes.decode('utf-8', errors='ignore')
            
            # Prepare safe HTML for visual evidence
            safe_html = prepare_safe_html(content_text, url) if response.status_code == 200 else ""
            
            return {
                "status_code": response.status_code,
                "content_size": len(content_bytes),
                "error": None,
                "safe_html": safe_html
            }
        except requests.exceptions.SSLError:
            # Try with http if https fails
            if url.startswith("https://"):
                try:
                    http_url = url.replace("https://", "http://")
                    response = requests.get(http_url, headers=headers, timeout=30, verify=False, allow_redirects=True, stream=True)
                    
                    # Download only first 150KB for visual analysis
                    content_chunks = []
                    total_bytes = 0
                    max_bytes = 150000
                    
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            content_chunks.append(chunk)
                            total_bytes += len(chunk)
                            if total_bytes >= max_bytes:
                                if total_bytes > max_bytes:
                                    excess = total_bytes - max_bytes
                                    content_chunks[-1] = content_chunks[-1][:-excess]
                                break
                    
                    content_bytes = b''.join(content_chunks)
                    content_text = content_bytes.decode('utf-8', errors='ignore')
                    safe_html = prepare_safe_html(content_text, http_url) if response.status_code == 200 else ""
                    
                    return {
                        "status_code": response.status_code,
                        "content_size": len(content_bytes),
                        "error": None,
                        "safe_html": safe_html
                    }
                except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
                    # Timeout/Connection error - mark as -1 (not 0)
                    return {"status_code": None, "content_size": -1, "error": str(e), "safe_html": ""}
                except Exception as e:
                    return {"status_code": None, "content_size": -1, "error": str(e), "safe_html": ""}
            return {"status_code": None, "content_size": -1, "error": "SSL Error", "safe_html": ""}
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            # Timeout/Connection error - mark as -1 (not 0) to distinguish from actual 0-byte response
            return {"status_code": None, "content_size": -1, "error": str(e), "safe_html": ""}
        except Exception as e:
            # Any other exception - mark as -1
            return {"status_code": None, "content_size": -1, "error": str(e), "safe_html": ""}
    
    # Perform 3 parallel GET requests using ThreadPoolExecutor
    clean_domain_for_log = url.split('://')[-1].split('/')[0] if '://' in url else clean_domain
    print(f"[Cloaking Detection] Testing domain: {clean_domain_for_log}")
    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = {
            "mobile": executor.submit(fetch_with_ua, "mobile", user_agents["mobile"]),
            "desktop": executor.submit(fetch_with_ua, "desktop", user_agents["desktop"]),
            "bot": executor.submit(fetch_with_ua, "bot", user_agents["bot"])
        }
        
        # Wait for all requests to complete (timeout slightly longer than request timeout)
        for persona, future in futures.items():
            try:
                result = future.result(timeout=35)  # Slightly longer than request timeout (30s)
                results[f"{persona}_result"] = result
            except Exception as e:
                # Mark as -1 for timeout/error (not 0)
                results[f"{persona}_result"] = {"status_code": None, "content_size": -1, "error": str(e), "safe_html": ""}
    
    # Compare Results and Calculate Deviation Score
    mobile_size = results["mobile_result"]["content_size"]
    desktop_size = results["desktop_result"]["content_size"]
    bot_size = results["bot_result"]["content_size"]
    
    # Handle -1 (error/timeout) - treat as None for comparison
    mobile_size = mobile_size if mobile_size != -1 else None
    desktop_size = desktop_size if desktop_size != -1 else None
    bot_size = bot_size if bot_size != -1 else None
    
    # Calculate deviation: if mobile/desktop is significantly different from bot
    # Only calculate if both requests succeeded (not -1/None)
    if bot_size is not None and bot_size > 0:
        # Calculate percentage difference
        mobile_deviation = 0
        desktop_deviation = 0
        if mobile_size is not None and mobile_size > 0:
            mobile_deviation = abs(mobile_size - bot_size) / bot_size * 100
        if desktop_size is not None and desktop_size > 0:
            desktop_deviation = abs(desktop_size - bot_size) / bot_size * 100
        results["deviation_score"] = max(mobile_deviation, desktop_deviation)
    elif mobile_size is not None and mobile_size > 0 and desktop_size is not None and desktop_size > 0:
        # If bot failed, compare mobile vs desktop (only if both succeeded)
        results["deviation_score"] = abs(mobile_size - desktop_size) / max(mobile_size, desktop_size) * 100
    else:
        results["deviation_score"] = 0
    
    # Determine verdict
    # Refined logic: Only flag Evasion if BOTH requests succeeded (status 200), 
    # BUT the content length difference is > 30% OR status codes differ
    
    is_cloaking = False
    verdict = "No Cloaking Detected"
    has_network_error = False
    
    # Check for network errors (timeouts/connection errors)
    mobile_error = results["mobile_result"].get("error") and results["mobile_result"]["content_size"] == -1
    desktop_error = results["desktop_result"].get("error") and results["desktop_result"]["content_size"] == -1
    bot_error = results["bot_result"].get("error") and results["bot_result"]["content_size"] == -1
    
    if mobile_error or desktop_error or bot_error:
        has_network_error = True
    
    # Only flag cloaking if NO network errors and both requests succeeded
    mobile_status = results["mobile_result"].get("status_code")
    bot_status = results["bot_result"].get("status_code")
    
    if has_network_error:
        # Network error occurred - inconclusive
        is_cloaking = False
        verdict = "Inconclusive - Network Error"
    elif mobile_status == 200 and bot_status == 200:
        # Both requests succeeded (200 OK) - can compare
        if mobile_size is not None and bot_size is not None and bot_size > 0:
            # Check for >30% content difference (Mobile or Desktop vs Bot)
            if mobile_size > bot_size * 1.3:
                is_cloaking = True
                verdict = f"High Probability of Cloaking: Mobile content >30% larger than Bot (Mobile: {mobile_size:,} vs Bot: {bot_size:,} bytes)"
            elif desktop_size is not None and desktop_size > bot_size * 1.3:
                is_cloaking = True
                verdict = f"High Probability of Cloaking: Desktop content >30% larger than Bot (Desktop: {desktop_size:,} vs Bot: {bot_size:,} bytes)"
        # Check for different status codes between desktop and bot (if both succeeded but different codes - suspicious)
        desktop_status = results["desktop_result"].get("status_code")
        if desktop_status and desktop_status != bot_status:
            is_cloaking = True
            verdict = f"High Probability of Cloaking: Different status codes (Desktop: {desktop_status} vs Bot: {bot_status})"
    elif mobile_status == 200 and (bot_status is None or bot_status != 200):
        # Mobile succeeded but bot failed - could be cloaking, but check if it's a network error
        if not bot_error and bot_status is not None:
            # Bot got a different status code (not a network error) - suspicious
            is_cloaking = True
            verdict = f"High Probability of Cloaking: Bot blocked ({bot_status}) while Mobile succeeded (200)"
        else:
            # Bot had network error - inconclusive
            verdict = "Inconclusive - Network Error"
    else:
        desktop_status = results["desktop_result"].get("status_code")
        if desktop_status == 200 and (bot_status is None or bot_status != 200):
            # Desktop succeeded but bot failed
            if not bot_error and bot_status is not None:
                is_cloaking = True
                verdict = f"High Probability of Cloaking: Bot blocked ({bot_status}) while Desktop succeeded (200)"
            else:
                verdict = "Inconclusive - Network Error"
    
    results["is_cloaking"] = is_cloaking
    results["verdict"] = verdict
    results["has_network_error"] = has_network_error
    
    print(f"[Cloaking Detection] Result for {clean_domain_for_log}: {verdict} (Deviation: {results['deviation_score']:.1f}%, Cloaking: {is_cloaking})")
    
    return results


@app.route("/api/ghost/cloaking", methods=["POST"])
def detect_cloaking_api():
    """
    API endpoint to detect cloaking behavior on a domain.
    
    POST body:
    {
        "domain": "example.com"
    }
    
    Returns:
    {
        "domain": "example.com",
        "mobile_result": {...},
        "desktop_result": {...},
        "bot_result": {...},
        "deviation_score": 75.5,
        "verdict": "High Probability of Cloaking",
        "is_cloaking": true
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        domain = data.get("domain")
        if not domain:
            return jsonify({"error": "Domain parameter is required"}), 400
        
        # Perform cloaking detection
        result = detect_cloaking(domain)
        
        return jsonify(result), 200
        
    except Exception as e:
        print(f"[Cloaking Detection API] Error: {e}")
        traceback.print_exc()
        return jsonify({
            "domain": data.get("domain", ""),
            "error": str(e),
            "is_cloaking": False,
            "verdict": "Error during detection"
        }), 500


@app.route("/api/ai/feedback", methods=["POST"])
def ai_feedback_api():
    """
    API endpoint to receive user feedback on AI recommendations.
    This enables lifelong learning - the system learns from user corrections.
    
    POST body:
    {
        "domain": "example.com",
        "feedback": 1  // 1 for accurate/positive, 0 for inaccurate/negative
    }
    
    Returns:
    {
        "success": true,
        "message": "Feedback saved successfully"
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        domain = data.get("domain")
        feedback = data.get("feedback")
        
        if not domain:
            return jsonify({"error": "Domain parameter is required"}), 400
        
        if feedback not in [0, 1]:
            return jsonify({"error": "Feedback must be 0 (inaccurate) or 1 (accurate)"}), 400
        
        # Update the most recent training record for this domain
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        current_time = time.time()
        
        # Find the most recent record for this domain without feedback
        cursor.execute('''
            SELECT id FROM ai_training_data
            WHERE domain = ? AND user_feedback IS NULL
            ORDER BY created_at DESC
            LIMIT 1
        ''', (domain,))
        
        row = cursor.fetchone()
        
        if row:
            record_id = row[0]
            # Update with user feedback
            cursor.execute('''
                UPDATE ai_training_data
                SET user_feedback = ?, updated_at = ?
                WHERE id = ?
            ''', (feedback, current_time, record_id))
            conn.commit()
            conn.close()
            print(f"[AI Feedback] Updated feedback for domain {domain}: {feedback}")
        else:
            # If no record found, create one
            # This might happen if the user gives feedback before analysis completes
            cursor.execute('''
                INSERT INTO ai_training_data (domain, features, verdict, user_feedback, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (domain, json.dumps({}), "unknown", feedback, current_time, current_time))
            conn.commit()
            conn.close()
            print(f"[AI Feedback] Created new feedback record for domain {domain}: {feedback}")
        
        return jsonify({
            "success": True,
            "message": "Thank you! Your feedback helps improve the AI model."
        }), 200
        
    except Exception as e:
        print(f"[AI Feedback] Error: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/api/vt/comments", methods=["GET"])
def get_vt_comments_api():
    """
    API endpoint to fetch VirusTotal comments.
    Query parameters:
    - limit: Number of comments (default: 10)
    - filter: Filter by tag (e.g., 'malware')
    - cursor: Pagination cursor
    """
    try:
        limit = request.args.get('limit', 10, type=int)
        filter_tag = request.args.get('filter', None)
        cursor = request.args.get('cursor', None)
        
        # Validate limit
        if limit < 1 or limit > 100:
            limit = 10
        
        comments_data, error = get_vt_comments_cached(limit=limit, filter_tag=filter_tag, cursor=cursor)
        
        if error:
            return jsonify({"error": error}), 400
        
        return jsonify(comments_data), 200
        
    except Exception as e:
        print(f"[VT Comments API] Error: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# --- Fix: Restore Main Search Function (which was missing) ---
@app.route("/lookup", methods=["POST"])
def lookup():
    news_data = get_cyber_news_cached()
    query = (request.form.get("query") or "").strip()
    cache_only = request.form.get("cache_only") == "1"
    
    if not query:
        recent_history = get_recent_history(limit=10)
        return render_template("index.html", result="No query entered", query="",
                               vt=None, ipdb=None,
                               malicious=0, suspicious=0, harmless=0, undetected=0, timeout=0,
                               news=news_data, recent_history=recent_history)

    # If cache_only flag is set, use cache only (no API calls)
    if cache_only:
        vt_data, ipdb_data = get_cached_lookup_data(query)
        if vt_data:
            vt_summary = vt_data if isinstance(vt_data, dict) else {}
            vt_err = None
            print(f"[Lookup] Using cached data for {query} (cache_only mode)")
        else:
            # Not in cache - return error message
            recent_history = get_recent_history(limit=10)
            return render_template("index.html", 
                                   result=f"Data for '{query}' not found in cache. Please perform a new lookup.",
                                   query=query,
                                   vt=None, ipdb=None,
                                   malicious=0, suspicious=0, harmless=0, undetected=0, timeout=0,
                                   news=news_data, recent_history=recent_history)
    else:
        # Normal lookup - will use cache if available, otherwise make API calls
        vt_summary, vt_err = vt_fetch(query)
        if not vt_summary:
            vt_summary = {} 

    ipdb_view = None
    ipdb_err = None
    if guess_type(query) == "ip":
        if cache_only:
            # Use cached AbuseIPDB data if available
            _, ipdb_data = get_cached_lookup_data(query)
            if ipdb_data:
                ipdb_view = ipdb_data
                ipdb_err = None
            else:
                # Not in cache - use normal lookup
                ipdb_view, ipdb_err = get_abuseipdb_info(query)
        else:
            ipdb_view, ipdb_err = get_abuseipdb_info(query)
        
        if not ipdb_view:
            ipdb_view = {
                "abuseScore": "None", "totalReports": "None", "country": "None",
                "isp": "None", "domain": "None", "flag_url": None, "reports": []
            }

    malicious = vt_summary.get("malicious", 0)
    suspicious = vt_summary.get("suspicious", 0)
    harmless = vt_summary.get("harmless", 0)
    undetected = vt_summary.get("undetected", 0)
    timeout = vt_summary.get("timeout", 0)

    vt_view = {
        "name": safe_text(vt_summary.get("name", "—")),
        "malicious": malicious,
        "suspicious": suspicious,
        "community_score": vt_summary.get("community_score", 0), 
        "error": vt_err,
        "kind": vt_summary.get("kind", "—"),
        "detections": vt_summary.get("detections", [])
    }

    if ipdb_view is not None:
        ipdb_view["error"] = ipdb_err

    # Get recent history AFTER lookups (so current search is included)
    recent_history = get_recent_history(limit=10)
    print(f"[Lookup] Recent history count: {len(recent_history)}")

    return render_template(
        "index.html",
        result=None,
        query=query,
        vt=vt_view,
        ipdb=ipdb_view,
        malicious=malicious,
        suspicious=suspicious,
        harmless=harmless,
        undetected=undetected,
        timeout=timeout,
        news=news_data,
        recent_history=recent_history
    )

@app.route("/lookup_fast", methods=["POST"])
def lookup_fast():
    """Fast lookup using only cached data (no API calls)."""
    query = (request.form.get("query") or (request.json.get("query") if request.is_json else None) or "").strip()
    if not query:
        return jsonify({"error": "No query provided"}), 400
    
    # Get data from cache only
    vt_data, ipdb_data = get_cached_lookup_data(query)
    
    if not vt_data:
        return jsonify({"error": "Data not found in cache"}), 404
    
    # Format VT data
    vt_summary = vt_data if isinstance(vt_data, dict) else {}
    malicious = vt_summary.get("malicious", 0)
    suspicious = vt_summary.get("suspicious", 0)
    harmless = vt_summary.get("harmless", 0)
    undetected = vt_summary.get("undetected", 0)
    timeout = vt_summary.get("timeout", 0)
    
    vt_view = {
        "name": safe_text(vt_summary.get("name", "—")),
        "malicious": malicious,
        "suspicious": suspicious,
        "community_score": vt_summary.get("community_score", 0),
        "error": None,
        "kind": vt_summary.get("kind", "—"),
        "detections": vt_summary.get("detections", [])
    }
    
    # Format AbuseIPDB data
    ipdb_view = None
    if ipdb_data:
        ipdb_view = ipdb_data
        ipdb_view["error"] = None
    
    return jsonify({
        "success": True,
        "query": query,
        "vt": vt_view,
        "ipdb": ipdb_view,
        "malicious": malicious,
        "suspicious": suspicious,
        "harmless": harmless,
        "undetected": undetected,
        "timeout": timeout
    })

# --- End Fix ---


# ======= COMPREHENSIVE SIEM LOG ANALYZER - REFACTORED INTO MODULAR FUNCTIONS =======
def extract_ips_from_text(texts_to_parse: list, original_log_text: str, analysis: dict) -> dict:
    """Extract IP addresses from log text using pre-compiled patterns."""
    seen_ips = set()
    
    # Extract source and destination IPs from all texts using pre-compiled patterns
    for parsing_text in texts_to_parse:
        # Try camelCase pattern first (e.g., "sourceIp="), then regular pattern
        src_match = _SRC_IP_CAMEL_PATTERN.search(parsing_text) or _SRC_IP_PATTERN.search(parsing_text)
        if src_match:
            ip = src_match.group(1)
            if ip and ip not in seen_ips:
                analysis["ips"].append({"type": "Source", "ip": ip, "context": "src"})
                analysis["src_ip"] = ip
                seen_ips.add(ip)
        
        dst_match = _DST_IP_CAMEL_PATTERN.search(parsing_text) or _DST_IP_PATTERN.search(parsing_text)
        if dst_match:
            ip = dst_match.group(1)
            if ip and ip not in seen_ips:
                analysis["ips"].append({"type": "Destination", "ip": ip, "context": "dst"})
                analysis["dst_ip"] = ip
                seen_ips.add(ip)
    
    # Extract all remaining IPs that weren't already found
    all_ips = _IP_REGEX_PATTERN.finditer(original_log_text)
    for match in all_ips:
        ip = match.group(0)
        if (ip and ip not in seen_ips and 
            not ip.startswith("127.") and not ip.startswith("0.") and
            not ip.startswith("0.0.0.") and len(ip.split('.')) == 4):
            analysis["ips"].append({"type": "Network", "ip": ip, "context": "generic"})
            seen_ips.add(ip)
    
    return analysis

def extract_network_info(texts_to_parse: list, analysis: dict) -> dict:
    """Extract network-related information (ports, protocol)."""
    # Extract ports
    if not analysis["src_port"]:
        for text in texts_to_parse:
            src_port_match = _SRC_PORT_PATTERN.search(text)
            if src_port_match:
                analysis["src_port"] = src_port_match.group(1)
                break
    
    if not analysis["dst_port"]:
        for text in texts_to_parse:
            dst_port_match = _DST_PORT_PATTERN.search(text)
            if dst_port_match:
                analysis["dst_port"] = dst_port_match.group(1)
                break
    
    # Extract protocol
    if not analysis["protocol"]:
        for text in texts_to_parse:
            protocol_match = _PROTOCOL_PATTERN.search(text)
            if protocol_match:
                analysis["protocol"] = protocol_match.group(1).upper()
                break
    
    return analysis

def extract_urls_and_domains(texts_to_parse: list, analysis: dict) -> dict:
    """Extract URLs and domains from log text."""
    all_urls = []
    for text in texts_to_parse:
        urls = _URL_REGEX_PATTERN.findall(text)
        all_urls.extend(urls)
    analysis["urls"] = list(set(all_urls))
    
    all_domains = []
    for text in texts_to_parse:
        domains = _DOMAIN_REGEX_PATTERN.findall(text)
        all_domains.extend(domains)
    analysis["domains"] = list(set([d.lower() for d in all_domains if "." in d and len(d) > 3]))[:10]
    
    return analysis

def extract_hashes(texts_to_parse: list, original_log_text: str, json_data: dict, analysis: dict) -> dict:
    """Extract hash values from log text."""
    def identify_hash_type(hash_str: str) -> str:
        length = len(hash_str)
        hash_types = {32: "MD5", 40: "SHA-1", 56: "SHA-224/SHA3-224", 
                      64: "SHA-256/SHA3-256", 96: "SHA-384/SHA3-384", 128: "SHA-512/SHA3-512"}
        return hash_types.get(length, f"Hash ({length} chars)")
    
    seen_hashes = set()
    all_hashes = []
    
    for text in texts_to_parse:
        hash_matches = _HASH_REGEX_PATTERN.finditer(text)
        for match in hash_matches:
            hash_candidate = match.group(0)
            if len(hash_candidate) in [32, 40, 56, 64, 96, 128] and hash_candidate not in seen_hashes:
                start_pos = match.start()
                end_pos = match.end()
                if (start_pos == 0 or not text[start_pos-1:start_pos].isalnum()) and \
                   (end_pos >= len(text) or not text[end_pos:end_pos+1].isalnum()):
                    all_hashes.append({
                        "hash": hash_candidate,
                        "type": identify_hash_type(hash_candidate),
                        "length": len(hash_candidate)
                    })
                    seen_hashes.add(hash_candidate)
    
    # Check JSON fields for hashes
    if json_data:
        try:
            hash_fields = ["hash", "file_hash", "md5", "sha1", "sha256", "sha512", 
                          "sha-1", "sha-256", "sha-512", "sha384", "sha-384",
                          "sha224", "sha-224", "checksum", "digest", "fingerprint"]
            for field in hash_fields:
                if field in json_data and isinstance(json_data[field], str):
                    hash_val = json_data[field].strip()
                    if hash_val and hash_val not in seen_hashes:
                        if len(hash_val) in [32, 40, 56, 64, 96, 128] and re.match(r"^[A-Fa-f0-9]+$", hash_val):
                            all_hashes.append({
                                "hash": hash_val,
                                "type": identify_hash_type(hash_val),
                                "length": len(hash_val)
                            })
                            seen_hashes.add(hash_val)
        except (json.JSONDecodeError, AttributeError, TypeError):
            pass
    
    analysis["hashes"] = all_hashes
    return analysis

def extract_user_info(texts_to_parse: list, analysis: dict) -> dict:
    """Extract user-related information (usernames, emails)."""
    def add_unique(collection, value):
        if value and value not in collection:
            collection.append(value)
    
    # Extract username
    if not analysis["username"]:
        for text in texts_to_parse:
            for pattern in _USERNAME_PATTERNS_COMPILED:
                user_match = pattern.search(text)
                if user_match:
                    username_candidate = user_match.group(1).strip()
                    if username_candidate:
                        analysis["username"] = username_candidate
                        break
            if analysis["username"]:
                break
    
    # Extract all usernames and emails
    for text in texts_to_parse:
        for pattern in _USERNAME_PATTERNS_COMPILED:
            user_match = pattern.search(text)
            if user_match:
                username_candidate = user_match.group(1).strip()
                add_unique(analysis["identifiers"]["usernames"], username_candidate)
        
        for email_match in _EMAIL_REGEX_PATTERN.finditer(text):
            add_unique(analysis["identifiers"]["emails"], email_match.group(0).strip())
        
        for custom_email_match in _MOBILEYE_EMAIL_REGEX_PATTERN.finditer(text):
            username = custom_email_match.group(1).strip()
            email = custom_email_match.group(0).strip()
            if username:
                analysis["user_information"]["username"] = username
                analysis["user_information"]["email"] = email
                add_unique(analysis["identifiers"]["usernames"], username)
                add_unique(analysis["identifiers"]["emails"], email)
                analysis["username"] = username
    
    if not analysis["username"] and analysis["identifiers"]["emails"]:
        analysis["username"] = analysis["identifiers"]["emails"][0]
    
    return analysis

def calculate_risk_score(original_log_text: str, analysis: dict) -> dict:
    """Calculate risk score and level based on log content."""
    risk_score = 0
    for match in _RISK_KEYWORDS_HIGH_PATTERN.finditer(original_log_text):
        keyword = match.group(1).upper()
        risk_score += 3
        analysis["risk_indicators"].append(keyword)
    
    for match in _RISK_KEYWORDS_MEDIUM_PATTERN.finditer(original_log_text):
        keyword = match.group(1).upper()
        risk_score += 1
        analysis["risk_indicators"].append(keyword)
    
    # Check for suspicious ports
    if analysis["dst_port"]:
        suspicious_ports = ["4444", "5555", "6666", "6667", "12345", "31337"]
        if analysis["dst_port"] in suspicious_ports:
            risk_score += 2
            analysis["risk_indicators"].append(f"SUSPICIOUS_PORT:{analysis['dst_port']}")
    
    # Check for suspicious HTTP status codes
    if analysis["http_status"]:
        if analysis["http_status"].startswith("4") or analysis["http_status"].startswith("5"):
            risk_score += 1
    
    # Determine risk level
    if risk_score >= 5:
        analysis["risk_level"] = "high"
    elif risk_score >= 2:
        analysis["risk_level"] = "medium"
    else:
        analysis["risk_level"] = "low"
    
    return analysis

def parse_office365_alert(json_data: dict, analysis: dict) -> dict:
    """
    Parse Office 365 Security & Compliance alert and extract relevant information.
    Handles alerts like email reported as not junk, phishing, malware, etc.
    """
    # Initialize Office 365 specific fields
    if "o365_alert" not in analysis:
        analysis["o365_alert"] = {}
    
    o365_data = analysis["o365_alert"]
    
    # Extract basic alert information
    o365_data["operation"] = json_data.get("Operation")
    o365_data["workload"] = json_data.get("Workload")
    o365_data["alert_type"] = json_data.get("AlertType")
    o365_data["category"] = json_data.get("Category")
    o365_data["severity"] = json_data.get("Severity")
    o365_data["status"] = json_data.get("Status")
    o365_data["alert_name"] = json_data.get("Name")
    o365_data["alert_id"] = json_data.get("AlertId")
    o365_data["object_id"] = json_data.get("ObjectId")
    o365_data["creation_time"] = json_data.get("CreationTime")
    o365_data["comments"] = json_data.get("Comments")
    o365_data["source"] = json_data.get("Source", "Office 365 Security & Compliance")
    o365_data["policy_id"] = json_data.get("PolicyId")
    o365_data["record_type"] = json_data.get("RecordType")
    o365_data["organization_id"] = json_data.get("OrganizationId")
    
    # Parse nested Data field (JSON string)
    data_field = json_data.get("Data")
    if isinstance(data_field, str):
        try:
            data_json = json.loads(data_field)
            o365_data["data"] = data_json
            
            # Extract key information from Data field
            o365_data["user_email"] = data_json.get("f3u")  # User email
            o365_data["alert_description"] = data_json.get("ad")  # Alert description
            o365_data["label_operation"] = data_json.get("lon")  # Label operation name
            o365_data["operation_type"] = data_json.get("op")  # Operation type
            o365_data["start_time"] = data_json.get("ts")  # Start time
            o365_data["end_time"] = data_json.get("te")  # End time
            o365_data["alert_name_internal"] = data_json.get("an")  # Alert name internal
            o365_data["severity_internal"] = data_json.get("sev")  # Severity internal
            o365_data["report_id"] = data_json.get("reid")  # Report ID
            o365_data["rule_id"] = data_json.get("rid")  # Rule ID
            
            # Set username from user email if available
            if o365_data["user_email"]:
                analysis["username"] = o365_data["user_email"]
                analysis["user_information"]["email"] = o365_data["user_email"]
                if o365_data["user_email"] not in analysis["identifiers"]["emails"]:
                    analysis["identifiers"]["emails"].append(o365_data["user_email"])
        except (json.JSONDecodeError, TypeError):
            o365_data["data"] = data_field  # Store as string if not JSON
    
    # Set event type and action based on alert information
    if o365_data.get("alert_name"):
        analysis["event_type"] = o365_data["alert_name"]
    
    if o365_data.get("operation"):
        analysis["action"] = o365_data["operation"]
    
    if o365_data.get("label_operation"):
        analysis["action"] = o365_data["label_operation"]
    
    # Set message/description
    if o365_data.get("alert_description"):
        analysis["message"] = o365_data["alert_description"]
    elif o365_data.get("comments"):
        analysis["message"] = o365_data["comments"]
    elif o365_data.get("alert_name"):
        analysis["message"] = o365_data["alert_name"]
    
    # Set timestamp
    if o365_data.get("creation_time"):
        analysis["timestamp"] = o365_data["creation_time"]
    elif o365_data.get("start_time"):
        analysis["timestamp"] = o365_data["start_time"]
    
    # Set risk level based on severity
    severity = o365_data.get("severity", "").lower()
    if severity == "high":
        analysis["risk_level"] = "high"
    elif severity == "medium":
        analysis["risk_level"] = "medium"
    elif severity == "low":
        analysis["risk_level"] = "low"
    
    # Add risk indicators based on alert type
    if o365_data.get("alert_type"):
        analysis["risk_indicators"].append(f"Alert Type: {o365_data['alert_type']}")
    
    if o365_data.get("category"):
        analysis["risk_indicators"].append(f"Category: {o365_data['category']}")
    
    # Format operation description for display
    operation_desc = ""
    if o365_data.get("label_operation"):
        operation_desc = o365_data["label_operation"]
    elif o365_data.get("alert_name"):
        operation_desc = o365_data["alert_name"]
    elif o365_data.get("operation"):
        operation_desc = o365_data["operation"]
    
    if operation_desc:
        o365_data["operation_description"] = operation_desc
        analysis["action"] = operation_desc
    
    return analysis

def parse_siem_log(log_text: str) -> dict:
    """
    Comprehensive SIEM log parser that extracts:
    - IP addresses (source, destination, etc.)
    - Timestamps
    - Protocols and ports
    - URLs and domains
    - User agents
    - HTTP methods and status codes
    - Usernames
    - Event types and actions
    - File paths and process names
    - Risk indicators
    """
    analysis = {
        "timestamp": None,
        "message": None,
        "event_type": None,
        "action": None,
        "username": None,
        "src_ip": None,
        "dst_ip": None,
        "protocol": None,
        "src_port": None,
        "dst_port": None,
        "ips": [],
        "urls": [],
        "domains": [],
        "hashes": [],
        "files": [],
        "user_agent": None,
        "http_method": None,
        "http_status": None,
        "url": None,
        "file_path": None,
        "process_name": None,
        "risk_level": "low",
        "risk_indicators": [],
        "exfil_score": 0,
        "exfil_level": "low",
        "exfil_factors": [],
        "log_format": "unknown",
        "identifiers": {
            "username": None,
            "emails": [],
            "usernames": []
        },
        "user_information": {
            "username": None,
            "email": None
        }
    }
    
    # Use pre-compiled regex patterns from module level
    
    def get_nested(data: dict, path: str):
        current = data
        for part in path.split('.'):
            if isinstance(current, dict):
                current = current.get(part)
            else:
                return None
        return current

    original_log_text = log_text
    json_data = None  # Cache JSON parsing result

    try:
        json_data = json.loads(log_text.strip())
        analysis["log_format"] = "json"
        
        # Check if this is an Office 365 Security & Compliance alert
        if json_data.get("Workload") == "SecurityComplianceCenter" or json_data.get("Operation") == "AlertTriggered" or json_data.get("Source") == "Office 365 Security & Compliance":
            analysis["log_format"] = "office365_alert"
            analysis = parse_office365_alert(json_data, analysis)
            # Continue with regular parsing as well to extract additional info
        
        # Common JSON log fields
        analysis["timestamp"] = json_data.get("timestamp") or json_data.get("time") or json_data.get("@timestamp") or json_data.get("event_time") or json_data.get("CreationTime")
        analysis["message"] = json_data.get("message") or json_data.get("msg") or json_data.get("log") or json_data.get("Comments")
        analysis["event_type"] = json_data.get("event_type") or json_data.get("event") or json_data.get("type") or json_data.get("event.action") or json_data.get("Operation")
        analysis["action"] = json_data.get("action") or json_data.get("event.action") or json_data.get("Operation")

        username_paths = [
            "username",
            "user",
            "user.name",
            "user.username",
            "user.user_name",
            "user.id",
            "user.email",
            "user.identity",
            "userName",
            "user_name",
            "userid",
            "userId",
            "principal.name",
            "principal.username",
            "actor.user",
            "actor.name",
            "actor.username",
            "identity.user",
            "identity.username",
            "identity.userName",
            "winlog.event_data.SubjectUserName",
            "winlog.event_data.TargetUserName",
            "winlog.event_data.TargetDomainName",
            "winlog.event_data.SubjectDomainName",
            "target.user.name",
            "target.user",
            "SubjectUserName",
            "TargetUserName"
        ]
        for path in username_paths:
            candidate = get_nested(json_data, path) if '.' in path else json_data.get(path)
            if isinstance(candidate, dict):
                continue
            if isinstance(candidate, str) and candidate.strip():
                analysis["username"] = candidate.strip()
                break

        analysis["src_ip"] = json_data.get("src_ip") or json_data.get("source.ip") or json_data.get("src") or json_data.get("source_address")
        analysis["dst_ip"] = json_data.get("dst_ip") or json_data.get("destination.ip") or json_data.get("dst") or json_data.get("dest_ip") or json_data.get("destination_address")
        analysis["protocol"] = json_data.get("protocol") or json_data.get("network.protocol") or json_data.get("network.transport")
        analysis["src_port"] = json_data.get("src_port") or json_data.get("source.port") or json_data.get("source_port")
        analysis["dst_port"] = json_data.get("dst_port") or json_data.get("destination.port") or json_data.get("dest_port") or json_data.get("destination_port")
        analysis["http_method"] = json_data.get("http.request.method") or json_data.get("method") or json_data.get("http_method")
        analysis["http_status"] = json_data.get("http.response.status_code") or json_data.get("status_code") or json_data.get("response_code")
        analysis["url"] = json_data.get("url") or json_data.get("http.request.original") or json_data.get("request")
        analysis["user_agent"] = json_data.get("user_agent") or json_data.get("user_agent.original") or json_data.get("http.request.headers.user-agent")
        analysis["file_path"] = json_data.get("file.path") or json_data.get("filepath") or json_data.get("file_name") or json_data.get("winlog.event_data.TargetFilename")
        analysis["process_name"] = json_data.get("process.name") or json_data.get("process_name") or json_data.get("process") or json_data.get("winlog.event_data.Image")
        
        # Extract from nested structures
        if "source" in json_data and isinstance(json_data["source"], dict):
            if not analysis["src_ip"]:
                analysis["src_ip"] = json_data["source"].get("ip")
            if not analysis["src_port"]:
                analysis["src_port"] = json_data["source"].get("port")
        
        if "destination" in json_data and isinstance(json_data["destination"], dict):
            if not analysis["dst_ip"]:
                analysis["dst_ip"] = json_data["destination"].get("ip")
            if not analysis["dst_port"]:
                analysis["dst_port"] = json_data["destination"].get("port")
            
        if "http" in json_data and isinstance(json_data["http"], dict):
            if not analysis["http_method"]:
                analysis["http_method"] = json_data["http"].get("request", {}).get("method")
            if not analysis["http_status"]:
                analysis["http_status"] = json_data["http"].get("response", {}).get("status_code")
            if not analysis["url"]:
                analysis["url"] = json_data["http"].get("request", {}).get("original")
            
        # Continue parsing from the message field if it exists, but keep original for display
        if analysis["message"]:
            log_text_for_parsing = analysis["message"]
        else:
            log_text_for_parsing = log_text
    except (json.JSONDecodeError, AttributeError):
        # Not JSON, continue with regex parsing
        log_text_for_parsing = log_text
    
    # Extract IPs with context - parse both original text and message field if it exists
    # This ensures we catch IPs in structured logs (JSON) and unstructured logs
    texts_to_parse = [original_log_text]  # Always parse original log
    if 'log_text_for_parsing' in locals() and log_text_for_parsing != original_log_text:
        texts_to_parse.append(log_text_for_parsing)  # Also parse message field if different
    
    # Use modular extraction functions
    analysis = extract_ips_from_text(texts_to_parse, original_log_text, analysis)
    analysis = extract_network_info(texts_to_parse, analysis)
    analysis = extract_urls_and_domains(texts_to_parse, analysis)
    
    # Extract hashes using modular function (handles JSON fields internally)
    analysis = extract_hashes(texts_to_parse, original_log_text, json_data, analysis)
    
    # Extract user agent using pre-compiled pattern
    if not analysis["user_agent"]:
        for text in texts_to_parse:
            ua_match = _USER_AGENT_REGEX_PATTERN.search(text)
            if ua_match:
                analysis["user_agent"] = ua_match.group(1).strip()
                break
    
    # Extract HTTP method using pre-compiled pattern
    if not analysis["http_method"]:
        for text in texts_to_parse:
            method_match = _HTTP_METHOD_REGEX_PATTERN.search(text)
            if method_match:
                analysis["http_method"] = method_match.group(1)
                break
    
    # Extract HTTP status using pre-compiled pattern
    if not analysis["http_status"]:
        for text in texts_to_parse:
            status_match = _HTTP_STATUS_REGEX_PATTERN.search(text)
            if status_match:
                analysis["http_status"] = status_match.group(1)
                break
    
    # Extract username using pre-compiled patterns
    if not analysis["username"]:
        for text in texts_to_parse:
            for pattern in _USERNAME_PATTERNS_COMPILED:
                user_match = pattern.search(text)
                if user_match:
                    username_candidate = user_match.group(1).strip()
                    if username_candidate:
                        analysis["username"] = username_candidate
                        break
            if analysis["username"]:
                break

    def add_unique(collection, value):
        if value and value not in collection:
            collection.append(value)

    for text in texts_to_parse:
        for pattern in _USERNAME_PATTERNS_COMPILED:
            user_match = pattern.search(text)
            if user_match:
                username_candidate = user_match.group(1).strip()
                add_unique(analysis["identifiers"]["usernames"], username_candidate)

        for email_match in _EMAIL_REGEX_PATTERN.finditer(text):
            add_unique(analysis["identifiers"]["emails"], email_match.group(0).strip())

        for mobileye_match in _MOBILEYE_EMAIL_REGEX_PATTERN.finditer(text):
            mobileye_username = mobileye_match.group(1).strip()
            mobileye_email = mobileye_match.group(0).strip()
            if mobileye_username:
                analysis["user_information"]["username"] = mobileye_username
                analysis["user_information"]["email"] = mobileye_email
                add_unique(analysis["identifiers"]["usernames"], mobileye_username)
                add_unique(analysis["identifiers"]["emails"], mobileye_email)
                analysis["username"] = mobileye_username

    if not analysis["username"] and analysis["identifiers"]["emails"]:
        analysis["username"] = analysis["identifiers"]["emails"][0]

    # Extract file paths using pre-compiled pattern
    if not analysis["file_path"]:
        for text in texts_to_parse:
            file_match = _FILE_PATH_REGEX_PATTERN.search(text)
            if file_match:
                analysis["file_path"] = file_match.group(0)
                break
    
    # Extract all files from log
    def get_file_type(filename: str) -> str:
        """Determine file type based on extension."""
        if not filename:
            return "unknown"
        filename_lower = filename.lower()
        
        # Executable files
        if filename_lower.endswith(('.exe', '.dll', '.sys', '.drv', '.com', '.scr', '.msi', '.bat', '.cmd', '.ps1', '.vbs', '.js', '.jar')):
            return "executable"
        # Script files
        elif filename_lower.endswith(('.sh', '.bash', '.py', '.pl', '.rb', '.php', '.asp', '.aspx', '.jsp')):
            return "script"
        # Document files
        elif filename_lower.endswith(('.doc', '.docx', '.pdf', '.xls', '.xlsx', '.ppt', '.pptx', '.txt', '.rtf', '.odt')):
            return "document"
        # Archive files
        elif filename_lower.endswith(('.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.cab', '.iso')):
            return "archive"
        # Image files
        elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg', '.ico', '.tiff')):
            return "image"
        # Video files
        elif filename_lower.endswith(('.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv', '.webm')):
            return "video"
        # Audio files
        elif filename_lower.endswith(('.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma')):
            return "audio"
        # Database files
        elif filename_lower.endswith(('.db', '.sqlite', '.mdb', '.accdb', '.sql')):
            return "database"
        # Config files
        elif filename_lower.endswith(('.config', '.conf', '.ini', '.xml', '.json', '.yaml', '.yml', '.properties')):
            return "config"
        # Log files
        elif filename_lower.endswith(('.log', '.txt')):
            return "log"
        else:
            return "other"
    
    # Enhanced file detection using pre-compiled patterns
    seen_files = set()
    all_files = []
    
    for text in texts_to_parse:
        # Windows paths
        for match in _WINDOWS_FILE_PATTERN.finditer(text):
            file_path = match.group(0)
            if file_path not in seen_files and len(file_path) > 3:
                seen_files.add(file_path)
                filename = file_path.split('\\')[-1] if '\\' in file_path else file_path
                all_files.append({
                    "path": file_path,
                    "name": filename,
                    "type": get_file_type(filename),
                    "location": "windows_path"
                })
        
        # Unix paths
        for match in _UNIX_FILE_PATTERN.finditer(text):
            file_path = match.group(0)
            if file_path not in seen_files and len(file_path) > 3:
                seen_files.add(file_path)
                filename = file_path.split('/')[-1] if '/' in file_path else file_path
                all_files.append({
                    "path": file_path,
                    "name": filename,
                    "type": get_file_type(filename),
                    "location": "unix_path"
                })
        
        # Quoted files
        for match in _QUOTED_FILE_PATTERN.finditer(text):
            file_path = match.group(1)
            if file_path not in seen_files and len(file_path) > 3:
                seen_files.add(file_path)
                filename = file_path.split('/')[-1] if '/' in file_path else file_path.split('\\')[-1]
                all_files.append({
                    "path": file_path,
                    "name": filename,
                    "type": get_file_type(filename),
                    "location": "quoted"
                })
        
        # Files in URLs
        for match in _URL_FILE_PATTERN.finditer(text):
            filename = match.group(1)
            if filename not in seen_files and len(filename) > 3:
                seen_files.add(filename)
                all_files.append({
                    "path": match.group(0),  # Full URL
                    "name": filename,
                    "type": get_file_type(filename),
                    "location": "url"
                })
        
        # Files after keywords
        for match in _KEYWORD_FILE_PATTERN.finditer(text):
            file_path = match.group(1)
            if file_path not in seen_files and len(file_path) > 3:
                seen_files.add(file_path)
                filename = file_path.split('/')[-1] if '/' in file_path else file_path.split('\\')[-1]
                all_files.append({
                    "path": file_path,
                    "name": filename,
                    "type": get_file_type(filename),
                    "location": "keyword_context"
                })
    
    # Also check JSON fields for files (reuse cached json_data if available)
    if json_data is None:
        try:
            json_data = json.loads(original_log_text.strip())
        except (json.JSONDecodeError, AttributeError, TypeError):
            json_data = None
    
    if json_data:
        try:
            file_fields = [
                "file", "filename", "filepath", "file_path", "file_name", "target_file", "source_file",
                "file.path", "file.name", "winlog.event_data.TargetFilename", "winlog.event_data.Image"
            ]
            for field in file_fields:
                if '.' in field:
                    parts = field.split('.')
                    current = json_data
                    for part in parts:
                        if isinstance(current, dict) and part in current:
                            current = current[part]
                        else:
                            current = None
                            break
                    if isinstance(current, str) and current.strip() and current.strip() not in seen_files:
                        file_path = current.strip()
                        seen_files.add(file_path)
                        filename = file_path.split('/')[-1] if '/' in file_path else file_path.split('\\')[-1]
                        all_files.append({
                            "path": file_path,
                            "name": filename,
                            "type": get_file_type(filename),
                            "location": "json_field"
                        })
                else:
                    if field in json_data and isinstance(json_data[field], str) and json_data[field].strip():
                        file_path = json_data[field].strip()
                        if file_path not in seen_files:
                            seen_files.add(file_path)
                            filename = file_path.split('/')[-1] if '/' in file_path else file_path.split('\\')[-1]
                            all_files.append({
                                "path": file_path,
                                "name": filename,
                                "type": get_file_type(filename),
                                "location": "json_field"
                            })
        except (json.JSONDecodeError, AttributeError, TypeError):
            pass
    
    analysis["files"] = all_files
    
    # Extract process name using pre-compiled pattern
    if not analysis["process_name"]:
        for text in texts_to_parse:
            process_match = _PROCESS_REGEX_PATTERN.search(text)
            if process_match:
                analysis["process_name"] = process_match.group(1)
                break
    
    # Extract timestamp (various formats) - check all texts using pre-compiled patterns
    if not analysis["timestamp"]:
        for text in texts_to_parse:
            for pattern in _TIMESTAMP_PATTERNS:
                ts_match = pattern.search(text)
                if ts_match:
                    analysis["timestamp"] = ts_match.group(1)
                    break
            if analysis["timestamp"]:
                break
    
    # Normalize timestamp to ISO format or Unix timestamp
    if analysis["timestamp"]:
        try:
            ts_str = str(analysis["timestamp"])
            # Try parsing as Unix timestamp (10-13 digits)
            if ts_str.isdigit() and len(ts_str) >= 10:
                ts_int = int(ts_str[:10])  # Use first 10 digits for seconds
                if len(ts_str) > 10:
                    # Milliseconds timestamp
                    ts_int = int(ts_str[:13]) / 1000
                analysis["timestamp_unix"] = ts_int
                analysis["timestamp_iso"] = datetime.fromtimestamp(ts_int).isoformat()
            else:
                # Try parsing as ISO format
                try:
                    dt = datetime.fromisoformat(ts_str.replace('Z', '+00:00'))
                    analysis["timestamp_unix"] = dt.timestamp()
                    analysis["timestamp_iso"] = dt.isoformat()
                except:
                    # Try common formats
                    for fmt in ['%Y-%m-%d %H:%M:%S', '%Y-%m-%dT%H:%M:%S', '%Y/%m/%d %H:%M:%S']:
                        try:
                            dt = datetime.strptime(ts_str[:19], fmt)
                            analysis["timestamp_unix"] = dt.timestamp()
                            analysis["timestamp_iso"] = dt.isoformat()
                            break
                        except:
                            continue
        except Exception as e:
            print(f"[LogParser] Error normalizing timestamp: {e}")
            analysis["timestamp_unix"] = None
            analysis["timestamp_iso"] = None
    else:
        analysis["timestamp_unix"] = None
        analysis["timestamp_iso"] = None
    
    # Extract event type and action - search in all texts using pre-compiled pattern
    if not analysis["event_type"]:
        for text in texts_to_parse:
            event_match = _EVENT_KEYWORDS_PATTERN.search(text)
            if event_match:
                analysis["event_type"] = event_match.group(1).upper()
                break
    
    # Risk assessment - use original log text for comprehensive risk analysis using pre-compiled patterns
    risk_score = 0
    for match in _RISK_KEYWORDS_HIGH_PATTERN.finditer(original_log_text):
        keyword = match.group(1).upper()
        risk_score += 3
        analysis["risk_indicators"].append(keyword)
    
    for match in _RISK_KEYWORDS_MEDIUM_PATTERN.finditer(original_log_text):
        keyword = match.group(1).upper()
        risk_score += 1
        analysis["risk_indicators"].append(keyword)
    
    # Check for suspicious ports
    if analysis["dst_port"]:
        suspicious_ports = ["4444", "5555", "6666", "6667", "12345", "31337"]
        if analysis["dst_port"] in suspicious_ports:
            risk_score += 2
            analysis["risk_indicators"].append(f"SUSPICIOUS_PORT:{analysis['dst_port']}")
    
    # Check for suspicious HTTP status codes
    if analysis["http_status"]:
        if analysis["http_status"].startswith("4") or analysis["http_status"].startswith("5"):
            risk_score += 1
    
    # Determine risk level
    if risk_score >= 5:
        analysis["risk_level"] = "high"
    elif risk_score >= 2:
        analysis["risk_level"] = "medium"
    else:
        analysis["risk_level"] = "low"
    
    # ========== ADVANCED EXFILTRATION PROBABILITY MODEL ==========
    # Note: is_private_ip is now defined at module level for reuse
    
    def sigmoid_normalize(raw_score: float, max_score: float = 20.0, steepness: float = 0.3) -> float:
        """
        Normalize raw score to 0-100 using S-curve (sigmoid-like function).
        This provides smooth transitions and allows sensitivity control.
        """
        import math
        # Scale to 0-1 range first
        normalized = min(1.0, raw_score / max_score)
        # Apply sigmoid transformation: 1 / (1 + e^(-k*(x-0.5)))
        # k controls steepness, x-0.5 centers the curve
        k = steepness * 10  # Adjust steepness multiplier
        sigmoid_value = 1 / (1 + math.exp(-k * (normalized - 0.5)))
        return min(100, max(0, sigmoid_value * 100))
    
    # ========== WEIGHT TABLE (Structured Weights) ==========
    WEIGHTS = {
        "flow_direction": {
            "internal_to_external": 4.0,  # High weight
            "external_to_internal": 0.5,  # Low weight
            "internal_to_internal": 0.0,
            "external_to_external": 0.0
        },
        "protocol": {
            "high_risk": ["HTTPS", "HTTP", "FTP", "SFTP", "SCP", "SMB", "TLS", "SSL"],  # Medium+ weight: 2.0
            "low_risk": ["DNS", "ICMP"],  # Low weight: 0.3
            "high_risk_weight": 2.0,
            "low_risk_weight": 0.3
        },
        "http_method": {
            "POST": 2.5,  # Positive weight
            "PUT": 2.5,
            "PATCH": 2.5,
            "GET": 0.2,  # Small or zero
            "DELETE": 1.0,
            "other": 0.0
        },
        "keywords": {
            "high_priority": ["export", "backup", "dump", "sync", "archive", "exfil"],  # High weight: 1.5 each
            "medium_priority": ["upload", "transfer", "copy", "move", "send"],  # Medium weight: 1.0 each
            "high_weight": 1.5,
            "medium_weight": 1.0
        },
        "risk_level": {
            "low": 1.0,  # Adds little
            "medium": 2.5,  # Adds more
            "high": 4.0  # Adds a lot
        },
        "risk_indicators": {
            "strong": ["ATTACK", "MALWARE", "EXPLOIT", "INTRUSION", "BREACH"],  # Weight: 2.0 each
            "moderate": ["UNAUTHORIZED", "SUSPICIOUS"],  # Weight: 1.0 each
            "strong_weight": 2.0,
            "moderate_weight": 1.0
        }
    }
    
    # ========== EVENT TYPE BASE SCORES ==========
    EVENT_TYPE_BASES = {
        "email": 2.0,  # Email events can be more sensitive
        "browser": 1.0,  # End-user browser events: medium risk
        "database": 3.0,  # DB queries or file shares: higher base
        "file_share": 3.0,
        "network": 1.5,
        "authentication": 1.0,
        "default": 1.0
    }
    
    # Determine event type
    event_type_key = "default"
    if analysis.get("event_type"):
        event_type_lower = analysis["event_type"].lower()
        if any(x in event_type_lower for x in ["mail", "email", "smtp", "imap"]):
            event_type_key = "email"
        elif any(x in event_type_lower for x in ["browser", "http", "web", "url"]):
            event_type_key = "browser"
        elif any(x in event_type_lower for x in ["database", "db", "sql", "query"]):
            event_type_key = "database"
        elif any(x in event_type_lower for x in ["file", "share", "smb", "nfs"]):
            event_type_key = "file_share"
        elif any(x in event_type_lower for x in ["network", "tcp", "udp", "connection"]):
            event_type_key = "network"
        elif any(x in event_type_lower for x in ["auth", "login", "logout"]):
            event_type_key = "authentication"
    
    base_score = EVENT_TYPE_BASES.get(event_type_key, EVENT_TYPE_BASES["default"])
    
    # ========== COMPONENT SCORES ==========
    infra_score = 0.0  # Infrastructure: protocol, port, internal/external
    content_score = 0.0  # Content: keywords, file paths, URLs
    behavior_score = 0.0  # Behavior: risk level, indicators (baseline would go here in Phase 2)
    exfil_factors = []
    confidence_factors = []
    
    # ========== INFRASTRUCTURE SCORE ==========
    # Flow direction (using optimized is_private_ip function)
    if analysis.get("src_ip") and analysis.get("dst_ip"):
        src_is_private = is_private_ip(analysis["src_ip"])
        dst_is_private = is_private_ip(analysis["dst_ip"])
        if src_is_private and not dst_is_private:
            weight = WEIGHTS["flow_direction"]["internal_to_external"]
            infra_score += weight
            exfil_factors.append(f"Traffic from internal IP ({analysis['src_ip']}) to external destination ({analysis['dst_ip']})")
            confidence_factors.append("flow_direction")
        elif not src_is_private and dst_is_private:
            weight = WEIGHTS["flow_direction"]["external_to_internal"]
            infra_score += weight
            exfil_factors.append(f"Traffic from external IP ({analysis['src_ip']}) to internal destination ({analysis['dst_ip']})")
    
    # Protocol
    if analysis.get("protocol"):
        protocol_upper = analysis["protocol"].upper()
        if protocol_upper in WEIGHTS["protocol"]["high_risk"]:
            weight = WEIGHTS["protocol"]["high_risk_weight"]
            infra_score += weight
            exfil_factors.append(f"Protocol is commonly used for data transfer ({protocol_upper})")
            confidence_factors.append("protocol")
        elif protocol_upper in WEIGHTS["protocol"]["low_risk"]:
            weight = WEIGHTS["protocol"]["low_risk_weight"]
            infra_score += weight
            exfil_factors.append(f"Low-risk protocol detected ({protocol_upper})")
    
    # Port
    exfil_ports = {
        "21": "FTP", "22": "SSH/SCP/SFTP", "80": "HTTP", "443": "HTTPS",
        "445": "SMB", "3306": "MySQL", "5432": "PostgreSQL",
        "8080": "HTTP (alt)", "8443": "HTTPS (alt)"
    }
    if analysis.get("dst_port"):
        port_str = str(analysis["dst_port"])
        if port_str in exfil_ports:
            infra_score += 1.5
            port_desc = exfil_ports[port_str]
            exfil_factors.append(f"Destination port {port_str} ({port_desc}) is commonly used for data transfer")
            confidence_factors.append("port")
    
    # HTTP Method
    if analysis.get("http_method"):
        method_upper = analysis["http_method"].upper()
        method_weight = WEIGHTS["http_method"].get(method_upper, WEIGHTS["http_method"]["other"])
        if method_weight > 0:
            infra_score += method_weight
            if method_upper in ["POST", "PUT", "PATCH"]:
                exfil_factors.append(f"HTTP method {method_upper} indicates possible data upload")
            confidence_factors.append("http_method")
    
    # ========== CONTENT SCORE ==========
    found_keywords = []
    texts_to_check = [original_log_text]
    if analysis.get("url"):
        texts_to_check.append(analysis["url"])
    if analysis.get("file_path"):
        texts_to_check.append(analysis["file_path"])
    
    for text_to_check in texts_to_check:
        # Use pre-compiled patterns for better performance
        for match in _EXFIL_KEYWORDS_HIGH_PATTERN.finditer(text_to_check):
            keyword = match.group(1).lower()
            if keyword not in found_keywords:
                found_keywords.append(keyword)
                content_score += WEIGHTS["keywords"]["high_weight"]
        
        for match in _EXFIL_KEYWORDS_MEDIUM_PATTERN.finditer(text_to_check):
            keyword = match.group(1).lower()
            if keyword not in found_keywords:
                found_keywords.append(keyword)
                content_score += WEIGHTS["keywords"]["medium_weight"]
    
    if found_keywords:
        location_info = []
        if any(re.search(rf"\b{k}\b", original_log_text, re.IGNORECASE) for k in found_keywords):
            location_info.append("log text")
        if analysis.get("url") and any(re.search(rf"\b{k}\b", analysis["url"], re.IGNORECASE) for k in found_keywords):
            location_info.append("URL")
        if analysis.get("file_path") and any(re.search(rf"\b{k}\b", analysis["file_path"], re.IGNORECASE) for k in found_keywords):
            location_info.append("file path")
        location_str = f" in {', '.join(location_info)}" if location_info else ""
        exfil_factors.append(f"Keyword(s) found{location_str}: {', '.join(found_keywords)}")
        confidence_factors.append("keywords")
    
    # ========== BEHAVIOR SCORE ==========
    # Risk level
    risk_level = analysis.get("risk_level", "low")
    behavior_score += WEIGHTS["risk_level"].get(risk_level, WEIGHTS["risk_level"]["low"])
    if risk_level != "low":
        exfil_factors.append(f"Event risk level is {risk_level}")
        confidence_factors.append("risk_level")
    
    # Risk indicators
    for indicator in analysis.get("risk_indicators", []):
        indicator_str = str(indicator).upper()
        if any(strong in indicator_str for strong in WEIGHTS["risk_indicators"]["strong"]):
            behavior_score += WEIGHTS["risk_indicators"]["strong_weight"]
            exfil_factors.append(f"Strong malicious indicator detected: {indicator}")
            confidence_factors.append("risk_indicators")
        elif any(mod in indicator_str for mod in WEIGHTS["risk_indicators"]["moderate"]):
            behavior_score += WEIGHTS["risk_indicators"]["moderate_weight"]
            exfil_factors.append(f"Moderate malicious indicator detected: {indicator}")
    
    # ========== CALCULATE RAW SCORE ==========
    raw_score = base_score + infra_score + content_score + behavior_score
    
    # ========== NORMALIZE USING S-CURVE ==========
    exfil_score = sigmoid_normalize(raw_score, max_score=20.0, steepness=0.3)
    
    # ========== CONFIDENCE SCORE ==========
    # Based on how many quality features were available
    available_features = len(set(confidence_factors))
    max_possible_features = 7  # flow, protocol, port, http_method, keywords, risk_level, risk_indicators
    confidence_score = min(100, (available_features / max_possible_features) * 100)
    
    if confidence_score >= 70:
        confidence_level = "high"
    elif confidence_score >= 40:
        confidence_level = "medium"
    else:
        confidence_level = "low"
    
    # ========== DETERMINE EXFILTRATION LEVEL ==========
    if exfil_score <= 30:
        exfil_level = "low"
    elif exfil_score <= 70:
        exfil_level = "medium"
    else:
        exfil_level = "high"
    
    # ========== STORE RESULTS ==========
    analysis["exfil_score"] = round(exfil_score, 1)
    analysis["exfil_level"] = exfil_level
    analysis["exfil_factors"] = exfil_factors
    analysis["exfil_confidence"] = round(confidence_score, 1)
    analysis["exfil_confidence_level"] = confidence_level
    analysis["exfil_component_scores"] = {
        "base": round(base_score, 2),
        "infrastructure": round(infra_score, 2),
        "content": round(content_score, 2),
        "behavior": round(behavior_score, 2),
        "raw_total": round(raw_score, 2)
    }
    
    # Generate activity summary
    activity_parts = []
    if analysis["event_type"]:
        activity_parts.append(analysis["event_type"])
    if analysis["protocol"]:
        activity_parts.append(analysis["protocol"])
    if analysis["src_ip"] and analysis["dst_ip"]:
        activity_parts.append(f"{analysis['src_ip']} → {analysis['dst_ip']}")
    elif analysis["src_ip"]:
        activity_parts.append(f"From {analysis['src_ip']}")
    elif analysis["dst_ip"]:
        activity_parts.append(f"To {analysis['dst_ip']}")
    if analysis["dst_port"]:
        activity_parts.append(f"Port {analysis['dst_port']}")
    if analysis["http_method"]:
        activity_parts.append(analysis["http_method"])
    
    analysis["activity_summary"] = " | ".join(activity_parts) if activity_parts else "Log entry analyzed"
    
    # Store original log text for display
    analysis["text"] = original_log_text
    
    return analysis

def map_to_mitre_attack(analysis: dict) -> list:
    """
    Map log analysis to MITRE ATT&CK tactics based on keywords, ports, methods, etc.
    Returns a list of identified tactics.
    """
    tactics = []
    log_text_lower = ""
    
    # Collect all text for analysis
    texts_to_check = []
    if analysis.get("message"):
        texts_to_check.append(str(analysis["message"]).lower())
    if analysis.get("event_type"):
        texts_to_check.append(str(analysis["event_type"]).lower())
    if analysis.get("action"):
        texts_to_check.append(str(analysis["action"]).lower())
    if analysis.get("process_name"):
        texts_to_check.append(str(analysis["process_name"]).lower())
    if analysis.get("file_path"):
        texts_to_check.append(str(analysis["file_path"]).lower())
    if analysis.get("user_agent"):
        texts_to_check.append(str(analysis["user_agent"]).lower())
    
    log_text_lower = " ".join(texts_to_check)
    
    # TA0002 - Execution
    execution_keywords = ['powershell', 'cmd.exe', 'bash', 'sh', 'python', 'perl', 'ruby', 
                          'wscript', 'cscript', 'mshta', 'rundll32', 'regsvr32', 'schtasks',
                          'at.exe', 'sc.exe', 'taskkill', 'net.exe', 'whoami', 'systeminfo']
    if any(keyword in log_text_lower for keyword in execution_keywords):
        tactics.append("Execution")
    
    # TA0011 - Command and Control
    c2_keywords = ['c2', 'beacon', 'command and control', 'command & control', 'c&c',
                   'callback', 'backdoor', 'reverse shell', 'bind shell', 'tunnel']
    c2_ports = ['4444', '8080', '443', '80', '53', '22']
    if any(keyword in log_text_lower for keyword in c2_keywords):
        tactics.append("Command and Control")
    elif analysis.get("dst_port") in c2_ports or analysis.get("src_port") in c2_ports:
        tactics.append("Command and Control")
    
    # TA0006 - Credential Access
    cred_keywords = ['brute', 'brute force', 'login failed', 'authentication failed',
                     'password', 'credential', 'hash', 'kerberos', 'ntlm', 'lsass',
                     'mimikatz', 'pass the hash', 'golden ticket']
    if any(keyword in log_text_lower for keyword in cred_keywords):
        tactics.append("Credential Access")
    
    # TA0010 - Exfiltration
    exfil_keywords = ['curl', 'wget', 'transfer', 'upload', 'download', 'exfil', 'exfiltration',
                      'data export', 'backup', 'sync', 'archive', 'dump', 'export']
    if any(keyword in log_text_lower for keyword in exfil_keywords):
        tactics.append("Exfiltration")
    
    # TA0005 - Defense Evasion
    evasion_keywords = ['disable', 'bypass', 'evade', 'obfuscate', 'encode', 'decode',
                       'base64', 'xor', 'encrypt', 'decrypt', 'clear logs', 'delete logs']
    if any(keyword in log_text_lower for keyword in evasion_keywords):
        tactics.append("Defense Evasion")
    
    # TA0007 - Discovery
    discovery_keywords = ['nmap', 'scan', 'enum', 'enumeration', 'netstat', 'arp',
                         'ipconfig', 'ifconfig', 'whoami', 'systeminfo', 'tasklist']
    if any(keyword in log_text_lower for keyword in discovery_keywords):
        tactics.append("Discovery")
    
    # TA0008 - Lateral Movement
    lateral_keywords = ['psexec', 'wmic', 'smb', 'rpc', 'rdp', 'ssh', 'remote',
                       'lateral', 'pivot', 'pass the hash']
    if any(keyword in log_text_lower for keyword in lateral_keywords):
        tactics.append("Lateral Movement")
    
    # TA0001 - Initial Access
    initial_keywords = ['phishing', 'spear', 'malware', 'trojan', 'virus', 'exploit',
                       'vulnerability', 'cve-', 'sql injection', 'xss', 'rce']
    if any(keyword in log_text_lower for keyword in initial_keywords):
        tactics.append("Initial Access")
    
    # TA0004 - Privilege Escalation
    priv_keywords = ['privilege', 'escalation', 'sudo', 'su ', 'runas', 'uac bypass',
                    'token', 'impersonation', 'getsystem']
    if any(keyword in log_text_lower for keyword in priv_keywords):
        tactics.append("Privilege Escalation")
    
    # TA0009 - Collection
    collection_keywords = ['keylog', 'screenshot', 'clipboard', 'browser', 'history',
                          'collect', 'gather', 'harvest']
    if any(keyword in log_text_lower for keyword in collection_keywords):
        tactics.append("Collection")
    
    # Remove duplicates and return
    return list(set(tactics))

def build_graph_data(analysis: dict, vt_results: dict, abuseipdb_results: dict, relationships: dict = None) -> list:
    """
    Build graph data for Cytoscape.js Threat Constellation Map.
    Returns a list of nodes and edges in Cytoscape format.
    Only includes relevant/threat nodes to reduce clutter.
    """
    if relationships is None:
        relationships = {}
    elements = []
    node_ids = set()
    
    # Central Node: Log Event
    central_id = "log_event"
    elements.append({
        "data": {
            "id": central_id,
            "label": "Log Event",
            "shortLabel": "",
            "type": "event",
            "color": "#5EEAD4"
        },
        "classes": "central-node"
    })
    node_ids.add(central_id)
    
    # Helper to check if IP is malicious/suspicious
    def is_threat_ip(ip: str) -> bool:
        """Check if IP is a known threat."""
        vt_ips = vt_results.get("ips", {})
        abuse_data = abuseipdb_results.get(ip, {})
        
        # Check VirusTotal
        if ip in vt_ips:
            vt_data = vt_ips[ip]
            if isinstance(vt_data, dict):
                if vt_data.get("malicious", 0) > 0 or vt_data.get("suspicious", 0) > 0:
                    return True
        
        # Check AbuseIPDB
        if abuse_data:
            abuse_score = abuse_data.get("abuseScore", 0)
            try:
                abuse_score_int = int(abuse_score) if abuse_score != "None" else 0
                if abuse_score_int > 20:
                    return True
            except (ValueError, TypeError):
                pass
        
        return False
    
    # Helper function to get IP color based on threat intelligence
    def get_ip_color(ip: str) -> str:
        """Determine IP node color based on VT and AbuseIPDB results."""
        vt_ips = vt_results.get("ips", {})
        abuse_data = abuseipdb_results.get(ip, {})
        
        # Check VirusTotal
        if ip in vt_ips:
            vt_data = vt_ips[ip]
            if isinstance(vt_data, dict):
                malicious = vt_data.get("malicious", 0)
                suspicious = vt_data.get("suspicious", 0)
                if malicious > 0:
                    return "#FF0040"  # Neon Red
                elif suspicious > 0:
                    return "#FF6B00"  # Neon Orange
        
        # Check AbuseIPDB
        if abuse_data:
            abuse_score = abuse_data.get("abuseScore", 0)
            try:
                abuse_score_int = int(abuse_score) if abuse_score != "None" else 0
                if abuse_score_int > 50:
                    return "#FF0040"  # Neon Red
                elif abuse_score_int > 20:
                    return "#FF6B00"  # Neon Orange
            except (ValueError, TypeError):
                pass
        
        return "#00A8FF"  # Neon Blue (safe/unknown)
    
    # IP Nodes - Only show Source/Destination and Threat IPs (NO DUPLICATES)
    seen_ips = set()
    added_ip_ids = set()  # Track which IP IDs we've already added
    src_ip_id = None
    dst_ip_id = None
    threat_ips = []
    other_ips = []
    
    for ip_info in analysis.get("ips", []):
        ip = ip_info.get("ip")
        if not ip or ip in seen_ips:
            continue
        seen_ips.add(ip)
        
        ip_type = ip_info.get("type", "Network")
        context = ip_info.get("context", "generic")
        is_src_dst = (context == "src" or context == "dst" or 
                     ip == analysis.get("src_ip") or ip == analysis.get("dst_ip"))
        is_threat = is_threat_ip(ip)
        
        # Prioritize: Source/Destination IPs and Threat IPs
        if is_src_dst or is_threat:
            threat_ips.append((ip_info, is_src_dst, is_threat))
        else:
            other_ips.append(ip_info)
    
    # Add threat IPs first (ensuring no duplicates)
    for ip_info, is_src_dst, is_threat in threat_ips:
        ip = ip_info.get("ip")
        ip_id = f"ip_{ip.replace('.', '_')}"
        
        # Skip if we already added this IP
        if ip_id in added_ip_ids:
            continue
        added_ip_ids.add(ip_id)
        
        ip_type = ip_info.get("type", "Network")
        context = ip_info.get("context", "generic")
        
        if context == "src" or ip == analysis.get("src_ip"):
            src_ip_id = ip_id
        elif context == "dst" or ip == analysis.get("dst_ip"):
            dst_ip_id = ip_id
        
        color = get_ip_color(ip)
        short_label = ip.split('.')[-1] if '.' in ip else ip[:8]
        
        elements.append({
            "data": {
                "id": ip_id,
                "label": ip,
                "shortLabel": short_label,
                "type": "ip",
                "ip_type": ip_type,
                "color": color
            },
            "classes": "ip-node"
        })
        node_ids.add(ip_id)
        
        # Simple edge label - only show on hover
        edge_label = "Source" if context == "src" else ("Destination" if context == "dst" else "Network")
        
        elements.append({
            "data": {
                "id": f"{central_id}_{ip_id}",
                "source": ip_id,
                "target": central_id,
                "label": "",  # No label by default - cleaner
                "hoverLabel": edge_label,  # Show on hover
                "type": "connection",
                "protocol": analysis.get("protocol"),
                "http_method": analysis.get("http_method"),
                "action": analysis.get("action")
            },
            "classes": "ip-edge"
        })
    
    # Add only top 2 other IPs if we have space (ensuring no duplicates)
    for ip_info in other_ips[:2]:
        ip = ip_info.get("ip")
        ip_id = f"ip_{ip.replace('.', '_')}"
        
        # Skip if we already added this IP
        if ip_id in added_ip_ids:
            continue
        added_ip_ids.add(ip_id)
        
        ip_type = ip_info.get("type", "Network")
        context = ip_info.get("context", "generic")
        color = get_ip_color(ip)
        short_label = ip.split('.')[-1] if '.' in ip else ip[:8]
        
        elements.append({
            "data": {
                "id": ip_id,
                "label": ip,
                "shortLabel": short_label,
                "type": "ip",
                "ip_type": ip_type,
                "color": color
            },
            "classes": "ip-node"
        })
        node_ids.add(ip_id)
        
        # Simple edge - no label by default
        elements.append({
            "data": {
                "id": f"{central_id}_{ip_id}",
                "source": ip_id,
                "target": central_id,
                "label": "",  # No label by default
                "hoverLabel": "Network",
                "type": "connection",
                "protocol": analysis.get("protocol"),
                "http_method": analysis.get("http_method")
            },
            "classes": "ip-edge"
        })
    
    # Hash Nodes (Files removed to reduce clutter)
    hash_nodes = []
    
    # Hash Nodes - Only show malicious/suspicious hashes (max 3, NO DUPLICATES)
    hash_priority = []
    seen_hashes = set()
    
    for hash_info in analysis.get("hashes", []):
        hash_value = hash_info.get("hash")
        if not hash_value or hash_value in seen_hashes:
            continue
        seen_hashes.add(hash_value)
        
        # Check if hash is malicious
        is_malicious = False
        vt_hashes = vt_results.get("hashes", {})
        if hash_value in vt_hashes:
            vt_data = vt_hashes[hash_value]
            if isinstance(vt_data, dict):
                if vt_data.get("malicious", 0) > 0 or vt_data.get("suspicious", 0) > 0:
                    is_malicious = True
        
        hash_priority.append((hash_info, is_malicious))
    
    # Sort: malicious first, then others
    hash_priority.sort(key=lambda x: (not x[1], x[0].get("hash", "")))
    
    # Only add top 3 hashes (ensuring no duplicates)
    added_hash_ids = set()
    for hash_info, is_malicious in hash_priority[:3]:
        hash_value = hash_info.get("hash")
        hash_id = f"hash_{hash_value[:16]}"
        
        # Skip if already added
        if hash_id in added_hash_ids:
            continue
        added_hash_ids.add(hash_id)
        
        hash_type = hash_info.get("type", "Hash")
        short_label = hash_type[:6] if len(hash_type) <= 6 else hash_type[:4] + ".."
        
        elements.append({
            "data": {
                "id": hash_id,
                "label": f"{hash_type}\n{hash_value[:16]}...",
                "shortLabel": short_label,
                "type": "hash",
                "hash_type": hash_type,
                "hash_value": hash_value,
                "color": "#A855F7"
            },
            "classes": "hash-node"
        })
        node_ids.add(hash_id)
        hash_nodes.append(hash_id)
        
        elements.append({
            "data": {
                "id": f"{central_id}_{hash_id}",
                "source": hash_id,
                "target": central_id,
                "label": "Contains",
                "type": "connection"
            },
            "classes": "hash-edge"
        })
    
    # File Nodes - Only show if suspicious or linked to threat IPs (max 2)
    # Skip files for now to reduce clutter - only show if critical
    # (Files are less relevant than IPs and Hashes for threat analysis)
    
    # Advanced Linking removed - files not shown to reduce clutter
    
    # Domain Nodes - Only show malicious/suspicious domains (max 3, NO DUPLICATES)
    domain_priority = []
    seen_domains = set()
    
    for domain in analysis.get("domains", []):
        if not domain or len(domain) < 3 or domain in seen_domains:
            continue
        seen_domains.add(domain)
        
        # Check if domain is malicious
        is_malicious = False
        vt_domains = vt_results.get("domains", {})
        if domain in vt_domains:
            vt_data = vt_domains[domain]
            if isinstance(vt_data, dict):
                if vt_data.get("malicious", 0) > 0 or vt_data.get("suspicious", 0) > 0:
                    is_malicious = True
        
        domain_priority.append((domain, is_malicious))
    
    # Sort: malicious first
    domain_priority.sort(key=lambda x: not x[1])
    
    # Only add top 3 domains (ensuring no duplicates)
    for domain, is_malicious in domain_priority[:3]:
        domain_id = f"domain_{domain.replace('.', '_')}"
        if domain_id in node_ids:
            continue
        
        domain_parts = domain.split('.')
        if len(domain_parts) > 2:
            short_label = domain_parts[-2]
        else:
            short_label = domain[:12] + "..." if len(domain) > 12 else domain
        
        elements.append({
            "data": {
                "id": domain_id,
                "label": domain,
                "shortLabel": short_label,
                "type": "domain",
                "color": "#06F3FF"
            },
            "classes": "domain-node"
        })
        node_ids.add(domain_id)
        
        elements.append({
            "data": {
                "id": f"{central_id}_{domain_id}",
                "source": domain_id,
                "target": central_id,
                "label": "Accessed",
                "type": "connection"
            },
            "classes": "domain-edge"
        })
    
    # User Nodes - Skip users to reduce clutter (not critical for threat visualization)
    
    # ========== CREATE COMMUNICATION FLOW: Source IP -> Destination IP ==========
    # Only show if we have clear source and destination
    if src_ip_id and dst_ip_id and src_ip_id in node_ids and dst_ip_id in node_ids and src_ip_id != dst_ip_id:
        # Simple label - show protocol if available, otherwise just connection
        comm_label = ""
        if analysis.get("protocol"):
            comm_label = analysis["protocol"]
        elif analysis.get("http_method"):
            comm_label = analysis["http_method"]
        
        # Check if edge already exists
        comm_edge_id = f"{src_ip_id}_{dst_ip_id}"
        if not any(e.get("data", {}).get("id") == comm_edge_id for e in elements):
            elements.append({
                "data": {
                    "id": comm_edge_id,
                    "source": src_ip_id,
                    "target": dst_ip_id,
                    "label": "",  # No label by default - cleaner
                    "hoverLabel": comm_label or "Connection",
                    "type": "communication",
                    "protocol": analysis.get("protocol"),
                    "http_method": analysis.get("http_method"),
                    "dst_port": analysis.get("dst_port"),
                    "http_status": analysis.get("http_status")
                },
                "classes": "communication-edge"
            })
    
    # ========== LINK DOMAINS TO DESTINATION IPs ==========
    # Only link top domain to destination IP for simplicity
    if dst_ip_id and dst_ip_id in node_ids and domain_priority:
        domain, is_malicious = domain_priority[0]  # Only top domain
        domain_id = f"domain_{domain.replace('.', '_')}"
        if domain_id in node_ids:
            # Check if edge already exists
            domain_ip_edge_id = f"{domain_id}_{dst_ip_id}"
            if not any(e.get("data", {}).get("id") == domain_ip_edge_id for e in elements):
                elements.append({
                    "data": {
                        "id": domain_ip_edge_id,
                        "source": domain_id,
                        "target": dst_ip_id,
                        "label": "",  # No label by default
                        "hoverLabel": "Resolves to",
                        "type": "domain-ip"
                    },
                    "classes": "domain-ip-edge"
                })
    
    # ========== SKIP URL NODES - Too cluttered, keep it simple ==========
    
    # ========== DEEP RELATIONSHIP ANALYSIS - Related Entities ==========
    if relationships:
        seen_related_entities = set()
        
        for rel_key, rel_data in relationships.items():
            if not rel_data or not isinstance(rel_data, dict):
                continue
            
            # Extract indicator and type from key (format: "indicator|||type")
            if "|||" not in rel_key:
                continue
            main_indicator, main_type = rel_key.split("|||", 1)
            
            # Find the main indicator node ID
            main_node_id = None
            if main_type == "ip":
                main_node_id = f"ip_{main_indicator.replace('.', '_')}"
            elif main_type == "domain":
                main_node_id = f"domain_{main_indicator.replace('.', '_')}"
            elif main_type == "hash":
                main_node_id = f"hash_{main_indicator}"
            
            if main_node_id not in node_ids:
                continue  # Skip if main node doesn't exist
            
            # Process communicating files
            for file_data in rel_data.get("communicating_files", [])[:5]:
                file_id = file_data.get("id")
                if not file_id or file_id in seen_related_entities:
                    continue
                seen_related_entities.add(file_id)
                
                attrs = file_data.get("attributes", {})
                stats = attrs.get("last_analysis_stats", {})
                malicious = stats.get("malicious", 0)
                
                # Only show malicious/suspicious files
                if malicious > 0 or stats.get("suspicious", 0) > 0:
                    file_node_id = f"file_{file_id[:16]}"
                    if file_node_id not in node_ids:
                        elements.append({
                            "data": {
                                "id": file_node_id,
                                "label": file_id[:32] + "..." if len(file_id) > 32 else file_id,
                                "shortLabel": file_id[:8],
                                "type": "related_file",
                                "color": "#F97316",  # Neon Orange for malware/files
                                "malicious": malicious,
                                "mainNodeId": main_node_id  # Store reference for highlighting
                            },
                            "classes": "related-file-node"
                        })
                        node_ids.add(file_node_id)
                    
                    # Create edge with dashed line
                    edge_id = f"{main_node_id}_{file_node_id}_comm"
                    if not any(e.get("data", {}).get("id") == edge_id for e in elements):
                        elements.append({
                            "data": {
                                "id": edge_id,
                                "source": main_node_id,
                                "target": file_node_id,
                                "label": "",
                                "hoverLabel": "Communicates With",
                                "type": "relationship"
                            },
                            "classes": "relationship-edge"
                        })
            
            # Process referrer files (for IPs)
            for file_data in rel_data.get("referrer_files", [])[:5]:
                file_id = file_data.get("id")
                if not file_id or file_id in seen_related_entities:
                    continue
                seen_related_entities.add(file_id)
                
                attrs = file_data.get("attributes", {})
                stats = attrs.get("last_analysis_stats", {})
                malicious = stats.get("malicious", 0)
                
                if malicious > 0 or stats.get("suspicious", 0) > 0:
                    file_node_id = f"file_{file_id[:16]}"
                    if file_node_id not in node_ids:
                        elements.append({
                            "data": {
                                "id": file_node_id,
                                "label": file_id[:32] + "..." if len(file_id) > 32 else file_id,
                                "shortLabel": file_id[:8],
                                "type": "related_file",
                                "color": "#F97316",  # Neon Orange for malware/files
                                "malicious": malicious,
                                "mainNodeId": main_node_id  # Store reference for highlighting
                            },
                            "classes": "related-file-node"
                        })
                        node_ids.add(file_node_id)
                    
                    edge_id = f"{main_node_id}_{file_node_id}_ref"
                    if not any(e.get("data", {}).get("id") == edge_id for e in elements):
                        elements.append({
                            "data": {
                                "id": edge_id,
                                "source": main_node_id,
                                "target": file_node_id,
                                "label": "",
                                "hoverLabel": "Referrer",
                                "type": "relationship"
                            },
                            "classes": "relationship-edge"
                        })
            
            # Process subdomains (for domains)
            for subdomain_data in rel_data.get("subdomains", [])[:5]:
                subdomain = subdomain_data.get("id") or subdomain_data.get("attributes", {}).get("id")
                if not subdomain or subdomain in seen_related_entities:
                    continue
                seen_related_entities.add(subdomain)
                
                subdomain_node_id = f"domain_{subdomain.replace('.', '_')}"
                if subdomain_node_id not in node_ids:
                    elements.append({
                        "data": {
                            "id": subdomain_node_id,
                            "label": subdomain,
                            "shortLabel": subdomain.split('.')[0] if '.' in subdomain else subdomain[:12],
                            "type": "related_domain",
                            "color": "#A855F7",  # Neon Purple for domains/URLs
                            "mainNodeId": main_node_id  # Store reference for highlighting
                        },
                        "classes": "related-domain-node"
                    })
                    node_ids.add(subdomain_node_id)
                
                edge_id = f"{main_node_id}_{subdomain_node_id}_sub"
                if not any(e.get("data", {}).get("id") == edge_id for e in elements):
                    elements.append({
                        "data": {
                            "id": edge_id,
                            "source": main_node_id,
                            "target": subdomain_node_id,
                            "label": "",
                            "hoverLabel": "Subdomain",
                            "type": "relationship"
                        },
                        "classes": "relationship-edge"
                    })
            
            # Process contacted URLs (for hash files)
            for url_data in rel_data.get("contacted_urls", [])[:5]:
                url = url_data.get("url") or url_data.get("id")
                if not url or url in seen_related_entities:
                    continue
                seen_related_entities.add(url)
                
                url_node_id = f"url_{hash(url) % 10000}"
                if url_node_id not in node_ids:
                    elements.append({
                        "data": {
                            "id": url_node_id,
                            "label": url[:40] + "..." if len(url) > 40 else url,
                            "shortLabel": url.split("/")[-1][:12] if "/" in url else url[:12],
                            "type": "related_url",
                            "color": "#A855F7",  # Neon Purple for domains/URLs
                            "mainNodeId": main_node_id  # Store reference for highlighting
                        },
                        "classes": "related-url-node"
                    })
                    node_ids.add(url_node_id)
                
                edge_id = f"{main_node_id}_{url_node_id}_url"
                if not any(e.get("data", {}).get("id") == edge_id for e in elements):
                    elements.append({
                        "data": {
                            "id": edge_id,
                            "source": main_node_id,
                            "target": url_node_id,
                            "label": "",
                            "hoverLabel": "Contacted URL",
                            "type": "relationship"
                        },
                        "classes": "relationship-edge"
                    })
            
            # Process contacted IPs (for hash files)
            for ip_data in rel_data.get("contacted_ips", [])[:5]:
                ip = ip_data.get("id")
                if not ip or ip in seen_related_entities or is_private_ip(ip):
                    continue
                seen_related_entities.add(ip)
                
                ip_node_id = f"ip_{ip.replace('.', '_')}"
                if ip_node_id not in node_ids:
                    elements.append({
                        "data": {
                            "id": ip_node_id,
                            "label": ip,
                            "shortLabel": ip.split('.')[-1],
                            "type": "related_ip",
                            "color": "#F97316",  # Neon Orange for malware-related IPs
                            "mainNodeId": main_node_id  # Store reference for highlighting
                        },
                        "classes": "related-ip-node"
                    })
                    node_ids.add(ip_node_id)
                
                edge_id = f"{main_node_id}_{ip_node_id}_contacted"
                if not any(e.get("data", {}).get("id") == edge_id for e in elements):
                    elements.append({
                        "data": {
                            "id": edge_id,
                            "source": main_node_id,
                            "target": ip_node_id,
                            "label": "",
                            "hoverLabel": "Contacted IP",
                            "type": "relationship"
                        },
                        "classes": "relationship-edge"
                    })
    
    return elements

@app.route("/analyze_log", methods=["POST"])
def analyze_log():
    """
    Analyzes SIEM log entries comprehensively and returns detailed analysis.
    """
    log_text = request.form.get("log_text", "")
    if not log_text:
        return redirect(url_for('home'))

    # Parse the log
    analysis = parse_siem_log(log_text)
    
    print(f"[LogAnalyzer] Analysis complete - Risk: {analysis['risk_level']}, IPs: {len(analysis['ips'])}, Format: {analysis['log_format']}")

    # ========== AUTOMATIC VIRUSTOTAL LOOKUPS (WITH PARALLEL PROCESSING) ==========
    vt_results = {
        "ips": {},
        "hashes": {},
        "domains": {}
    }
    
    # Collect all items to lookup (deduplicated)
    ips_to_lookup = []
    seen_ips = set()
    for ip_info in analysis.get("ips", []):
        ip = ip_info.get("ip")
        if ip and ip not in seen_ips and not is_private_ip(ip):
            seen_ips.add(ip)
            ips_to_lookup.append(ip)
    
    hashes_to_lookup = []
    seen_hashes = set()
    for hash_info in analysis.get("hashes", []):
        hash_value = hash_info.get("hash")
        if hash_value and hash_value not in seen_hashes:
            seen_hashes.add(hash_value)
            hashes_to_lookup.append(hash_value)
    
    domains_to_lookup = []
    seen_domains = set()
    for domain in analysis.get("domains", []):
        if domain and domain not in seen_domains and len(domain) > 3:
            if domain.lower() not in ["localhost", "local", "example.com", "test.com"]:
                seen_domains.add(domain)
                domains_to_lookup.append(domain)
    
    # Extract domains from URLs and add to lookup list
    urls_to_process = analysis.get("urls", [])
    IGNORED_DOMAINS = {
        'ns.adobe.com', 'www.w3.org', 'purl.org', 'schemas.xmlsoap.org', 
        'adobe.com', 'w3.org', 'schemas.microsoft.com'
    }
    for url in urls_to_process:
        try:
            parsed = urlparse(url)
            domain = parsed.netloc
            # Remove port if present
            if ':' in domain:
                domain = domain.split(':')[0]
            # Add domain if valid and not already seen
            if domain and domain not in seen_domains and len(domain) > 3:
                if domain.lower() not in IGNORED_DOMAINS and domain.lower() not in ["localhost", "local", "example.com", "test.com"]:
                    seen_domains.add(domain)
                    domains_to_lookup.append(domain)
        except Exception as e:
            print(f"[LogAnalyzer] Error extracting domain from URL {url}: {e}")
    
    # Parallel API calls using ThreadPoolExecutor
    def lookup_vt_ip(ip):
        """Wrapper for VT IP lookup."""
        print(f"[LogAnalyzer] Auto-lookup IP in VT: {ip}")
        result, error = vt_fetch(ip)
        if result and not error and isinstance(result, dict) and ("malicious" in result or "suspicious" in result or "harmless" in result):
            return ("ip", ip, result)
        return None
    
    def lookup_vt_hash(hash_value):
        """Wrapper for VT hash lookup."""
        print(f"[LogAnalyzer] Auto-lookup Hash in VT: {hash_value[:16]}...")
        result, error = vt_fetch(hash_value)
        if result and not error and isinstance(result, dict) and ("malicious" in result or "suspicious" in result or "harmless" in result):
            return ("hash", hash_value, result)
        return None
    
    def lookup_vt_domain(domain):
        """Wrapper for VT domain lookup."""
        print(f"[LogAnalyzer] Auto-lookup Domain in VT: {domain}")
        result, error = vt_fetch(domain)
        if result and not error and isinstance(result, dict) and ("malicious" in result or "suspicious" in result or "harmless" in result):
            return ("domain", domain, result)
        return None
    
    def lookup_abuseipdb(ip):
        """Wrapper for AbuseIPDB lookup."""
        print(f"[LogAnalyzer] Auto-lookup IP in AbuseIPDB: {ip}")
        result, error = get_abuseipdb_info(ip)
        if result and not error:
            return (ip, result)
        return None
    
    # Execute all lookups in parallel
    with ThreadPoolExecutor(max_workers=5) as executor:
        # Submit all VT lookups
        futures = []
        for ip in ips_to_lookup:
            futures.append(executor.submit(lookup_vt_ip, ip))
        for hash_value in hashes_to_lookup:
            futures.append(executor.submit(lookup_vt_hash, hash_value))
        for domain in domains_to_lookup:
            futures.append(executor.submit(lookup_vt_domain, domain))
        for ip in ips_to_lookup:
            futures.append(executor.submit(lookup_abuseipdb, ip))
        
        # Collect results as they complete
        abuseipdb_results = {}
        for future in as_completed(futures):
            try:
                result = future.result()
                if result:
                    if len(result) == 3:  # VT result
                        result_type, key, value = result
                        vt_results[result_type + "s"][key] = value
                    elif len(result) == 2:  # AbuseIPDB result
                        ip, data = result
                        abuseipdb_results[ip] = data
            except Exception as e:
                print(f"[LogAnalyzer] Error in parallel lookup: {e}")

    # Add VT and AbuseIPDB results to analysis
    analysis["vt_results"] = vt_results
    analysis["abuseipdb_results"] = abuseipdb_results
    
    # ========== DEEP RELATIONSHIP ANALYSIS (Phase 3) ==========
    # Fetch relationships for main indicators in background
    relationships = {}
    
    def fetch_relationships_for_indicator(indicator, ind_type):
        """Wrapper for relationship fetching."""
        print(f"[LogAnalyzer] Fetching relationships for {ind_type}: {indicator}")
        rel_data, error = vt_fetch_relationships(indicator, ind_type)
        if rel_data and not error:
            return (indicator, ind_type, rel_data)
        return None
    
    # Collect relationship lookups for main threat indicators
    relationship_futures = []
    with ThreadPoolExecutor(max_workers=3) as rel_executor:
        # Fetch relationships for threat IPs
        for ip in ips_to_lookup[:3]:  # Limit to top 3 IPs to avoid overload
            if ip in vt_results.get("ips", {}) and not is_private_ip(ip):
                vt_data = vt_results["ips"][ip]
                if isinstance(vt_data, dict) and (vt_data.get("malicious", 0) > 0 or vt_data.get("suspicious", 0) > 0):
                    relationship_futures.append(rel_executor.submit(fetch_relationships_for_indicator, ip, "ip"))
        
        # Fetch relationships for threat domains
        for domain in domains_to_lookup[:2]:  # Limit to top 2 domains
            if domain in vt_results.get("domains", {}):
                vt_data = vt_results["domains"][domain]
                if isinstance(vt_data, dict) and (vt_data.get("malicious", 0) > 0 or vt_data.get("suspicious", 0) > 0):
                    relationship_futures.append(rel_executor.submit(fetch_relationships_for_indicator, domain, "domain"))
        
        # Fetch relationships for threat hashes
        for hash_value in hashes_to_lookup[:2]:  # Limit to top 2 hashes
            if hash_value in vt_results.get("hashes", {}):
                vt_data = vt_results["hashes"][hash_value]
                if isinstance(vt_data, dict) and (vt_data.get("malicious", 0) > 0 or vt_data.get("suspicious", 0) > 0):
                    relationship_futures.append(rel_executor.submit(fetch_relationships_for_indicator, hash_value, "hash"))
        
        # Collect relationship results
        for future in as_completed(relationship_futures):
            try:
                result = future.result()
                if result:
                    indicator, ind_type, rel_data = result
                    # Use a separator that won't conflict with indicator content
                    relationships[f"{indicator}|||{ind_type}"] = rel_data
            except Exception as e:
                print(f"[LogAnalyzer] Error fetching relationships: {e}")
    
    analysis["relationships"] = relationships
    print(f"[LogAnalyzer] Relationship analysis complete - {len(relationships)} indicators with relationships")
    
    # ========== MITRE ATT&CK MAPPING (Phase 4) ==========
    mitre_tactics = map_to_mitre_attack(analysis)
    analysis["mitre_tactics"] = mitre_tactics
    print(f"[LogAnalyzer] MITRE ATT&CK mapping: {mitre_tactics}")
    
    # ========== GENERATE HYPOTHETICAL EXPLANATION (Based on VT & AbuseIPDB) ==========
    def generate_hypothetical_explanation(analysis: dict, original_log_text: str, vt_results: dict, abuseipdb_results: dict) -> str:
        """Generate a hypothetical explanation based on VirusTotal and AbuseIPDB intelligence."""
        explanation_parts = []
        summary_sentences = []
        
        # Use module-level is_private_ip function
        
        # ========== VIRUSTOTAL INTELLIGENCE ==========
        vt_ips = vt_results.get("ips", {})
        vt_hashes = vt_results.get("hashes", {})
        vt_domains = vt_results.get("domains", {})
        
        # Analyze IPs from VirusTotal
        malicious_ips = []
        suspicious_ips = []
        for ip, vt_data in vt_ips.items():
            if isinstance(vt_data, dict) and "error" not in vt_data:
                malicious = vt_data.get("malicious", 0)
                suspicious = vt_data.get("suspicious", 0)
                if malicious > 0:
                    malicious_ips.append((ip, malicious, suspicious))
                elif suspicious > 0:
                    suspicious_ips.append((ip, malicious, suspicious))
        
        if malicious_ips:
            for ip, mal_count, sus_count in malicious_ips[:3]:
                explanation_parts.append(f"🔴 <strong>MALICIOUS IP DETECTED</strong>: <code>{ip}</code> flagged by <strong>{mal_count} security engines</strong> on VirusTotal as malicious.")
                if sus_count > 0:
                    explanation_parts.append(f"Additionally, <strong>{sus_count} engines</strong> marked it as suspicious.")
            summary_sentences.append(f"<strong>Threat Intelligence:</strong> {len(malicious_ips)} malicious IP(s) detected")
        
        if suspicious_ips and not malicious_ips:
            for ip, mal_count, sus_count in suspicious_ips[:3]:
                explanation_parts.append(f"🟡 <strong>SUSPICIOUS IP DETECTED</strong>: <code>{ip}</code> flagged by <strong>{sus_count} security engines</strong> on VirusTotal as suspicious.")
            summary_sentences.append(f"<strong>Threat Intelligence:</strong> {len(suspicious_ips)} suspicious IP(s) detected")
        
        # Analyze Hashes from VirusTotal
        malicious_hashes = []
        suspicious_hashes = []
        for hash_val, vt_data in vt_hashes.items():
            if isinstance(vt_data, dict) and "error" not in vt_data:
                malicious = vt_data.get("malicious", 0)
                suspicious = vt_data.get("suspicious", 0)
                hash_type = next((h.get("type") for h in analysis.get("hashes", []) if h.get("hash") == hash_val), "Hash")
                if malicious > 0:
                    malicious_hashes.append((hash_val[:16] + "...", hash_type, malicious, suspicious))
                elif suspicious > 0:
                    suspicious_hashes.append((hash_val[:16] + "...", hash_type, malicious, suspicious))
        
        if malicious_hashes:
            for hash_short, hash_type, mal_count, sus_count in malicious_hashes[:2]:
                explanation_parts.append(f"🔴 <strong>MALICIOUS FILE HASH DETECTED</strong>: <code>{hash_type}</code> hash <code>{hash_short}</code> flagged by <strong>{mal_count} security engines</strong> as malicious malware.")
            summary_sentences.append(f"<strong>Malware Detection:</strong> {len(malicious_hashes)} malicious file hash(es) identified")
        
        if suspicious_hashes and not malicious_hashes:
            for hash_short, hash_type, mal_count, sus_count in suspicious_hashes[:2]:
                explanation_parts.append(f"🟡 <strong>SUSPICIOUS FILE HASH</strong>: <code>{hash_type}</code> hash <code>{hash_short}</code> flagged by <strong>{sus_count} engines</strong> as suspicious.")
        
        # Analyze Domains from VirusTotal
        malicious_domains = []
        suspicious_domains = []
        for domain, vt_data in vt_domains.items():
            if isinstance(vt_data, dict) and "error" not in vt_data:
                malicious = vt_data.get("malicious", 0)
                suspicious = vt_data.get("suspicious", 0)
                if malicious > 0:
                    malicious_domains.append((domain, malicious, suspicious))
                elif suspicious > 0:
                    suspicious_domains.append((domain, malicious, suspicious))
        
        if malicious_domains:
            for domain, mal_count, sus_count in malicious_domains[:2]:
                explanation_parts.append(f"🔴 <strong>MALICIOUS DOMAIN DETECTED</strong>: <code>{domain}</code> flagged by <strong>{mal_count} security engines</strong> as malicious or associated with threats.")
            summary_sentences.append(f"<strong>Domain Threat:</strong> {len(malicious_domains)} malicious domain(s) identified")
        
        if suspicious_domains and not malicious_domains:
            for domain, mal_count, sus_count in suspicious_domains[:2]:
                explanation_parts.append(f"🟡 <strong>SUSPICIOUS DOMAIN</strong>: <code>{domain}</code> flagged by <strong>{sus_count} engines</strong> as suspicious.")
        
        # ========== ABUSEIPDB INTELLIGENCE ==========
        high_abuse_ips = []
        medium_abuse_ips = []
        for ip, abuse_data in abuseipdb_results.items():
            if isinstance(abuse_data, dict) and "error" not in abuse_data:
                abuse_score = abuse_data.get("abuseScore", 0)
                total_reports = abuse_data.get("totalReports", 0)
                country = abuse_data.get("country", "Unknown")
                isp = abuse_data.get("isp", "Unknown")
                
                try:
                    abuse_score_int = int(abuse_score) if abuse_score != "None" else 0
                    reports_int = int(total_reports) if total_reports != "None" else 0
                except (ValueError, TypeError):
                    abuse_score_int = 0
                    reports_int = 0
                
                if abuse_score_int >= 75 or reports_int >= 10:
                    high_abuse_ips.append((ip, abuse_score_int, reports_int, country, isp))
                elif abuse_score_int >= 25 or reports_int >= 3:
                    medium_abuse_ips.append((ip, abuse_score_int, reports_int, country, isp))
        
        if high_abuse_ips:
            for ip, score, reports, country, isp in high_abuse_ips[:2]:
                explanation_parts.append(f"🔴 <strong>HIGH ABUSE SCORE IP</strong>: <code>{ip}</code> has an abuse confidence score of <strong>{score}%</strong> with <strong>{reports} abuse reports</strong> on AbuseIPDB.")
                if country != "None":
                    explanation_parts.append(f"IP is located in <strong>{country}</strong> and belongs to ISP: <strong>{isp}</strong>.")
            summary_sentences.append(f"<strong>Abuse Intelligence:</strong> {len(high_abuse_ips)} high-risk IP(s) with abuse history")
        
        if medium_abuse_ips and not high_abuse_ips:
            for ip, score, reports, country, isp in medium_abuse_ips[:2]:
                explanation_parts.append(f"🟡 <strong>MODERATE ABUSE SCORE IP</strong>: <code>{ip}</code> has an abuse confidence score of <strong>{score}%</strong> with <strong>{reports} abuse reports</strong>.")
        
        # ========== NETWORK FLOW ANALYSIS (Based on threat intelligence) ==========
        if analysis.get("src_ip") and analysis.get("dst_ip"):
            src_ip = analysis["src_ip"]
            dst_ip = analysis["dst_ip"]
            src_is_private = is_private_ip(src_ip)
            dst_is_private = is_private_ip(dst_ip)
            
            try:
                src_abuse_score = abuseipdb_results.get(src_ip, {}).get("abuseScore", 0)
                dst_abuse_score = abuseipdb_results.get(dst_ip, {}).get("abuseScore", 0)
                src_abuse_int = int(src_abuse_score) if src_abuse_score != "None" else 0
                dst_abuse_int = int(dst_abuse_score) if dst_abuse_score != "None" else 0
            except (ValueError, TypeError):
                src_abuse_int = 0
                dst_abuse_int = 0
            
            src_is_threat = (src_ip in vt_ips and vt_ips[src_ip].get("malicious", 0) > 0) or src_abuse_int >= 75
            dst_is_threat = (dst_ip in vt_ips and vt_ips[dst_ip].get("malicious", 0) > 0) or dst_abuse_int >= 75
            
            if src_is_private and not dst_is_private:
                if dst_is_threat:
                    explanation_parts.append(f"⚠️ <strong>CRITICAL: Internal system connecting to KNOWN THREAT</strong> - Internal IP <code>{src_ip}</code> is connecting to external malicious destination <code>{dst_ip}</code>.")
                    summary_sentences.append("<strong>Threat Level:</strong> CRITICAL - Connection to known malicious destination")
                else:
                    explanation_parts.append(f"Outbound connection from internal IP <code>{src_ip}</code> to external destination <code>{dst_ip}</code>.")
            elif not src_is_private and dst_is_private:
                if src_is_threat:
                    explanation_parts.append(f"⚠️ <strong>CRITICAL: KNOWN THREAT connecting to internal system</strong> - Malicious external IP <code>{src_ip}</code> is attempting to connect to internal destination <code>{dst_ip}</code>.")
                    summary_sentences.append("<strong>Threat Level:</strong> CRITICAL - Inbound connection from known threat")
                else:
                    explanation_parts.append(f"Inbound connection from external IP <code>{src_ip}</code> to internal destination <code>{dst_ip}</code>.")
        
        # ========== PROTOCOL & PORT (Only if relevant to threat) ==========
        if analysis.get("protocol") and (malicious_ips or malicious_hashes or malicious_domains or high_abuse_ips):
            protocol = analysis["protocol"].upper()
            port = analysis.get("dst_port")
            if port:
                explanation_parts.append(f"Connection uses <strong>{protocol}</strong> protocol on port <code>{port}</code>.")
        
        # ========== HTTP METHOD (Only if relevant) ==========
        if analysis.get("http_method") and (malicious_ips or malicious_domains):
            method = analysis["http_method"].upper()
            if method in ["POST", "PUT", "PATCH"]:
                explanation_parts.append(f"HTTP <strong>{method}</strong> method detected, indicating data upload/submission to potentially malicious destination.")
        
        # ========== EXFILTRATION PROBABILITY (Based on threat intelligence) ==========
        if analysis.get("exfil_score", 0) > 50:
            if malicious_ips or malicious_hashes or high_abuse_ips:
                explanation_parts.append(f"🔴 <strong>HIGH EXFILTRATION RISK ({analysis.get('exfil_score', 0)}%)</strong> - Combined with threat intelligence, this strongly indicates potential data exfiltration to malicious destinations.")
            else:
                explanation_parts.append(f"🟡 <strong>MODERATE EXFILTRATION RISK ({analysis.get('exfil_score', 0)}%)</strong> - Event shows characteristics of data transfer activity.")
        
        # ========== USER CONTEXT (If available) ==========
        usernames = []
        if analysis.get("identifiers", {}).get("usernames"):
            usernames.extend(analysis["identifiers"]["usernames"])
        if analysis.get("username") and analysis["username"] not in usernames:
            usernames.append(analysis["username"])
        
        if usernames and (malicious_ips or malicious_hashes or high_abuse_ips):
            if len(usernames) == 1:
                explanation_parts.append(f"⚠️ Activity associated with user <strong>{usernames[0]}</strong> - this user's account may be compromised or involved in malicious activity.")
            else:
                explanation_parts.append(f"⚠️ Activity involving <strong>{len(usernames)} users</strong> - multiple accounts may be affected.")
        
        # ========== OVERALL ASSESSMENT ==========
        if not explanation_parts:
            # Fallback to basic analysis if no threat intelligence available
            if analysis.get("src_ip") and analysis.get("dst_ip"):
                explanation_parts.append(f"Network connection detected between <code>{analysis['src_ip']}</code> and <code>{analysis['dst_ip']}</code>.")
            if analysis.get("protocol"):
                explanation_parts.append(f"Using <strong>{analysis['protocol']}</strong> protocol.")
            explanation_parts.append("No threat intelligence data available from VirusTotal or AbuseIPDB for this event.")
        
        # Combine into final explanation
        explanation = " ".join(explanation_parts)
        
        # Add summary at the top
        if summary_sentences:
            explanation = "<div style='margin-bottom: 12px; padding: 12px; background: rgba(239, 68, 68, 0.15); border-left: 4px solid var(--danger-text); border-radius: 4px;'>" + "<br>".join(summary_sentences) + "</div>" + explanation
        elif not malicious_ips and not malicious_hashes and not malicious_domains and not high_abuse_ips:
            explanation = "<div style='margin-bottom: 12px; padding: 12px; background: rgba(34, 197, 94, 0.1); border-left: 4px solid var(--success-text); border-radius: 4px;'><strong>Threat Intelligence:</strong> No known threats detected in VirusTotal or AbuseIPDB databases</div>" + explanation
        
        return explanation
    
    # Generate hypothetical explanation based on threat intelligence
    analysis["hypothetical_explanation"] = generate_hypothetical_explanation(analysis, log_text, vt_results, abuseipdb_results)
    
    # Build graph data for Threat Constellation Map
    graph_data = build_graph_data(analysis, vt_results, abuseipdb_results, analysis.get("relationships", {}))
    
    # Calculate aggregated threat stats from VT results
    total_malicious = 0
    total_suspicious = 0
    total_harmless = 0
    total_undetected = 0
    
    for ip, result in vt_results["ips"].items():
        if isinstance(result, dict) and "error" not in result:
            total_malicious += result.get("malicious", 0)
            total_suspicious += result.get("suspicious", 0)
            total_harmless += result.get("harmless", 0)
            total_undetected += result.get("undetected", 0)
    
    for hash_val, result in vt_results["hashes"].items():
        if isinstance(result, dict) and "error" not in result:
            total_malicious += result.get("malicious", 0)
            total_suspicious += result.get("suspicious", 0)
            total_harmless += result.get("harmless", 0)
            total_undetected += result.get("undetected", 0)
    
    for domain, result in vt_results["domains"].items():
        if isinstance(result, dict) and "error" not in result:
            total_malicious += result.get("malicious", 0)
            total_suspicious += result.get("suspicious", 0)
            total_harmless += result.get("harmless", 0)
            total_undetected += result.get("undetected", 0)
    
    print(f"[LogAnalyzer] VT lookups complete - Malicious: {total_malicious}, Suspicious: {total_suspicious}")

    news_data = get_cyber_news()
    
    recent_history = get_recent_history(limit=10)
    return render_template(
        "index.html",
        result=None,
        query="",
        vt=None,
        ipdb=None,
        malicious=total_malicious,
        suspicious=total_suspicious,
        harmless=total_harmless,
        undetected=total_undetected,
        timeout=0,
        news=news_data,
        log_analysis_results=analysis,
        graph_data=graph_data,
        recent_history=recent_history
    )
# ================================

# ======= FILE TYPE DETECTION & PDF FORENSICS =======
def analyze_pdf(file_path):
    """
    Analyzes a PDF file for phishing indicators.
    Extracts text, links, IOCs, and performs threat intelligence lookups.
    """
    print(f"[PDF Analysis] Starting analysis of PDF file: {file_path}")
    
    if not PDFMINER_AVAILABLE:
        return {
            "error": "pdfminer.six library not installed. Please run: pip install pdfminer.six"
        }
    
    try:
        # Extract text from PDF
        print("[PDF Analysis] Extracting text from PDF...")
        pdf_text = extract_text(file_path, laparams=LAParams())
        
        # Extract URLs and IPs from PDF text
        print("[PDF Analysis] Extracting IOCs from PDF text...")
        ip_matches = _IP_REGEX_PATTERN.finditer(pdf_text)
        unique_ips = list(set([match.group(0) for match in ip_matches]))
        unique_ips = [ip for ip in unique_ips if not (ip.startswith(('127.', '192.168.', '10.', '172.')) or (ip.startswith('172.') and len(ip.split('.')) > 1 and ip.split('.')[1].isdigit() and 16 <= int(ip.split('.')[1]) <= 31))]
        
        url_matches = _URL_REGEX_PATTERN.findall(pdf_text)
        unique_urls = list(set(url_matches))
        
        print(f"[PDF Analysis] Found {len(unique_ips)} IPs, {len(unique_urls)} URLs")
        
        # Threat Intelligence Lookups
        print("[PDF Analysis] Performing threat intelligence lookups...")
        vt_results_ips = {}
        vt_results_urls = {}
        
        def lookup_ip(ip):
            try:
                result, error = get_cached_data(ip, _vt_fetch_internal, 'vt', max_age_hours=24)
                return ip, result, error
            except Exception as e:
                return ip, None, str(e)
        
        def lookup_url(url):
            try:
                parsed = urlparse(url)
                if not parsed.scheme or not parsed.netloc:
                    return url, None, "Invalid URL format"
                
                # Try full URL lookup first
                try:
                    result, error = vt_fetch(url)
                    if result and not error:
                        return url, result, error
                except:
                    pass
                
                # Fall back to domain lookup
                domain = parsed.netloc
                if domain:
                    result, error = vt_fetch(domain)
                    return url, result, error
                return url, None, "Invalid URL"
            except Exception as e:
                return url, None, str(e)
        
        # Lookup IPs (limit to 10)
        if unique_ips:
            with ThreadPoolExecutor(max_workers=3) as executor:
                ip_futures = {executor.submit(lookup_ip, ip): ip for ip in unique_ips[:10]}
                for future in as_completed(ip_futures):
                    ip, result, error = future.result()
                    if result:
                        vt_results_ips[ip] = result
                    elif error:
                        vt_results_ips[ip] = {"error": error}
        
        # Lookup URLs (limit to 30)
        if unique_urls:
            with ThreadPoolExecutor(max_workers=3) as executor:
                url_futures = {executor.submit(lookup_url, url): url for url in unique_urls[:30]}
                for future in as_completed(url_futures):
                    url, result, error = future.result()
                    if result:
                        vt_results_urls[url] = result
                    elif error:
                        vt_results_urls[url] = {"error": error}
        
        # Calculate score
        score = 0
        
        for ip, result in vt_results_ips.items():
            if isinstance(result, dict) and "error" not in result:
                if result.get("malicious", 0) > 0:
                    score += 100
                elif result.get("suspicious", 0) > 0:
                    score += 50
        
        for url, result in vt_results_urls.items():
            if isinstance(result, dict) and "error" not in result:
                if result.get("malicious", 0) > 0:
                    score += 100
                elif result.get("suspicious", 0) > 0:
                    score += 50
        
        # Format results
        ips_list = []
        for ip in unique_ips[:10]:
            result = vt_results_ips.get(ip, {})
            ips_list.append({
                "ip": ip,
                "malicious": result.get("malicious", 0) if isinstance(result, dict) else 0,
                "suspicious": result.get("suspicious", 0) if isinstance(result, dict) else 0,
                "error": result.get("error") if isinstance(result, dict) and "error" in result else None
            })
        
        urls_list = []
        for url in unique_urls[:30]:
            result = vt_results_urls.get(url, {})
            if isinstance(result, dict):
                malicious = result.get("malicious", 0)
                suspicious = result.get("suspicious", 0)
                harmless = result.get("harmless", 0)
                undetected = result.get("undetected", 0)
                total = malicious + suspicious + harmless + undetected
                
                urls_list.append({
                    "url": url,
                    "malicious": malicious,
                    "suspicious": suspicious,
                    "harmless": harmless,
                    "undetected": undetected,
                    "total": total,
                    "error": result.get("error") if "error" in result else None
                })
            else:
                urls_list.append({
                    "url": url,
                    "malicious": 0,
                    "suspicious": 0,
                    "harmless": 0,
                    "undetected": 0,
                    "total": 0,
                    "error": "No result"
                })
        
        urls_list.sort(key=lambda x: (x.get("malicious", 0) * 1000 + x.get("suspicious", 0)), reverse=True)
        
        # Determine verdict
        if score >= 100:
            verdict = "MALICIOUS"
        elif score >= 50:
            verdict = "SUSPICIOUS"
        else:
            verdict = "CLEAN"
        
        body_preview = pdf_text[:500] if pdf_text else ""
        
        print(f"[PDF Analysis] Analysis complete. Verdict: {verdict}, Score: {score}")
        
        return {
            "headers": {
                "Subject": "PDF File Analysis",
                "From": "N/A",
                "To": "N/A",
                "Date": "N/A",
                "Message-ID": "N/A"
            },
            "body_preview": body_preview,
            "urls": urls_list,
            "ips": ips_list,
            "keywords_found": [],
            "attachments": [],
            "score": score,
            "verdict": verdict,
            "total_urls_found": len(unique_urls),
            "urls_checked": min(len(unique_urls), 30),
            "file_type": "PDF"
        }
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"[PDF Analysis] Error analyzing PDF:")
        print(error_trace)
        return {
            "error": f"Error analyzing PDF file: {str(e)}"
        }

def detect_and_route_file(file_path):
    """
    Detects file type using magic bytes and determines routing.
    Returns: (file_type, extension_mismatch, spoofing_warning, spoofing_score)
    file_type: 'pdf', 'msg', or 'unknown'
    """
    print(f"[File Detection] Analyzing file: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    # Read first 2048 bytes for magic byte detection
    try:
        with open(file_path, 'rb') as f:
            magic_bytes = f.read(2048)
    except Exception as e:
        print(f"[File Detection] Error reading file: {e}")
        return 'unknown', False, None, 0
    
    detected_type = 'unknown'
    extension_mismatch = False
    spoofing_warning = None
    spoofing_score = 0
    
    # Priority 1: Check for PDF signature (%PDF) - must check first!
    if magic_bytes.startswith(b'%PDF'):
        detected_type = 'pdf'
        print("[File Detection] ✓ Detected PDF via signature (%PDF at start)")
    elif b'%PDF' in magic_bytes[:2048]:  # Check entire buffer for PDF (might have metadata)
        detected_type = 'pdf'
        pdf_pos = magic_bytes.find(b'%PDF')
        print(f"[File Detection] ✓ Detected PDF via signature (found %PDF at position {pdf_pos})")
    
    # Priority 2: Strict OLE structure validation using olefile - only if not PDF
    if detected_type != 'pdf' and OLEFILE_AVAILABLE:
        try:
            if olefile.isOleFile(file_path):
                detected_type = 'msg'
                extension_mismatch = False
                spoofing_warning = None
                spoofing_score = 0
                print("[File Detection] ✓ Detected valid OLE structure via olefile.isOleFile() - definitive MSG file")
                return detected_type, extension_mismatch, spoofing_warning, spoofing_score
            else:
                print("[File Detection] ⚠️ olefile.isOleFile() returned False - file is NOT a valid OLE container")
        except Exception as e:
            print(f"[File Detection] olefile validation error: {e}")
    
    # Priority 3: Try filetype library - only if not PDF and not OLE
    if detected_type == 'unknown' and FILETYPE_AVAILABLE:
        try:
            kind = filetype.guess(magic_bytes)
            if kind:
                filetype_detected = kind.extension
                filetype_mime = kind.mime
                
                # If filetype detects PDF, trust it (even over extension)
                if filetype_mime == 'application/pdf' or filetype_detected == 'pdf':
                    detected_type = 'pdf'
                    print(f"[File Detection] ✓ Detected PDF via filetype library: {filetype_mime}")
                # If filetype detects MSG/OLE, trust it (but only if OLE validation wasn't available)
                elif ('outlook' in filetype_mime or filetype_detected == 'msg') and not OLEFILE_AVAILABLE:
                    detected_type = 'msg'
                    print(f"[File Detection] ✓ Detected MSG via filetype library: {filetype_mime}")
                # Otherwise, use filetype result if we don't have a detection yet
                elif detected_type == 'unknown':
                    detected_type = filetype_detected
                    print(f"[File Detection] ✓ Detected type via filetype library: {filetype_detected} ({filetype_mime})")
        except Exception as e:
            print(f"[File Detection] filetype library error: {e}")
    
    # Priority 4: Fallback to extension ONLY if still unknown and extension matches expected type
    # NOTE: We do NOT fallback to .msg extension if OLE validation failed - that would be unsafe
    if detected_type == 'unknown':
        if file_extension == '.pdf':
            detected_type = 'pdf'
            print(f"[File Detection] ⚠️ No magic bytes found, falling back to extension: {file_extension}")
        elif file_extension == '.msg' and not OLEFILE_AVAILABLE:
            # Only fallback to .msg if olefile is not available (shouldn't happen in production)
            detected_type = 'msg'
            print(f"[File Detection] ⚠️ No magic bytes found and olefile unavailable, falling back to extension: {file_extension}")
    
    # Check for extension mismatch (spoofing)
    if file_extension == '.msg' and detected_type == 'pdf':
        extension_mismatch = True
        spoofing_warning = "🚨 CRITICAL: EXTENSION SPOOFING DETECTED! File has .msg extension but is actually a PDF. This is a common phishing technique."
        spoofing_score = 100  # Critical spoofing
        print("[File Detection] 🚨 CRITICAL SPOOFING: .msg extension but PDF content detected!")
    elif file_extension == '.pdf' and detected_type == 'msg':
        extension_mismatch = True
        spoofing_warning = f"⚠️ Extension Mismatch: File has .pdf extension but is actually MSG format."
        spoofing_score = 20
    elif file_extension and detected_type != 'unknown' and file_extension != f'.{detected_type}':
        extension_mismatch = True
        spoofing_warning = f"⚠️ Extension Mismatch: File has {file_extension} extension but detected type is {detected_type}."
        spoofing_score = 20
        print(f"[File Detection] ⚠️ Extension mismatch: {file_extension} vs {detected_type}")
    
    return detected_type, extension_mismatch, spoofing_warning, spoofing_score

# ======= CRASH-PROOF MSG ANALYZER =======
def _get_header(msg_obj, key, default="N/A"):
    """
    Helper function to safely extract header values.
    Uses getattr and converts to str immediately.
    """
    value = getattr(msg_obj, key, default)
    return str(value) if value is not None else default

def extract_data_from_broken_ole(file_path):
    """
    Surgical OLE Extraction: Safely extracts data from corrupted OLE files.
    Iterates through all streams in the OLE container, skipping corrupted ones.
    Returns a structured result dictionary with extracted IOCs.
    """
    print(f"[Surgical OLE Extraction] Starting extraction from broken OLE file: {file_path}")
    
    if not OLEFILE_AVAILABLE:
        print("[Surgical OLE Extraction] olefile not available, falling back to raw extraction")
        return None
    
    all_extracted_text = ""
    ole = None
    
    try:
        # Open OLE file
        ole = olefile.OleFileIO(file_path)
        print("[Surgical OLE Extraction] Successfully opened OLE file")
        
        # Iterate through all streams
        stream_list = ole.listdir()
        print(f"[Surgical OLE Extraction] Found {len(stream_list)} stream(s) to process")
        
        for stream_path in stream_list:
            stream_name = '/'.join(stream_path) if isinstance(stream_path, tuple) else str(stream_path)
            print(f"[Surgical OLE Extraction] Attempting to read stream: {stream_name}")
            
            try:
                # Try to read the stream
                stream_data = ole.openstream(stream_path).read()
                
                # Try to decode as UTF-16-LE (common in MSG files)
                try:
                    decoded_text = stream_data.decode('utf-16-le', errors='ignore')
                    if len(decoded_text) > 10:  # Only use if we got substantial content
                        all_extracted_text += decoded_text + "\n"
                        print(f"[Surgical OLE Extraction] ✓ Decoded stream {stream_name} as UTF-16-LE ({len(decoded_text)} chars)")
                        continue
                except:
                    pass
                
                # Try to decode as UTF-8
                try:
                    decoded_text = stream_data.decode('utf-8', errors='ignore')
                    if len(decoded_text) > 10:
                        all_extracted_text += decoded_text + "\n"
                        print(f"[Surgical OLE Extraction] ✓ Decoded stream {stream_name} as UTF-8 ({len(decoded_text)} chars)")
                        continue
                except:
                    pass
                
                # If decoding failed, try to extract printable ASCII strings
                ascii_pattern = re.compile(rb'[\x20-\x7E]{6,}')
                ascii_matches = ascii_pattern.findall(stream_data)
                if ascii_matches:
                    ascii_text = b' '.join(ascii_matches[:100]).decode('ascii', errors='ignore')
                    all_extracted_text += ascii_text + "\n"
                    print(f"[Surgical OLE Extraction] ✓ Extracted ASCII strings from stream {stream_name} ({len(ascii_text)} chars)")
                
            except Exception as stream_error:
                # Skip corrupted streams and continue
                print(f"[Surgical OLE Extraction] ⚠️ Stream {stream_name} is corrupted, skipping: {stream_error}")
                continue
        
        print(f"[Surgical OLE Extraction] Total extracted text: {len(all_extracted_text)} characters")
        
        # Extract IOCs from the accumulated text
        ip_matches = _IP_REGEX_PATTERN.finditer(all_extracted_text)
        unique_ips = list(set([match.group(0) for match in ip_matches]))
        unique_ips = [ip for ip in unique_ips if not (ip.startswith(('127.', '192.168.', '10.', '172.')) or (ip.startswith('172.') and len(ip.split('.')) > 1 and ip.split('.')[1].isdigit() and 16 <= int(ip.split('.')[1]) <= 31))]
        
        url_matches = _URL_REGEX_PATTERN.findall(all_extracted_text)
        unique_urls = list(set(url_matches))
        
        email_matches = _EMAIL_REGEX_PATTERN.findall(all_extracted_text)
        unique_emails = list(set(email_matches))
        
        print(f"[Surgical OLE Extraction] Found {len(unique_ips)} IPs, {len(unique_urls)} URLs, {len(unique_emails)} emails")
        
        # Perform basic threat intelligence lookups (simplified)
        ips_list = []
        for ip in unique_ips[:10]:
            ips_list.append({
                "ip": str(ip),
                "malicious": 0,
                "suspicious": 0,
                "error": "Not checked in partial mode"
            })
        
        urls_list = []
        for url in unique_urls[:30]:
            urls_list.append({
                "url": str(url),
                "malicious": 0,
                "suspicious": 0,
                "harmless": 0,
                "undetected": 0,
                "total": 0,
                "error": "Not checked in partial mode"
            })
        
        # Basic scoring
        score = 0
        if unique_ips or unique_urls:
            score += 10  # Indicate IOCs found
        
        body_preview = all_extracted_text[:500] if all_extracted_text else ""
        
        return {
            "headers": {
                "Subject": "PARTIAL (CORRUPTED)",
                "From": "N/A",
                "To": "N/A",
                "Date": "N/A",
                "Message-ID": "N/A"
            },
            "body_preview": str(body_preview),
            "urls": urls_list,
            "ips": ips_list,
            "keywords_found": [],
            "attachments": [],
            "score": int(score),
            "verdict": "PARTIAL (CORRUPTED)",
            "total_urls_found": int(len(unique_urls)),
            "urls_checked": int(min(len(unique_urls), 30)),
            "file_type": "MSG",
            "status": "partial",
            "warning": "⚠️ File structure is corrupted. Extracted data from valid OLE streams only."
        }
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"[Surgical OLE Extraction] Error during OLE extraction: {e}")
        print(f"[Surgical OLE Extraction] Traceback: {error_trace}")
        return None  # Signal failure to caller
    
    finally:
        # Ensure OLE file is closed
        if ole is not None:
            try:
                ole.close()
                print("[Surgical OLE Extraction] Closed OLE file")
            except:
                pass

def extract_ioc_from_raw_file(file_path):
    """
    Fail-Safe Raw Extraction Mode: Bypasses extract-msg entirely.
    Extracts IOCs from binary file content using string extraction and regex.
    Returns a structured dictionary compatible with the analysis result format.
    """
    print(f"[Raw Extraction] Starting raw IOC extraction from file: {file_path}")
    
    try:
        # Step 1: Read file in binary mode
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        print(f"[Raw Extraction] Read {len(raw_data)} bytes from file")
        
        # Step 2: String Extraction - Try multiple encodings
        extracted_strings = []
        text_content = ""
        
        # Try standard encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                decoded = raw_data.decode(encoding, errors='ignore')
                if len(decoded) > 100:  # Only use if we got substantial content
                    text_content = decoded
                    print(f"[Raw Extraction] Successfully decoded using {encoding} ({len(decoded)} chars)")
                    break
            except Exception as e:
                print(f"[Raw Extraction] Failed to decode with {encoding}: {e}")
                continue
        
        # Step 3: If decoding failed, extract printable ASCII strings using regex
        if not text_content or len(text_content) < 50:
            print("[Raw Extraction] Decoding failed, extracting ASCII strings using regex...")
            # Find sequences of printable ASCII characters (6+ chars)
            ascii_pattern = re.compile(rb'[\x20-\x7E]{6,}')
            ascii_matches = ascii_pattern.findall(raw_data)
            if ascii_matches:
                # Join and decode as ASCII
                text_content = b' '.join(ascii_matches[:1000]).decode('ascii', errors='ignore')  # Limit to avoid huge strings
                print(f"[Raw Extraction] Extracted {len(ascii_matches)} ASCII string sequences ({len(text_content)} chars)")
        
        if not text_content or len(text_content) < 10:
            print("[Raw Extraction] ⚠️ Could not extract sufficient text content")
            return {
                "status": "partial",
                "warning": "Parsed as raw data due to corruption. Could not extract readable text content.",
                "headers": {
                    "Subject": "N/A",
                    "From": "N/A",
                    "To": "N/A",
                    "Date": "N/A",
                    "Message-ID": "N/A"
                },
                "body_preview": "",
                "urls": [],
                "ips": [],
                "keywords_found": [],
                "attachments": [],
                "score": 0,
                "verdict": "UNKNOWN",
                "total_urls_found": 0,
                "urls_checked": 0,
                "file_type": "MSG"
            }
        
        # Step 4: Extract IOCs using existing regex patterns
        print("[Raw Extraction] Extracting IPs and URLs from text content...")
        
        # Extract IPs
        ip_matches = _IP_REGEX_PATTERN.finditer(text_content)
        unique_ips = list(set([match.group(0) for match in ip_matches]))
        # Filter private IPs
        unique_ips = [ip for ip in unique_ips if not (ip.startswith(('127.', '192.168.', '10.', '172.')) or (ip.startswith('172.') and len(ip.split('.')) > 1 and ip.split('.')[1].isdigit() and 16 <= int(ip.split('.')[1]) <= 31))]
        
        # Extract URLs
        url_matches = _URL_REGEX_PATTERN.findall(text_content)
        unique_urls = list(set(url_matches))
        
        print(f"[Raw Extraction] Found {len(unique_ips)} IPs and {len(unique_urls)} URLs")
        
        # Step 5: Extract basic headers from text (if present)
        headers = {
            "Subject": "N/A",
            "From": "N/A",
            "To": "N/A",
            "Date": "N/A",
            "Message-ID": "N/A"
        }
        
        # Try to extract headers using regex
        subject_match = re.search(r'(?i)subject:\s*([^\r\n]+)', text_content)
        if subject_match:
            headers["Subject"] = str(subject_match.group(1).strip()[:200])  # Limit length
        
        from_match = re.search(r'(?i)from:\s*([^\r\n]+)', text_content)
        if from_match:
            headers["From"] = str(from_match.group(1).strip()[:200])
        
        to_match = re.search(r'(?i)to:\s*([^\r\n]+)', text_content)
        if to_match:
            headers["To"] = str(to_match.group(1).strip()[:200])
        
        date_match = re.search(r'(?i)date:\s*([^\r\n]+)', text_content)
        if date_match:
            headers["Date"] = str(date_match.group(1).strip()[:200])
        
        # Step 6: Perform threat intelligence lookups (limited to avoid API overload)
        print("[Raw Extraction] Performing threat intelligence lookups...")
        vt_results_ips = {}
        vt_results_urls = {}
        
        def lookup_ip(ip):
            try:
                result, error = get_cached_data(ip, _vt_fetch_internal, 'vt', max_age_hours=24)
                return ip, result, error
            except Exception as e:
                return ip, None, str(e)
        
        def lookup_url(url):
            try:
                parsed = urlparse(url)
                if not parsed.scheme or not parsed.netloc:
                    return url, None, "Invalid URL format"
                
                try:
                    result, error = vt_fetch(url)
                    if result and not error:
                        return url, result, error
                except:
                    pass
                
                domain = parsed.netloc
                if domain:
                    result, error = vt_fetch(domain)
                    return url, result, error
                return url, None, "Invalid URL"
            except Exception as e:
                return url, None, str(e)
        
        # Lookup IPs (limit to 5 for raw extraction)
        if unique_ips:
            with ThreadPoolExecutor(max_workers=2) as executor:
                ip_futures = {executor.submit(lookup_ip, ip): ip for ip in unique_ips[:5]}
                for future in as_completed(ip_futures):
                    ip, result, error = future.result()
                    if result:
                        vt_results_ips[ip] = result
                    elif error:
                        vt_results_ips[ip] = {"error": str(error)}
        
        # Lookup URLs (limit to 10 for raw extraction)
        if unique_urls:
            with ThreadPoolExecutor(max_workers=2) as executor:
                url_futures = {executor.submit(lookup_url, url): url for url in unique_urls[:10]}
                for future in as_completed(url_futures):
                    url, result, error = future.result()
                    if result:
                        vt_results_urls[url] = result
                    elif error:
                        vt_results_urls[url] = {"error": str(error)}
        
        # Step 7: Calculate score
        score = 0
        for ip, result in vt_results_ips.items():
            if isinstance(result, dict) and "error" not in result:
                malicious = result.get("malicious", 0) or 0
                suspicious = result.get("suspicious", 0) or 0
                if malicious > 0:
                    score += 100
                elif suspicious > 0:
                    score += 50
        
        for url, result in vt_results_urls.items():
            if isinstance(result, dict) and "error" not in result:
                malicious = result.get("malicious", 0) or 0
                suspicious = result.get("suspicious", 0) or 0
                if malicious > 0:
                    score += 100
                elif suspicious > 0:
                    score += 50
        
        # Step 8: Format results (ensure all values are JSON-serializable)
        ips_list = []
        for ip in unique_ips[:5]:
            result = vt_results_ips.get(ip, {})
            ips_list.append({
                "ip": str(ip),
                "malicious": int(result.get("malicious", 0)) if isinstance(result, dict) and "error" not in result else 0,
                "suspicious": int(result.get("suspicious", 0)) if isinstance(result, dict) and "error" not in result else 0,
                "error": str(result.get("error")) if isinstance(result, dict) and "error" in result else None
            })
        
        urls_list = []
        for url in unique_urls[:10]:
            result = vt_results_urls.get(url, {})
            if isinstance(result, dict) and "error" not in result:
                urls_list.append({
                    "url": str(url),
                    "malicious": int(result.get("malicious", 0)),
                    "suspicious": int(result.get("suspicious", 0)),
                    "harmless": int(result.get("harmless", 0)),
                    "undetected": int(result.get("undetected", 0)),
                    "total": int(result.get("malicious", 0) + result.get("suspicious", 0) + result.get("harmless", 0) + result.get("undetected", 0)),
                    "error": None
                })
            else:
                urls_list.append({
                    "url": str(url),
                    "malicious": 0,
                    "suspicious": 0,
                    "harmless": 0,
                    "undetected": 0,
                    "total": 0,
                    "error": str(result.get("error")) if isinstance(result, dict) and "error" in result else "No result"
                })
        
        urls_list.sort(key=lambda x: (x.get("malicious", 0) * 1000 + x.get("suspicious", 0)), reverse=True)
        
        # Determine verdict
        if score >= 100:
            verdict = "MALICIOUS"
        elif score >= 50:
            verdict = "SUSPICIOUS"
        else:
            verdict = "CLEAN"
        
        body_preview = str(text_content[:500]) if text_content else ""
        
        print(f"[Raw Extraction] Raw extraction complete. Verdict: {verdict}, Score: {score}")
        
        return {
            "status": "partial",
            "warning": "Parsed as raw data due to corruption. Standard MSG parsing failed, but IOCs were extracted from binary content.",
            "headers": headers,
            "body_preview": body_preview,
            "urls": urls_list,
            "ips": ips_list,
            "keywords_found": [],
            "attachments": [],
            "score": int(score),
            "verdict": str(verdict),
            "total_urls_found": int(len(unique_urls)),
            "urls_checked": int(min(len(unique_urls), 10)),
            "file_type": "MSG"
        }
    
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"[Raw Extraction] ERROR during raw extraction:")
        print(error_trace)
        print(f"[Raw Extraction] Error: {str(e)}")
        
        return {
            "status": "error",
            "error": f"Raw extraction failed: {str(e)}",
            "headers": {
                "Subject": "N/A",
                "From": "N/A",
                "To": "N/A",
                "Date": "N/A",
                "Message-ID": "N/A"
            },
            "body_preview": "",
            "urls": [],
            "ips": [],
            "keywords_found": [],
            "attachments": [],
            "score": 0,
            "verdict": "UNKNOWN",
            "total_urls_found": 0,
            "urls_checked": 0,
            "file_type": "MSG"
        }

def analyze_msg_robust(file_path):
    """
    Robust MSG file analyzer using msg-parser library (Python equivalent of msg-reader).
    Inspired by: https://github.com/Rasalas/msg-reader
    
    This implementation uses msg-parser which:
    - Works cross-platform (Linux, Windows, macOS)
    - Handles corrupted MSG files gracefully
    - Extracts headers, body, and attachments reliably
    - Converts MSG to JSON format for easy parsing
    
    Note: For .eml files, use extract_email_attachments instead.
    """
    print(f"[MSG Reader] Starting robust analysis of MSG file: {file_path}")
    
    # Check file extension - msg-parser only works with .msg files
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.eml':
        print("[MSG Reader] EML file detected - msg-parser does not support EML files")
        print("[MSG Reader] EML files should be processed via extract_email_attachments")
        # Return empty result for .eml files - they will be handled by extract_email_attachments
        return {
            "headers": {},
            "body_preview": "",
            "urls": [],
            "ips": [],
            "keywords_found": [],
            "attachments": [],
            "score": 0,
            "verdict": "CLEAN",
            "total_urls_found": 0,
            "urls_checked": 0,
            "file_type": "EML",
            "warning": "EML files should be processed via extract_email_attachments, not analyze_msg_robust"
        }
    
    # Try msg-parser first (preferred, more robust) - only for .msg files
    if MSG_PARSER_AVAILABLE:
        try:
            print("[MSG Reader] Using msg-parser library (msg-reader equivalent)...")
            msg = MsOxMessage(file_path)
            
            # Get message as JSON for structured parsing
            try:
                msg_json = json.loads(msg.get_message_as_json())
            except:
                msg_json = {}
            
            # Extract headers
            headers = {
                "Subject": msg_json.get('subject', 'N/A'),
                "From": msg_json.get('sender', {}).get('email', 'N/A') if isinstance(msg_json.get('sender'), dict) else str(msg_json.get('sender', 'N/A')),
                "To": ', '.join([r.get('email', '') for r in msg_json.get('recipients', []) if isinstance(r, dict)]) or 'N/A',
                "Date": msg_json.get('date', 'N/A'),
                "Message-ID": msg_json.get('messageId', 'N/A')
            }
            
            # Extract body (prefer HTML, fallback to plain text)
            body_text = ''
            html_body = ''
            if msg_json.get('body'):
                body_text = str(msg_json['body'])
            if msg_json.get('htmlBody'):
                html_body = str(msg_json['htmlBody'])
            
            # Combine bodies for IOC extraction
            all_text = (body_text + ' ' + html_body).strip()
            body_preview = body_text[:500] if body_text else (html_body[:500] if html_body else '')
            
            # Extract IOCs from body
            unique_ips = []
            unique_urls = []
            if all_text:
                ip_matches = _IP_REGEX_PATTERN.finditer(all_text)
                unique_ips = list(set([match.group(0) for match in ip_matches]))
                unique_ips = [ip for ip in unique_ips if not is_private_ip(ip)]
                
                url_matches = _URL_REGEX_PATTERN.findall(all_text)
                unique_urls = list(set([u for u in url_matches if u.startswith(('http://', 'https://'))]))
                
                # Extract URLs from HTML
                if html_body:
                    try:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(html_body, 'html.parser')
                        for link in soup.find_all('a', href=True):
                            href = link.get('href', '')
                            if href.startswith(('http://', 'https://')):
                                unique_urls.append(href)
                    except:
                        pass
            
            # Extract attachments
            attachments_list = []
            attachments_data = msg_json.get('attachments', [])
            
            for idx, att_data in enumerate(attachments_data):
                try:
                    att_name = att_data.get('name', f'Attachment_{idx + 1}')
                    att_content = att_data.get('content', b'')
                    
                    # Calculate hash
                    if isinstance(att_content, str):
                        att_content = att_content.encode('utf-8', errors='ignore')
                    att_hash = hashlib.sha256(att_content).hexdigest() if att_content else None
                    
                    # Detect file type
                    detected_mime = None
                    detected_ext = None
                    detected_type = None
                    
                    if FILETYPE_AVAILABLE and att_content:
                        try:
                            kind = filetype.guess(att_content)
                            if kind:
                                detected_mime = kind.mime
                                detected_ext = kind.extension
                                detected_type = kind.extension.upper()
                        except:
                            pass
                    
                    # Extract text from attachments if DOCX/PDF
                    extracted_text = None
                    extracted_urls = []
                    
                    if detected_type and detected_type.lower() == 'docx' and att_content:
                        try:
                            extracted_text = extract_text_from_docx(att_content)
                            extracted_urls = extract_urls_from_docx(att_content)
                        except:
                            pass
                    elif detected_type and detected_type.lower() == 'pdf' and PDFMINER_AVAILABLE and att_content:
                        try:
                            # Cache PDF for HTML viewing later
                            cache_pdf_attachment(att_hash, att_content, att_name)
                            # Also cache in general attachment cache for client-side viewing
                            cache_attachment(att_hash, att_content, att_name, detected_mime)
                            
                            # Save to temp file for PDFMiner
                            fd, tmp_path = tempfile.mkstemp(suffix='.pdf')
                            try:
                                with os.fdopen(fd, 'wb') as f:
                                    f.write(att_content)
                                extracted_text = extract_text(tmp_path, laparams=LAParams())
                                # Extract URLs from PDF text
                                if extracted_text:
                                    url_matches = _URL_REGEX_PATTERN.findall(extracted_text)
                                    extracted_urls = list(set(url_matches))
                            finally:
                                try:
                                    os.unlink(tmp_path)
                                except:
                                    pass
                        except:
                            pass
                    else:
                        # Cache other attachment types for client-side viewing
                        if att_hash and att_content:
                            cache_attachment(att_hash, att_content, att_name, detected_mime)
                    
                    # VirusTotal lookup for attachment hash
                    vt_stats = {}
                    if att_hash:
                        try:
                            vt_result, vt_error = vt_fetch(att_hash)
                            if vt_result:
                                vt_stats = {k: str(v) for k, v in vt_result.items()}
                        except:
                            pass
                    
                    attachments_list.append({
                        "filename": str(att_name),
                        "hash": str(att_hash) if att_hash else None,
                        "detected_mime": str(detected_mime) if detected_mime else None,
                        "detected_ext": str(detected_ext) if detected_ext else None,
                        "detected_type": str(detected_type) if detected_type else None,
                        "extension_mismatch": False,
                        "spoofing_warning": None,
                        "has_dangerous_ext": bool(att_name.lower().endswith(('.exe', '.bat', '.ps1', '.vbs', '.js', '.jar', '.scr', '.cmd', '.dll', '.com', '.pif'))),
                        "vt_stats": vt_stats,
                        "extracted_urls": extracted_urls,
                        "extracted_text": extracted_text[:50000] if extracted_text else None,  # Limit size
                        "error": None
                    })
                    
                except Exception as att_error:
                    print(f"[MSG Reader] ⚠️ Error processing attachment {idx}: {att_error}")
                    attachments_list.append({
                        "filename": f"Attachment_{idx + 1}",
                        "hash": None,
                        "error": str(att_error)
                    })
            
            # Calculate score based on attachments
            score = 0
            for att in attachments_list:
                vt_stats = att.get('vt_stats', {})
                if isinstance(vt_stats, dict):
                    malicious = int(vt_stats.get('malicious', 0) or 0)
                    suspicious = int(vt_stats.get('suspicious', 0) or 0)
                    if malicious > 0:
                        score += 100
                    elif suspicious > 0:
                        score += 50
                
                if att.get('has_dangerous_ext'):
                    score += 25
            
            # Determine verdict
            if score >= 100:
                verdict = "MALICIOUS"
            elif score >= 50:
                verdict = "SUSPICIOUS"
            elif len(attachments_list) > 0 or len(unique_urls) > 0:
                verdict = "CLEAN"
            else:
                verdict = "UNKNOWN"
            
            # Format URLs and IPs for return
            urls_list = [{"url": u, "malicious": 0, "suspicious": 0, "harmless": 0, "undetected": 0, "total": 0} for u in unique_urls[:30]]
            ips_list = [{"ip": ip, "malicious": 0, "suspicious": 0} for ip in unique_ips[:10]]
            
            print(f"[MSG Reader] ✓ Analysis complete: {len(attachments_list)} attachments, {len(unique_urls)} URLs, {len(unique_ips)} IPs")
            
            return {
                "headers": {k: str(v) for k, v in headers.items()},
                "body_preview": str(body_preview),
                "urls": urls_list,
                "ips": ips_list,
                "keywords_found": [],
                "attachments": attachments_list,
                "score": int(score),
                "verdict": str(verdict),
                "total_urls_found": int(len(unique_urls)),
                "urls_checked": int(len(urls_list)),
                "file_type": "MSG"
            }
            
        except Exception as msg_parser_error:
            print(f"[MSG Reader] ⚠️ msg-parser failed: {msg_parser_error}")
            print(f"[MSG Reader] Falling back to extract-msg...")
            traceback.print_exc()
            # Fall through to extract-msg fallback
    
    # Fallback to extract-msg if msg-parser not available or failed
        if not EXTRACT_MSG_AVAILABLE:
            return {
            "error": "No MSG parser available. Please install: pip install msg-parser or pip install extract-msg"
            }
        
    # Fallback: Use extract-msg library (original implementation)
    print("[MSG Reader] Using extract-msg as fallback...")
    
    # CRITICAL PRE-CHECK: Validate OLE structure before attempting to parse
    if OLEFILE_AVAILABLE:
        try:
            if not olefile.isOleFile(file_path):
                print("[MSG Reader] ❌ OLE validation failed: File is not a valid OLE container")
                return {
                    "error": "Invalid file structure. This file is not a valid Outlook OLE container. The file may be corrupted or not a valid .msg file."
                }
            else:
                print("[MSG Reader] ✓ OLE validation passed: File is a valid OLE container")
        except Exception as ole_check_error:
            print(f"[MSG Reader] ⚠️ OLE validation error: {ole_check_error}")
            # Continue anyway if validation check itself fails (but log the warning)
    
    # Initialize defaults
    headers = {
            "Subject": "N/A",
            "From": "N/A",
            "To": "N/A",
            "Date": "N/A",
            "Message-ID": "N/A"
        }
    body = ''
    html_body = ''
    body_preview = ''
    unique_ips = []
    unique_urls = []
    keywords_found = []
    attachments_list = []
    score = 0
    temp_attachment_dir = None
    msg = None  # Initialize msg to None for finally block
    
    # Try to load message - use direct try/finally for cleanup
    print("[MSG Reader] Loading MSG file with extract-msg...")
    try:
        # Use extract_msg.Message directly with try/finally for cleanup
        msg = extract_msg.Message(file_path)
        print("[MSG Reader] ✓ Successfully loaded MSG file")
        
        # Extract headers using safe helper function
        print("[MSG Reader] Extracting headers using safe helper...")
        headers = {
            "Subject": _get_header(msg, 'subject'),
            "From": _get_header(msg, 'sender'),
            "To": _get_header(msg, 'to'),
            "Date": _get_header(msg, 'date'),
            "Message-ID": _get_header(msg, 'messageId')
        }
        
        # Safe body extraction
        print("[MSG Reader] Extracting body content safely...")
        body_raw = getattr(msg, 'body', None)
        body = '' if body_raw is None else (body_raw.decode('utf-8', errors='ignore') if isinstance(body_raw, bytes) else str(body_raw))
        
        html_body_raw = getattr(msg, 'htmlBody', None)
        html_body = '' if html_body_raw is None else (html_body_raw.decode('utf-8', errors='ignore') if isinstance(html_body_raw, bytes) else str(html_body_raw))
        
        # Combine both bodies for IOC extraction
        all_text = body + ' ' + html_body
        body_preview = body[:500] if body else ''
        
        # Extract IOCs
        print("[MSG Reader] Extracting IOCs...")
        ip_matches = _IP_REGEX_PATTERN.finditer(all_text)
        unique_ips = list(set([match.group(0) for match in ip_matches]))
        unique_ips = [ip for ip in unique_ips if not (ip.startswith(('127.', '192.168.', '10.', '172.')) or (ip.startswith('172.') and len(ip.split('.')) > 1 and ip.split('.')[1].isdigit() and 16 <= int(ip.split('.')[1]) <= 31))]
        
        url_matches = _URL_REGEX_PATTERN.findall(all_text)
        unique_urls = list(set(url_matches))
        
        # Extract URLs from HTML
        if html_body:
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_body, 'html.parser')
                for link in soup.find_all('a', href=True):
                    href = link.get('href', '')
                    if href.startswith(('http://', 'https://')):
                        unique_urls.append(href)
            except:
                    pass
            
            unique_urls = list(set(unique_urls))
            
            # Initialize totals
            total_urls_found = len(unique_urls)
            urls_checked = min(len(unique_urls), 30)
            
            # Threat Intelligence Lookups
            print("[MSG Reader] Performing threat intelligence lookups...")
            vt_results_ips = {}
            vt_results_urls = {}
            
            def lookup_ip(ip):
                try:
                    result, error = get_cached_data(ip, _vt_fetch_internal, 'vt', max_age_hours=24)
                    return ip, result, error
                except Exception as e:
                    return ip, None, str(e)
            
            def lookup_url(url):
                try:
                    parsed = urlparse(url)
                    if not parsed.scheme or not parsed.netloc:
                        return url, None, "Invalid URL format"
                    
                    try:
                        result, error = vt_fetch(url)
                        if result and not error:
                            return url, result, error
                    except:
                        pass
                    
                    domain = parsed.netloc
                    if domain:
                        result, error = vt_fetch(domain)
                        return url, result, error
                    return url, None, "Invalid URL"
                except Exception as e:
                    return url, None, str(e)
            
            # Lookup IPs
            if unique_ips:
                with ThreadPoolExecutor(max_workers=3) as executor:
                    ip_futures = {executor.submit(lookup_ip, ip): ip for ip in unique_ips[:10]}
                    for future in as_completed(ip_futures):
                        ip, result, error = future.result()
                        if result:
                            vt_results_ips[ip] = result
                        elif error:
                            vt_results_ips[ip] = {"error": error}
            
            # Lookup URLs
            if unique_urls:
                with ThreadPoolExecutor(max_workers=3) as executor:
                    url_futures = {executor.submit(lookup_url, url): url for url in unique_urls[:30]}
                    for future in as_completed(url_futures):
                        url, result, error = future.result()
                        if result:
                            vt_results_urls[url] = result
                        elif error:
                            vt_results_urls[url] = {"error": error}
            
            # Keyword Analysis
            body_lower = body.lower()
            keywords_found = []
            for keyword in PHISHING_KEYWORDS:
                count = body_lower.count(keyword.lower())
                if count > 0:
                    keywords_found.append({"keyword": keyword, "count": count})
            
            # Calculate score
            score = 0
            for ip, result in vt_results_ips.items():
                if isinstance(result, dict) and "error" not in result:
                    if result.get("malicious", 0) > 0:
                        score += 100
                    elif result.get("suspicious", 0) > 0:
                        score += 50
            
            for url, result in vt_results_urls.items():
                if isinstance(result, dict) and "error" not in result:
                    if result.get("malicious", 0) > 0:
                        score += 100
                    elif result.get("suspicious", 0) > 0:
                        score += 50
            
            keyword_points = min(len(keywords_found) * 10, 30)
            score += keyword_points
            
            # CRASH-PROOF ATTACHMENT ANALYSIS
            print("[MSG Reader] Starting crash-proof attachment analysis...")
            attachments_list = []  # Initialize attachments list
            try:
                # Ensure msg exists before accessing attachments
                if msg is None:
                    print("[MSG Reader] ⚠️ msg is None, skipping attachment analysis")
                    attachments = []
                else:
                    # SAFE ATTACHMENT ACCESS - Wrap in try/except to prevent crashes
                    try:
                        attachments = getattr(msg, 'attachments', [])
                        # Try to access attachments to see if they're accessible
                        if attachments:
                            _ = len(attachments)  # This will fail if attachments are corrupted
                        print(f"[MSG Reader] ✓ Successfully accessed attachments list ({len(attachments)} items)")
                    except Exception as att_access_error:
                        print(f"[MSG Reader] ⚠️ Cannot access attachments: {att_access_error}")
                        attachments = []  # Continue without attachments
                
                print(f"[MSG Reader] Found {len(attachments)} attachment(s)")
                
                # Create temporary directory for attachment extraction fallback
                temp_attachment_dir = tempfile.mkdtemp()
                
                if attachments:
                    # Process each attachment with error isolation
                    for att_index, att in enumerate(attachments):  # Use safe variable, not msg.attachments directly
                        try:
                            print(f"[MSG Reader] Processing attachment {att_index + 1}/{len(attachments)}...")
                            
                            # Extract filename - Use getFilename() or fallback
                            filename = 'Unknown'
                            try:
                                if hasattr(att, 'getFilename'):
                                    try:
                                        filename = att.getFilename() or 'Unknown'
                                    except Exception as get_filename_error:
                                        print(f"[MSG Reader] getFilename() failed: {get_filename_error}")
                                        pass
                                if filename == 'Unknown':
                                    filename = getattr(att, 'longFilename', None) or getattr(att, 'shortFilename', None) or f'Attachment_{att_index + 1}'
                            except Exception as filename_error:
                                print(f"[MSG Reader] Error extracting filename: {filename_error}")
                                filename = f'Attachment_{att_index + 1}'
                            
                            print(f"[MSG Reader] Processing attachment: {filename}")
                            
                            att_data = None
                            
                            # Attempt A: Try att.data (direct memory access)
                            try:
                                att_data = att.data if hasattr(att, 'data') else None
                                if att_data and isinstance(att_data, bytes):
                                    print(f"[MSG Reader] ✓ Got attachment {filename} via att.data (Attempt A)")
                            except (AttributeError, TypeError) as mem_error:
                                print(f"[MSG Reader] Attempt A (att.data) failed for {filename}: {mem_error}")
                            
                            # Attempt B: Fallback - Save to temp file and read (works for complex OLE objects)
                            if att_data is None:
                                try:
                                    temp_att_path = os.path.join(temp_attachment_dir, f"att_{hashlib.md5(filename.encode()).hexdigest()[:8]}.tmp")
                                    print(f"[MSG Reader] Attempt B: Saving attachment {filename} to temp file...")
                                    
                                    # Check if att is Info object or embedded msg - use save() with customPath
                                    if hasattr(att, 'save'):
                                        try:
                                            # Try with customPath parameter first
                                            att.save(customPath=temp_attachment_dir)
                                            # Find the saved file
                                            saved_files = [f for f in os.listdir(temp_attachment_dir) if os.path.isfile(os.path.join(temp_attachment_dir, f))]
                                            if saved_files:
                                                temp_att_path = os.path.join(temp_attachment_dir, saved_files[0])
                                            print(f"[MSG Reader] ✓ Saved attachment {filename} to {temp_att_path}")
                                        except Exception as custom_path_error:
                                            print(f"[MSG Reader] customPath save failed: {custom_path_error}, trying direct path...")
                                            try:
                                                # Fallback to direct path
                                                att.save(temp_att_path)
                                                print(f"[MSG Reader] ✓ Saved attachment {filename} to direct path {temp_att_path}")
                                            except Exception as direct_path_error:
                                                print(f"[MSG Reader] Direct path save also failed: {direct_path_error}")
                                                raise direct_path_error
                                    else:
                                        print(f"[MSG Reader] Attachment {filename} has no save() method")
                                        attachments_list.append({
                                            "filename": str(filename),
                                            "hash": None,
                                            "detected_mime": None,
                                            "detected_ext": None,
                                            "detected_type": None,
                                            "extension_mismatch": False,
                                            "spoofing_warning": None,
                                            "has_dangerous_ext": False,
                                            "vt_stats": {},
                                            "error": "Unreadable Attachment: No save() method available"
                                        })
                                        continue
                                    
                                    # Read the saved file
                                    if os.path.exists(temp_att_path):
                                        try:
                                            with open(temp_att_path, 'rb') as f:
                                                att_data = f.read()
                                            print(f"[MSG Reader] ✓ Got attachment {filename} via temp file extraction (Attempt B) - {len(att_data)} bytes")
                                            # Clean up temp file immediately
                                            try:
                                                os.unlink(temp_att_path)
                                            except:
                                                pass
                                        except Exception as read_error:
                                            print(f"[MSG Reader] Failed to read saved file {temp_att_path}: {read_error}")
                                            att_data = None
                                    else:
                                        print(f"[MSG Reader] ⚠️ Saved file {temp_att_path} does not exist")
                                        att_data = None
                                except Exception as save_error:
                                    print(f"[MSG Reader] Attempt B (temp file save) failed for {filename}: {save_error}")
                            
                            # Attempt C: If still None, skip this attachment but log it
                            if att_data is None or not isinstance(att_data, bytes):
                                print(f"[MSG Reader] ⚠️ Skipped empty attachment {filename} (Attempt C - no data extracted)")
                                attachments_list.append({
                                    "filename": str(filename),
                                    "hash": None,
                                    "detected_mime": None,
                                    "detected_ext": None,
                                    "detected_type": None,
                                    "extension_mismatch": False,
                                    "spoofing_warning": None,
                                    "has_dangerous_ext": False,
                                    "vt_stats": {},
                                    "error": "Unreadable Attachment: Could not extract data (all extraction methods failed)"
                                })
                                continue  # Continue with next attachment
                            
                            # Calculate SHA256 hash
                            file_hash = hashlib.sha256(att_data).hexdigest()
                            
                            # Detect true file type using magic bytes
                            detected_mime = None
                            detected_ext = None
                            detected_type = None
                            
                            if FILETYPE_AVAILABLE:
                                try:
                                    kind = filetype.guess(att_data)
                                    if kind:
                                        detected_mime = kind.mime
                                        detected_ext = f".{kind.extension}"
                                        detected_type = kind.extension
                                except:
                                    pass
                            
                            if detected_type is None and att_data.startswith(b'%PDF'):
                                detected_mime = 'application/pdf'
                                detected_ext = '.pdf'
                                detected_type = 'pdf'
                            
                            # Spoofing Check: Extension vs Detected Type
                            filename_ext = os.path.splitext(filename)[1].lower() if filename else ''
                            extension_mismatch = False
                            spoofing_warning = None
                            
                            if detected_ext and filename_ext:
                                filename_ext_clean = filename_ext.lstrip('.')
                                detected_ext_clean = detected_ext.lstrip('.')
                                if filename_ext_clean != detected_ext_clean:
                                    extension_mismatch = True
                                    spoofing_warning = f"⚠️ TYPE MISMATCH! Filename shows {filename_ext} but actual type is {detected_ext}"
                                    if detected_type in ['exe', 'dosexec'] or detected_mime == 'application/x-dosexec':
                                        score += 100  # Critical
                                    else:
                                        score += 80
                            
                            # Check dangerous extensions
                            dangerous_extensions = ['.exe', '.bat', '.ps1', '.vbs', '.js', '.jar', '.scr', '.com', '.pif', '.cmd']
                            has_dangerous_ext = filename_ext in dangerous_extensions
                            if has_dangerous_ext:
                                score += 20
                            
                            # VirusTotal Lookup - CRITICAL: Always use FULL hash
                            vt_result = None
                            vt_error = None
                            try:
                                print(f"[MSG Reader] Running VirusTotal lookup for FULL hash: {file_hash} (length: {len(file_hash)})")
                                vt_result, vt_error = vt_fetch(file_hash)
                                if vt_result and isinstance(vt_result, dict):
                                    malicious = vt_result.get("malicious", 0) or 0
                                    suspicious = vt_result.get("suspicious", 0) or 0
                                    harmless = vt_result.get("harmless", 0) or 0
                                    print(f"[MSG Reader] ✓ VT Result for {file_hash[:16]}...: {malicious} malicious, {suspicious} suspicious, {harmless} harmless")
                                    if malicious > 0:
                                        score += 100
                                    elif suspicious > 0:
                                        score += 50
                                else:
                                    print(f"[MSG Reader] ⚠️ VT Lookup failed for {file_hash}: {vt_error}")
                            except Exception as vt_lookup_error:
                                print(f"[MSG Reader] ⚠️ VirusTotal lookup exception for {file_hash}: {vt_lookup_error}")
                                vt_error = str(vt_lookup_error)
                            
                            # Ensure hash is stored as full string
                            full_hash = str(file_hash)
                            
                            # Append to attachments list
                            attachments_list.append({
                                "filename": str(filename),
                                "hash": full_hash,  # FULL HASH - always store complete hash
                                "detected_mime": str(detected_mime) if detected_mime else None,
                                "detected_ext": str(detected_ext) if detected_ext else None,
                                "detected_type": str(detected_type) if detected_type else None,
                                "extension_mismatch": bool(extension_mismatch),
                                "spoofing_warning": str(spoofing_warning) if spoofing_warning else None,
                                "has_dangerous_ext": bool(has_dangerous_ext),
                                "vt_stats": {k: str(v) for k, v in vt_result.items()} if vt_result else {},
                                "error": str(vt_error) if vt_error else None
                            })
                            
                            print(f"[MSG Reader] ✓ Successfully processed attachment {filename} - Full Hash: {full_hash}")
                            
                        except Exception as att_error:
                            # Error Isolation: If one attachment fails, log and continue with next attachment
                            error_trace = traceback.format_exc()
                            print(f"[MSG Reader] ⚠️ Error processing attachment {att_index + 1}: {att_error}")
                            print(f"[MSG Reader] Error trace: {error_trace}")
                            attachments_list.append({
                                "filename": str(filename) if 'filename' in locals() else f"Attachment_{att_index + 1}",
                                "hash": None,
                                "detected_mime": None,
                                "detected_ext": None,
                                "detected_type": None,
                                "extension_mismatch": False,
                                "spoofing_warning": None,
                                "has_dangerous_ext": False,
                                "vt_stats": {},
                                "error": f"Failed to analyze attachment: {str(att_error)}"
                            })
                            continue  # Don't crash the email analysis - continue with next attachment
                
            except Exception as attachment_error:
                # Log error but continue with analysis - attachments are optional
                error_trace = traceback.format_exc()
                print(f"[MSG Reader] ⚠️ Error during attachment analysis (continuing with email body analysis):")
                print(error_trace)
                print(f"[MSG Reader] Attachment error: {str(attachment_error)}")
                # Don't add to attachments_list here - it's already handled in the inner try/except
                # Just continue with the rest of the analysis (headers, body, IOCs)
            
            finally:
                # Cleanup temp directory
                if temp_attachment_dir and os.path.exists(temp_attachment_dir):
                    try:
                        import shutil
                        shutil.rmtree(temp_attachment_dir)
                        print(f"[MSG Reader] Cleaned up temp attachment directory")
                    except Exception as cleanup_error:
                        print(f"[MSG Reader] Warning: Failed to cleanup temp directory: {cleanup_error}")
            
            # Format results
            ips_list = []
            for ip in unique_ips[:10]:
                result = vt_results_ips.get(ip, {})
                ips_list.append({
                    "ip": str(ip),
                    "malicious": int(result.get("malicious", 0)) if isinstance(result, dict) else 0,
                    "suspicious": int(result.get("suspicious", 0)) if isinstance(result, dict) else 0,
                    "error": str(result.get("error")) if isinstance(result, dict) and "error" in result else None
                })
            
            urls_list = []
            for url in unique_urls[:30]:
                result = vt_results_urls.get(url, {})
                if isinstance(result, dict):
                    urls_list.append({
                        "url": str(url),
                        "malicious": int(result.get("malicious", 0)),
                        "suspicious": int(result.get("suspicious", 0)),
                        "harmless": int(result.get("harmless", 0)),
                        "undetected": int(result.get("undetected", 0)),
                        "total": int(result.get("malicious", 0) + result.get("suspicious", 0) + result.get("harmless", 0) + result.get("undetected", 0)),
                        "error": str(result.get("error")) if "error" in result else None
                    })
                else:
                    urls_list.append({
                        "url": str(url),
                        "malicious": 0,
                        "suspicious": 0,
                        "harmless": 0,
                        "undetected": 0,
                        "total": 0,
                        "error": "No result"
                    })
            
            urls_list.sort(key=lambda x: (x.get("malicious", 0) * 1000 + x.get("suspicious", 0)), reverse=True)
            
            # Final verdict
            if score >= 100:
                verdict = "MALICIOUS"
            elif score >= 50:
                verdict = "SUSPICIOUS"
            else:
                verdict = "CLEAN"
            
            # Return full JSON structure - ensure all values are JSON-serializable
            # Sanitize headers (convert dates and other objects to strings)
            sanitized_headers = {}
            for key, value in headers.items():
                if value is None:
                    sanitized_headers[key] = "N/A"
                elif isinstance(value, (datetime, type)):
                    sanitized_headers[key] = str(value)
                else:
                    sanitized_headers[key] = str(value) if not isinstance(value, (str, int, float, bool)) else value
            
            # Sanitize attachments (ensure all values are JSON-serializable)
            sanitized_attachments = []
            for att in attachments_list:
                sanitized_att = {}
                for key, value in att.items():
                    if value is None:
                        sanitized_att[key] = None
                    elif isinstance(value, (datetime, type)):
                        sanitized_att[key] = str(value)
                    elif isinstance(value, dict):
                        # Sanitize nested dicts (like vt_stats)
                        sanitized_dict = {}
                        for k, v in value.items():
                            if v is None:
                                sanitized_dict[k] = None
                            elif isinstance(v, (datetime, type)):
                                sanitized_dict[k] = str(v)
                            else:
                                sanitized_dict[k] = v if isinstance(v, (str, int, float, bool, type(None))) else str(v)
                        sanitized_att[key] = sanitized_dict
                    else:
                        sanitized_att[key] = value if isinstance(value, (str, int, float, bool, type(None), list, dict)) else str(value)
                sanitized_attachments.append(sanitized_att)
            
            # Log attachments count before returning
            print(f"[MSG Reader] ✓ Returning result with {len(sanitized_attachments)} attachment(s)")
            if sanitized_attachments:
                print(f"[MSG Reader] Attachment filenames: {[att.get('filename', 'Unknown') for att in sanitized_attachments[:5]]}")
            
            return {
                "headers": sanitized_headers,
                "body_preview": str(body_preview) if body_preview else "",
                "urls": urls_list,
                "ips": ips_list,
                "keywords_found": keywords_found,
                "attachments": sanitized_attachments,
                "score": int(score),
                "verdict": str(verdict),
                "total_urls_found": int(total_urls_found),
                "urls_checked": int(urls_checked),
                "file_type": "MSG"
            }
    except Exception as load_error:
            # Handle errors when opening MSG file
            error_trace = traceback.format_exc()
            error_msg = str(load_error).lower()
            print(f"[MSG Reader] Failed to load MSG: {str(load_error)}")
            print(f"[MSG Reader] Error trace:")
            print(error_trace)
            
            # Try fallback method for corrupted files
            if 'ole' in error_msg or 'sector' in error_msg or 'stream' in error_msg or 'empty stream' in error_msg:
                print("[MSG Reader] OLE error detected. Checking if file is actually a PDF...")
                
                # CRITICAL: Check if this is actually a PDF masquerading as MSG
                try:
                    with open(file_path, 'rb') as f:
                        first_bytes = f.read(4096)  # Read first 4KB
                    
                    # Check for PDF signature (might be after some metadata)
                    is_pdf = False
                    if b'%PDF' in first_bytes[:2048]:  # Check first 2KB
                        is_pdf = True
                    print("[MSG Reader] 🚨 DETECTED: File is actually a PDF, not MSG! Routing to PDF analyzer...")
                    
                    # Check using filetype library if available
                    if not is_pdf and FILETYPE_AVAILABLE:
                        try:
                            kind = filetype.guess(first_bytes)
                            if kind and (kind.extension == 'pdf' or kind.mime == 'application/pdf'):
                                is_pdf = True
                            print(f"[MSG Reader] 🚨 DETECTED via filetype: File is actually PDF ({kind.mime})!")
                        except:
                            pass
                    
                    # If it's actually a PDF, analyze it as PDF instead
                    if is_pdf:
                        print("[MSG Reader] Analyzing as PDF (with extension spoofing warning)...")
                        try:
                            result = analyze_pdf(file_path)
                            # Add critical spoofing warning
                            result['spoofing_warning'] = "🚨 CRITICAL: EXTENSION SPOOFING DETECTED! File has .msg extension but is actually a PDF. This is a common phishing technique."
                            result['score'] = result.get('score', 0) + 100  # Add critical spoofing score
                            # Recalculate verdict with spoofing score
                            if result['score'] >= 100:
                                result['verdict'] = "MALICIOUS"
                            elif result['score'] >= 50:
                                result['verdict'] = "SUSPICIOUS"
                            return result
                        except Exception as pdf_error:
                            print(f"[MSG Reader] Error analyzing as PDF: {pdf_error}")
                            return {
                                "error": f"File type mismatch: File has .msg extension but is actually a PDF. PDF analysis failed: {str(pdf_error)}",
                                "spoofing_warning": "🚨 CRITICAL: File extension (.msg) does not match actual content (PDF)."
                            }
                    
                    print("[MSG Reader] Confirmed: Not a PDF. Proceeding with corrupted MSG extraction...")
                except Exception as pdf_check_error:
                    print(f"[MSG Reader] Error during PDF check: {pdf_check_error}")
                
                # Fallback to raw extraction for genuinely corrupted MSG files
                print("[MSG Robust Analysis] Attempting raw file extraction as fallback...")
                
                # CRITICAL: Try to extract attachments even from corrupted OLE file
                # First, try to open with extract_msg just for attachments (even if main parsing failed)
                fallback_attachments = []
                print("[MSG Robust Analysis] Attempting to extract attachments using extract_msg (attachment-only mode)...")
                try:
                    msg_att_only = extract_msg.Message(file_path)
                    try:
                        att_list = getattr(msg_att_only, 'attachments', [])
                        if att_list:
                            print(f"[MSG Robust Analysis] Found {len(att_list)} attachment(s) in attachment-only mode")
                            # Process attachments (simplified version)
                            for att_idx, att in enumerate(att_list):
                                try:
                                    filename = 'Unknown'
                                    if hasattr(att, 'getFilename'):
                                        try:
                                            filename = att.getFilename() or f'Attachment_{att_idx + 1}'
                                        except:
                                            filename = f'Attachment_{att_idx + 1}'
                                    
                                    att_data = None
                                    try:
                                        att_data = att.data if hasattr(att, 'data') else None
                                    except:
                                        pass
                                    
                                    if att_data and isinstance(att_data, bytes) and len(att_data) > 0:
                                        file_hash = hashlib.sha256(att_data).hexdigest()
                                        
                                        # Detect type
                                        detected_mime = None
                                        detected_ext = None
                                        if FILETYPE_AVAILABLE:
                                            try:
                                                kind = filetype.guess(att_data)
                                                if kind:
                                                    detected_mime = kind.mime
                                                    detected_ext = f".{kind.extension}"
                                            except:
                                                pass
                                        
                                        # VT lookup
                                        vt_result = None
                                        try:
                                            vt_result, _ = vt_fetch(file_hash)
                                        except:
                                            pass
                                        
                                        fallback_attachments.append({
                                            "filename": str(filename),
                                            "hash": str(file_hash),
                                            "detected_mime": str(detected_mime) if detected_mime else None,
                                            "detected_ext": str(detected_ext) if detected_ext else None,
                                            "detected_type": None,
                                            "extension_mismatch": False,
                                            "spoofing_warning": None,
                                            "has_dangerous_ext": False,
                                            "vt_stats": {k: str(v) for k, v in vt_result.items()} if vt_result else {},
                                            "error": None
                                        })
                                        print(f"[MSG Robust Analysis] ✓ Extracted attachment: {filename} ({len(att_data)} bytes)")
                                except Exception as att_process_error:
                                    print(f"[MSG Robust Analysis] Error processing attachment {att_idx + 1}: {att_process_error}")
                                    continue
                    finally:
                        try:
                            msg_att_only.close()
                        except:
                            pass
                except Exception as att_only_error:
                    print(f"[MSG Robust Analysis] Attachment-only extraction failed: {att_only_error}")
                
                # If extract_msg failed, try OLE file directly
                if not fallback_attachments and OLEFILE_AVAILABLE:
                    print("[MSG Robust Analysis] Attempting to extract attachments from corrupted OLE file...")
                    try:
                        ole = olefile.OleFileIO(file_path)
                        stream_list = ole.listdir()
                        
                        # Look for attachment streams (common patterns in MSG files)
                        attachment_patterns = ['__attach', 'attach', 'attachment', '__substg1.0_3701', '__substg1.0_3704']
                        for stream_path in stream_list:
                            stream_name = '/'.join(stream_path) if isinstance(stream_path, tuple) else str(stream_path)
                            
                            # Check if this stream might be an attachment
                            is_attachment_stream = any(pattern.lower() in stream_name.lower() for pattern in attachment_patterns)
                            
                            if is_attachment_stream:
                                try:
                                    stream_data = ole.openstream(stream_path).read()
                                    if stream_data and len(stream_data) > 100:  # Only process substantial streams
                                        # Try to extract filename from stream name or data
                                        filename = stream_name.split('/')[-1] or f"Attachment_{len(fallback_attachments) + 1}"
                                        
                                        # Calculate hash
                                        file_hash = hashlib.sha256(stream_data).hexdigest()
                                        
                                        # Try to detect file type
                                        detected_mime = None
                                        detected_ext = None
                                        detected_type = None
                                        
                                        if FILETYPE_AVAILABLE:
                                            try:
                                                kind = filetype.guess(stream_data)
                                                if kind:
                                                    detected_mime = kind.mime
                                                    detected_ext = f".{kind.extension}"
                                                    detected_type = kind.extension
                                            except:
                                                pass
                                        
                                        # VirusTotal lookup - CRITICAL: Always run VT for each hash (FULL HASH)
                                        vt_result = None
                                        vt_error = None
                                        try:
                                            print(f"[MSG Robust Analysis] Running VirusTotal lookup for FULL hash: {file_hash} (length: {len(file_hash)})")
                                            vt_result, vt_error = vt_fetch(file_hash)
                                            if vt_result:
                                                malicious = vt_result.get('malicious', 0) or 0
                                                suspicious = vt_result.get('suspicious', 0) or 0
                                                harmless = vt_result.get('harmless', 0) or 0
                                                print(f"[MSG Robust Analysis] ✓ VT Result for {file_hash[:16]}...: {malicious} malicious, {suspicious} suspicious, {harmless} harmless")
                                            else:
                                                print(f"[MSG Robust Analysis] ⚠️ VT Lookup failed for {file_hash}: {vt_error}")
                                        except Exception as vt_ex:
                                            print(f"[MSG Robust Analysis] ⚠️ VT Exception for {file_hash}: {vt_ex}")
                                            vt_error = str(vt_ex)
                                        
                                        # Ensure hash is stored as full string
                                        full_hash = str(file_hash)
                                        
                                        fallback_attachments.append({
                                            "filename": str(filename),
                                            "hash": full_hash,  # FULL HASH - always store complete hash
                                            "detected_mime": str(detected_mime) if detected_mime else None,
                                            "detected_ext": str(detected_ext) if detected_ext else None,
                                            "detected_type": str(detected_type) if detected_type else None,
                                            "extension_mismatch": False,
                                            "spoofing_warning": None,
                                            "has_dangerous_ext": False,
                                            "vt_stats": {k: str(v) for k, v in vt_result.items()} if vt_result else {},
                                            "error": str(vt_error) if vt_error else None
                                        })
                                        print(f"[MSG Robust Analysis] ✓ Extracted attachment: {filename} ({len(stream_data)} bytes) - Hash: {full_hash}")
                                except Exception as att_extract_error:
                                    print(f"[MSG Robust Analysis] Failed to extract attachment from stream {stream_name}: {att_extract_error}")
                                    continue
                        
                        ole.close()
                        print(f"[MSG Robust Analysis] Extracted {len(fallback_attachments)} attachment(s) from corrupted OLE file")
                    except Exception as ole_att_error:
                        print(f"[MSG Robust Analysis] Failed to extract attachments from OLE: {ole_att_error}")
                
                try:
                    with open(file_path, 'rb') as f:
                        raw_content = f.read()
                    
                    text_content = ""
                    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                    for encoding in encodings:
                        try:
                            text_content = raw_content.decode(encoding, errors='ignore')
                            if len(text_content) > 100:
                                break
                        except:
                            continue
                    
                    if not text_content or len(text_content) < 50:
                        ascii_pattern = re.compile(rb'[\x20-\x7E]{10,}')
                        ascii_matches = ascii_pattern.findall(raw_content)
                        text_content = b' '.join(ascii_matches).decode('ascii', errors='ignore')
                    
                    # Extract basic IOCs
                    ip_matches = _IP_REGEX_PATTERN.finditer(text_content)
                    unique_ips = list(set([match.group(0) for match in ip_matches]))
                    unique_ips = [ip for ip in unique_ips if not (ip.startswith(('127.', '192.168.', '10.', '172.')) or (ip.startswith('172.') and len(ip.split('.')) > 1 and ip.split('.')[1].isdigit() and 16 <= int(ip.split('.')[1]) <= 31))]
                    
                    url_matches = _URL_REGEX_PATTERN.findall(text_content)
                    unique_urls = list(set(url_matches))
                    
                    # Extract headers from text
                    subject_match = re.search(r'(?i)subject:\s*([^\r\n]+)', text_content)
                    if subject_match:
                        headers["Subject"] = subject_match.group(1).strip()[:100]
                    
                    from_match = re.search(r'(?i)from:\s*([^\r\n]+)', text_content)
                    if from_match:
                        headers["From"] = from_match.group(1).strip()[:100]
                    
                    body_preview = text_content[:500]
                    
                    # Calculate score based on attachments
                    fallback_score = 0
                    for att in fallback_attachments:
                        vt_stats = att.get('vt_stats', {})
                        if isinstance(vt_stats, dict):
                            malicious = int(vt_stats.get('malicious', 0) or 0)
                            suspicious = int(vt_stats.get('suspicious', 0) or 0)
                            if malicious > 0:
                                fallback_score += 100
                            elif suspicious > 0:
                                fallback_score += 50
                    
                    # Determine verdict based on score
                    if fallback_score >= 100:
                        fallback_verdict = "MALICIOUS"
                    elif fallback_score >= 50:
                        fallback_verdict = "SUSPICIOUS"
                    elif len(fallback_attachments) > 0:
                        fallback_verdict = "CLEAN"  # Has attachments but no threats detected
                    else:
                        fallback_verdict = "UNKNOWN"
                    
                    print(f"[MSG Robust Analysis] Calculated score: {fallback_score}, Verdict: {fallback_verdict}")
                    
                    return {
                        "headers": {k: str(v) for k, v in headers.items()},
                        "body_preview": str(body_preview),
                        "urls": [],
                        "ips": [],
                        "keywords_found": [],
                        "attachments": fallback_attachments,  # Include extracted attachments!
                        "score": int(fallback_score),
                        "verdict": str(fallback_verdict),
                        "total_urls_found": int(len(unique_urls)),
                        "urls_checked": 0,
                        "warning": f"⚠️ File appears corrupted (OLE error). Extracted limited information from raw file content. Found {len(fallback_attachments)} attachment(s). Original error: {str(load_error)}"
                    }
                except Exception as fallback_error:
                    return {
                        "error": f"Unable to parse .msg file. The file appears to be corrupted or damaged. Error: {str(load_error)}. Please try opening the file in Outlook and re-saving it, or export it as .eml format."
                    }
            else:
                return {
                    "error": f"Error loading email file: {str(load_error)}"
                }
    finally:
        # Ensure msg is closed even if there was an error
        if msg is not None:
            try:
                msg.close()
                print("[MSG Reader] Closed MSG file")
            except:
                pass
            
            # Cleanup temp directory if exists
            if temp_attachment_dir and os.path.exists(temp_attachment_dir):
                try:
                    import shutil
                    shutil.rmtree(temp_attachment_dir)
                    print(f"[MSG Reader] Cleaned up temp attachment directory")
                except:
                    pass

# ======= PHISHING HUNTER MODULE - EMAIL FORENSICS =======
def extract_attachment_robust(att, att_index, temp_dir):
    """
    Robust attachment extraction with multiple fallback methods.
    Returns: (att_data, filename, error_message)
    """
    filename = f'Attachment_{att_index + 1}'
    att_data = None
    error_message = None
    
    # Step 1: Get filename using multiple methods
    try:
        if hasattr(att, 'getFilename'):
            filename = att.getFilename() or filename
        if filename == f'Attachment_{att_index + 1}':
            filename = getattr(att, 'longFilename', None) or \
                      getattr(att, 'shortFilename', None) or \
                      getattr(att, 'name', None) or filename
    except Exception as e:
        print(f"[Attachment] Filename extraction error: {e}")
    
    # Step 2: Method A - Direct data access
    try:
        if hasattr(att, 'data'):
            att_data = att.data
            if att_data and isinstance(att_data, bytes) and len(att_data) > 0:
                print(f"[Attachment] ✓ Method A: Got {len(att_data)} bytes via att.data")
                return att_data, filename, None
    except Exception as e:
        print(f"[Attachment] Method A failed: {e}")
    
    # Step 3: Method B - Save to temp directory
    try:
        if hasattr(att, 'save'):
            # Try saving with customPath
            try:
                att.save(customPath=temp_dir)
                # Find the saved file
                saved_files = [f for f in os.listdir(temp_dir) 
                             if os.path.isfile(os.path.join(temp_dir, f))]
                
                if saved_files:
                    # Use the most recently created file
                    saved_files.sort(key=lambda x: os.path.getctime(os.path.join(temp_dir, x)), 
                                   reverse=True)
                    saved_path = os.path.join(temp_dir, saved_files[0])
                    
                    with open(saved_path, 'rb') as f:
                        att_data = f.read()
                    
                    print(f"[Attachment] ✓ Method B: Got {len(att_data)} bytes via save()")
                    
                    # Update filename if different
                    if saved_files[0] != filename:
                        filename = saved_files[0]
                    
                    # Clean up
                    try:
                        os.unlink(saved_path)
                    except:
                        pass
                    
                    return att_data, filename, None
            except Exception as custom_save_error:
                print(f"[Attachment] customPath save failed: {custom_save_error}")
                
                # Try direct path save as fallback
                try:
                    temp_path = os.path.join(temp_dir, secure_filename(filename))
                    att.save(temp_path)
                    
                    if os.path.exists(temp_path):
                        with open(temp_path, 'rb') as f:
                            att_data = f.read()
                        
                        print(f"[Attachment] ✓ Method B2: Got {len(att_data)} bytes via direct save")
                        
                        # Clean up
                        try:
                            os.unlink(temp_path)
                        except:
                            pass
                        
                        return att_data, filename, None
                except Exception as direct_save_error:
                    print(f"[Attachment] Direct save also failed: {direct_save_error}")
    except Exception as e:
        print(f"[Attachment] Method B failed: {e}")
    
    # Step 4: Method C - Try accessing as embedded message
    try:
        if hasattr(att, 'type') and att.type == 'msg':
            # This is an embedded MSG file
            if hasattr(att, 'msg'):
                embedded_msg = att.msg
                # Try to serialize the embedded message
                buffer = io.BytesIO()
                # This might not work for all versions, but worth trying
                try:
                    embedded_msg.save(buffer)
                    att_data = buffer.getvalue()
                    print(f"[Attachment] ✓ Method C: Got {len(att_data)} bytes from embedded MSG")
                    return att_data, filename, None
                except:
                    pass
    except Exception as e:
        print(f"[Attachment] Method C failed: {e}")
    
    # Step 5: Method D - Raw OLE stream access (last resort)
    try:
        if hasattr(att, '_AttachmentBase__attachmentDir'):
            att_dir = att._AttachmentBase__attachmentDir
            # Try to access the raw data stream
            if hasattr(att, '_AttachmentBase__msg') and hasattr(att._AttachmentBase__msg, 'path'):
                msg_path = att._AttachmentBase__msg.path
                # This is very low-level - access OLE file directly
                if OLEFILE_AVAILABLE:
                    try:
                        ole = olefile.OleFileIO(msg_path)
                        # Look for attachment data streams
                        for stream in ole.listdir():
                            stream_name = '/'.join(stream) if isinstance(stream, tuple) else str(stream)
                            if att_dir in stream_name and '3701' in stream_name:  # 3701 is attachment data property
                                att_data = ole.openstream(stream).read()
                                print(f"[Attachment] ✓ Method D: Got {len(att_data)} bytes via OLE stream")
                                ole.close()
                                return att_data, filename, None
                        ole.close()
                    except Exception as ole_error:
                        print(f"[Attachment] OLE access failed: {ole_error}")
    except Exception as e:
        print(f"[Attachment] Method D failed: {e}")
    
    # All methods failed
    error_message = "All extraction methods failed - attachment may be corrupted or encrypted"
    return None, filename, error_message

def extract_strings_from_binary(data: bytes) -> str:
    """
    The Hammer: Extract ALL printable strings from any binary blob.
    Crucial for finding uncompressed links in corrupted files.
    """
    all_text = ""
    
    try:
        # 1. ASCII Strings via Regex (5+ printable chars)
        ascii_matches = _ASCII_STRINGS.findall(data)
        if ascii_matches:
            # Join first 10000 strings to avoid memory issues
            all_text += b' '.join(ascii_matches[:10000]).decode('ascii', errors='ignore') + " "
        
        # 2. UTF-8 Decode (Best Effort)
        try:
            utf8_text = data.decode('utf-8', errors='ignore')
            if len(utf8_text) > len(all_text):
                all_text = utf8_text
        except:
            pass
        
        # 3. UTF-16LE Decode (Outlook Default)
        try:
            utf16_text = data.decode('utf-16-le', errors='ignore')
            if len(utf16_text) > len(all_text):
                all_text = utf16_text
        except:
            pass
        
    except Exception as e:
        print(f"[String Extraction] Error: {e}")
    
    return all_text

def scrape_visual_headers(text: str) -> dict:
    """
    Extract email headers from raw text by looking for header patterns.
    This is useful when MSG structure is broken but headers are visible in text.
    
    Args:
        text: Raw text content to search for headers
        
    Returns:
        Dictionary with header keys (Subject, From, To, Date) and their values
    """
    headers = {}
    
    if not text:
        return headers
    
    # Extract Subject
    subject_match = re.search(r'Subject:\s*([^\r\n]+)', text, re.IGNORECASE)
    if subject_match:
        headers["Subject"] = subject_match.group(1).strip()
    
    # Extract From (handles both simple and formatted email addresses)
    from_match = re.search(r'From:\s*([^\r\n<]+)', text, re.IGNORECASE)
    if from_match:
        from_value = from_match.group(1).strip()
        # Remove email address if present, keep name
        from_value = re.sub(r'<[^>]+>', '', from_value).strip()
        if from_value:
            headers["From"] = from_value
    
    # Extract To
    to_match = re.search(r'To:\s*([^\r\n<]+)', text, re.IGNORECASE)
    if to_match:
        to_value = to_match.group(1).strip()
        to_value = re.sub(r'<[^>]+>', '', to_value).strip()
        if to_value:
            headers["To"] = to_value
    
    # Extract Date (various formats)
    date_patterns = [
        r'Date:\s*([^\r\n]+)',
        r'Sent:\s*([^\r\n]+)',
        r'Received:\s*([^\r\n]+)'
    ]
    for pattern in date_patterns:
        date_match = re.search(pattern, text, re.IGNORECASE)
        if date_match:
            headers["Date"] = date_match.group(1).strip()
            break
    
    return headers

def extract_domains_from_received_headers(received_headers):
    """
    Extract all domains from Received headers.
    
    Args:
        received_headers: List of Received header strings or single string
        
    Returns:
        List of unique domains found in Received headers
    """
    domains = set()
    
    if not received_headers:
        return list(domains)
    
    # Handle both list and single string
    if isinstance(received_headers, str):
        received_headers = [received_headers]
    
    # Patterns to match domains in Received headers
    # Examples: "from domain.com", "by domain.com", "from [IP] (domain.com)"
    domain_patterns = [
        r'from\s+([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)+)',
        r'by\s+([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)+)',
        r'\(([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)+)\)',
        r'@([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)+)',
    ]
    
    for received_header in received_headers:
        if not received_header:
            continue
        
        received_str = str(received_header)
        
        # Try each pattern
        for pattern in domain_patterns:
            matches = re.finditer(pattern, received_str, re.IGNORECASE)
            for match in matches:
                domain = match.group(1).lower()
                # Filter out common non-domain patterns
                if domain and len(domain) > 3 and '.' in domain:
                    # Skip IP addresses and common patterns
                    if not re.match(r'^\d+\.\d+\.\d+\.\d+$', domain):
                        domains.add(domain)
    
    return list(domains)

def parse_authentication_results(auth_results_header):
    """
    Parse Authentication-Results header to extract SPF/DKIM/DMARC results.
    
    Args:
        auth_results_header: Authentication-Results header string or list
        
    Returns:
        Dict with parsed authentication results
    """
    results = {
        "spf_result": None,  # pass, fail, softfail, neutral, none, temperror, permerror
        "spf_domain": None,
        "dkim_result": None,
        "dkim_domain": None,
        "dmarc_result": None,
        "dmarc_domain": None,
        "raw": auth_results_header
    }
    
    if not auth_results_header:
        return results
    
    # Handle list of headers
    if isinstance(auth_results_header, list):
        auth_results_header = ' '.join(auth_results_header)
    
    auth_str = str(auth_results_header)
    
    # Parse SPF result: spf=pass (domain.com) or spf=fail, spf=softfail, etc.
    # Also handles: "spf=pass", "spf=fail", "spf=softfail", "spf=none"
    spf_patterns = [
        r'spf=(\w+)(?:\s+\(([^)]+)\))?',
        r'spf\s+(\w+)(?:\s+\(([^)]+)\))?',
        r'spf=(\w+)',
    ]
    
    for pattern in spf_patterns:
        spf_match = re.search(pattern, auth_str, re.IGNORECASE)
        if spf_match:
            results["spf_result"] = spf_match.group(1).lower()
            if len(spf_match.groups()) >= 2 and spf_match.group(2):
                results["spf_domain"] = spf_match.group(2).strip()
            # Try to extract domain from header.from=domain.com pattern
            if not results["spf_domain"]:
                from_match = re.search(r'header\.from=([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', auth_str, re.IGNORECASE)
                if from_match:
                    results["spf_domain"] = from_match.group(1).lower()
            break
    
    # Parse DKIM result: dkim=pass, dkim=fail, dkim=none, etc.
    dkim_patterns = [
        r'dkim=(\w+)(?:\s+\(([^)]+)\))?',
        r'dkim\s+(\w+)(?:\s+\(([^)]+)\))?',
        r'dkim=(\w+)',
    ]
    
    for pattern in dkim_patterns:
        dkim_match = re.search(pattern, auth_str, re.IGNORECASE)
        if dkim_match:
            results["dkim_result"] = dkim_match.group(1).lower()
            if len(dkim_match.groups()) >= 2 and dkim_match.group(2):
                results["dkim_domain"] = dkim_match.group(2).strip()
            # Try to extract domain from header.d=domain.com pattern
            if not results["dkim_domain"]:
                header_d_match = re.search(r'header\.d=([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', auth_str, re.IGNORECASE)
                if header_d_match:
                    results["dkim_domain"] = header_d_match.group(1).lower()
            break
    
    # Parse DMARC result: dmarc=pass, dmarc=fail, dmarc=none, etc.
    dmarc_patterns = [
        r'dmarc=(\w+)(?:\s+\(([^)]+)\))?',
        r'dmarc\s+(\w+)(?:\s+\(([^)]+)\))?',
        r'dmarc=(\w+)',
    ]
    
    for pattern in dmarc_patterns:
        dmarc_match = re.search(pattern, auth_str, re.IGNORECASE)
        if dmarc_match:
            results["dmarc_result"] = dmarc_match.group(1).lower()
            if len(dmarc_match.groups()) >= 2 and dmarc_match.group(2):
                results["dmarc_domain"] = dmarc_match.group(2).strip()
            # Try to extract domain from header.from=domain.com pattern (DMARC checks From header)
            if not results["dmarc_domain"]:
                from_match = re.search(r'header\.from=([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', auth_str, re.IGNORECASE)
                if from_match:
                    results["dmarc_domain"] = from_match.group(1).lower()
            break
    
    return results

def extract_deep_headers_from_wrapper(file_path, msg_reader_result=None):
    """
    Deep Header Extraction: Detects wrapper/report emails and extracts headers from embedded messages.
    
    This function:
    1. Detects if the email is a wrapper/report containing another email as attachment (message/rfc822)
    2. Extracts headers from the embedded email (the actual phishing email)
    3. Returns the embedded headers for DKIM validation
    
    Args:
        file_path: Path to the email file (.msg or .eml)
        msg_reader_result: Optional result from msg_reader_wrapper (to avoid re-parsing)
        
    Returns:
        Tuple of (embedded_headers, is_wrapper, wrapper_info)
        - embedded_headers: Dict with headers from embedded email (if found)
        - is_wrapper: Boolean indicating if this is a wrapper email
        - wrapper_info: Dict with information about the wrapper detection
    """
    embedded_headers = {}
    is_wrapper = False
    wrapper_info = {
        "detected": False,
        "method": None,
        "embedded_attachment_count": 0,
        "wrapper_subject": None,
        "wrapper_from": None
    }
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        # Method 1: Check msg_reader_result for embedded message attachments
        if msg_reader_result and msg_reader_result.get('success'):
            attachments = msg_reader_result.get('attachments', [])
            wrapper_info["wrapper_subject"] = msg_reader_result.get('headers', {}).get('subject', '')
            wrapper_info["wrapper_from"] = msg_reader_result.get('headers', {}).get('from', '')
            
            # Look for message/rfc822 attachments
            for att in attachments:
                mime_type = att.get('mime_type', '').lower()
                if 'message/rfc822' in mime_type or 'message/rfc822' in str(att.get('content_type', '')).lower():
                    wrapper_info["embedded_attachment_count"] += 1
                    is_wrapper = True
                    wrapper_info["detected"] = True
                    wrapper_info["method"] = "msg_reader_attachment"
                    
                    print(f"[Deep Headers] 🔍 Wrapper detected: Found message/rfc822 attachment: {att.get('filename', 'unknown')}")
                    
                    # Extract headers from embedded message
                    att_content = att.get('content')
                    if att_content:
                        try:
                            # Try parsing as EML format
                            from email import policy
                            from email.parser import BytesParser
                            
                            embedded_msg = BytesParser(policy=policy.default).parsebytes(att_content)
                            
                            # Extract ALL headers from embedded message (full Message Details)
                            embedded_headers = {}
                            # Get all header names and values
                            for header_name in embedded_msg.keys():
                                header_values = embedded_msg.get_all(header_name, [])
                                if len(header_values) == 1:
                                    embedded_headers[header_name] = header_values[0]
                                else:
                                    embedded_headers[header_name] = header_values  # Multiple values (like Received)
                            
                            # Ensure key headers are present
                            if 'Subject' not in embedded_headers:
                                embedded_headers['Subject'] = embedded_msg.get('Subject', '')
                            if 'From' not in embedded_headers:
                                embedded_headers['From'] = embedded_msg.get('From', '')
                            if 'To' not in embedded_headers:
                                embedded_headers['To'] = embedded_msg.get('To', '')
                            if 'Received' not in embedded_headers:
                                embedded_headers['Received'] = embedded_msg.get_all('Received', [])
                            
                            # Clean up headers (remove None values, convert lists to strings if needed)
                            cleaned_headers = {}
                            for k, v in embedded_headers.items():
                                if v:
                                    if isinstance(v, list):
                                        cleaned_headers[k] = v
                                    else:
                                        cleaned_headers[k] = str(v)
                            embedded_headers = cleaned_headers
                            
                            print(f"[Deep Headers] ✓ Extracted headers from embedded email:")
                            print(f"[Deep Headers]   From: {embedded_headers.get('From', 'N/A')}")
                            print(f"[Deep Headers]   Subject: {embedded_headers.get('Subject', 'N/A')}")
                            print(f"[Deep Headers]   DKIM-Signature: {'Present' if embedded_headers.get('DKIM-Signature') else 'Missing'}")
                            
                            return embedded_headers, is_wrapper, wrapper_info
                            
                        except Exception as parse_error:
                            print(f"[Deep Headers] ⚠️ Error parsing embedded message: {parse_error}")
                            # Try fallback: extract headers from raw text
                            try:
                                att_text = att_content.decode('utf-8', errors='ignore')
                                embedded_headers = scrape_visual_headers(att_text)
                                if embedded_headers:
                                    print(f"[Deep Headers] ✓ Extracted headers via text scraping")
                                    return embedded_headers, is_wrapper, wrapper_info
                            except:
                                pass
        
        # Method 2: Parse file directly using Python email library (for EML files)
        if file_ext == '.eml' and not embedded_headers:
            try:
                from email import policy
                from email.parser import BytesParser
                
                with open(file_path, 'rb') as f:
                    msg = BytesParser(policy=policy.default).parse(f)
                
                # Check if multipart with message/rfc822 parts
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        
                        if content_type == 'message/rfc822':
                            is_wrapper = True
                            wrapper_info["detected"] = True
                            wrapper_info["method"] = "eml_multipart"
                            wrapper_info["embedded_attachment_count"] += 1
                            
                            print(f"[Deep Headers] 🔍 Wrapper detected: Found message/rfc822 part in EML")
                            
                            # Extract embedded message
                            embedded_msg = part.get_payload()[0] if isinstance(part.get_payload(), list) else part.get_payload()
                            
                            if hasattr(embedded_msg, 'get'):
                                # Extract ALL headers from embedded message (full Message Details)
                                embedded_headers = {}
                                for header_name in embedded_msg.keys():
                                    header_values = embedded_msg.get_all(header_name, [])
                                    if len(header_values) == 1:
                                        embedded_headers[header_name] = header_values[0]
                                    else:
                                        embedded_headers[header_name] = header_values
                                
                                    # Clean up headers
                                    cleaned_headers = {}
                                    for k, v in embedded_headers.items():
                                        if v:
                                            if isinstance(v, list):
                                                cleaned_headers[k] = v
                                        else:
                                            cleaned_headers[k] = str(v)
                                embedded_headers = cleaned_headers
                                
                                print(f"[Deep Headers] ✓ Extracted headers from embedded EML message")
                                print(f"[Deep Headers]   From: {embedded_headers.get('From', 'N/A')}")
                                print(f"[Deep Headers]   DKIM-Signature: {'Present' if embedded_headers.get('DKIM-Signature') else 'Missing'}")
                                
                                return embedded_headers, is_wrapper, wrapper_info
                            
            except Exception as eml_error:
                print(f"[Deep Headers] ⚠️ Error parsing EML directly: {eml_error}")
        
        # Method 3: Check for MSG embedded messages using extract-msg (fallback)
        if file_ext == '.msg' and not embedded_headers and EXTRACT_MSG_AVAILABLE:
            try:
                import extract_msg
                msg_obj = extract_msg.Message(file_path)
                
                # Check attachments for embedded messages
                if hasattr(msg_obj, 'attachments') and msg_obj.attachments:
                    for att in msg_obj.attachments:
                        if hasattr(att, 'type') and att.type and 'message' in str(att.type).lower():
                            is_wrapper = True
                            wrapper_info["detected"] = True
                            wrapper_info["method"] = "extract_msg"
                            wrapper_info["embedded_attachment_count"] += 1
                            
                            print(f"[Deep Headers] 🔍 Wrapper detected via extract-msg: Found message attachment")
                            
                            # Try to extract headers from embedded MSG
                            try:
                                if hasattr(att, 'data') and att.data:
                                    # Save embedded MSG to temp file and parse
                                    fd, temp_embedded = tempfile.mkstemp(suffix='.msg')
                                    try:
                                        with os.fdopen(fd, 'wb') as tf:
                                            tf.write(att.data)
                                        
                                        embedded_msg_obj = extract_msg.Message(temp_embedded)
                                        
                                        embedded_headers = {
                                            "Subject": getattr(embedded_msg_obj, 'subject', ''),
                                            "From": getattr(embedded_msg_obj, 'sender', ''),
                                            "To": getattr(embedded_msg_obj, 'to', ''),
                                            "Date": str(getattr(embedded_msg_obj, 'date', '')),
                                            "Message-ID": getattr(embedded_msg_obj, 'messageId', '')
                                        }
                                        
                                        embedded_headers = {k: v for k, v in embedded_headers.items() if v}
                                        
                                        print(f"[Deep Headers] ✓ Extracted headers from embedded MSG")
                                        
                                        return embedded_headers, is_wrapper, wrapper_info
                                        
                                    finally:
                                        try:
                                            os.unlink(temp_embedded)
                                        except:
                                            pass
                            except Exception as nested_error:
                                print(f"[Deep Headers] ⚠️ Error extracting from nested MSG: {nested_error}")
                
                msg_obj.close()
            except Exception as extract_msg_error:
                print(f"[Deep Headers] ⚠️ Error using extract-msg: {extract_msg_error}")
        
    except Exception as deep_error:
        print(f"[Deep Headers] ⚠️ Error in deep header extraction: {deep_error}")
        traceback.print_exc()
    
    return embedded_headers, is_wrapper, wrapper_info

def analyze_email_content(file_path):
    """
    Onion Architecture V4: Multi-Layer Security Analysis Engine
    
    Implements a robust, crash-proof analysis pipeline that handles:
    1. Spoofing/Polyglots: Files named .msg that are actually PDFs
    2. Corrupted/Complex MSGs: Real OLE files with malformed headers but valid attachments
    
    Architecture Layers:
    - Layer 1: Foundation (Raw data & safety net)
    - Layer 2: Smart Routing (PDF vs OLE/MSG decision)
    - Layer 3: Data Merging & Enrichment (Text aggregation, header recovery)
    - Layer 4: Final IOC Sweep & Intelligence (VirusTotal lookups, verdict calculation)
    """
    print(f"[Onion] ====== Starting Onion Architecture Analysis ======")
    print(f"[Onion] File: {file_path}")
    
    # ========== INITIALIZATION ==========
    results = {
        "urls": set(),
        "ips": set(),
        "emails": set(),
        "verdict": "CLEAN",
        "score": 0,
        "warnings": [],
        "file_type": "UNKNOWN",
        "body_preview": "",
        "attachments": [],
        "score_reasons": [],  # Track all reasons for threat score
        "verdict_reasons": []  # Track all reasons for verdict
    }
    
    headers = {
        "Subject": "N/A",
        "From": "N/A",
        "To": "N/A",
        "Date": "N/A",
        "Message-ID": "N/A"
    }
    
    raw_data = b""
    file_hash = ""
    base_text = ""
    pdf_text = ""
    msg_body_text = ""
    msg_body_html = None  # Store HTML body separately
    
    # ========== LAYER 1: THE FOUNDATION ==========
    print("[Onion] Layer 1: Foundation - Reading raw data and creating safety net...")
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        if not raw_data:
            return {
                "error": "File is empty or could not be read",
                "verdict": "CORRUPTED",
                "score": 0,
                "file_type": "UNKNOWN"
            }
        
        # Calculate SHA256 hash
        file_hash = hashlib.sha256(raw_data).hexdigest()
        print(f"[Onion] ✓ File hash calculated: {file_hash[:16]}...")
        
        # Safety Net: Extract strings immediately (this ensures we have something even if parsers crash)
        base_text = extract_strings_from_binary(raw_data)
        print(f"[Onion] ✓ Safety net: Extracted {len(base_text)} characters from raw strings")
        
    except Exception as layer1_error:
        print(f"[Onion] ✗ Layer 1 error: {layer1_error}")
        traceback.print_exc()
        return {
            "error": f"Failed to read file: {str(layer1_error)}",
            "verdict": "CORRUPTED",
            "score": 0,
            "file_type": "UNKNOWN",
            "file_hash": "",
            "headers": headers,
            "urls": [],
            "ips": [],
            "attachments": []
        }
    
    # ========== LAYER 2: SMART ROUTING (THE BRAIN) ==========
    print("[Onion] Layer 2: Smart Routing - Analyzing file structure...")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    pdf_start = raw_data.find(b'%PDF', 0, 4096)  # Check first 4KB
    ole_header = raw_data.startswith(b'\xD0\xCF\x11\xE0')
    
    routing_strategy = None
    
    # Strategy A: PDF / Polyglot Detection
    if pdf_start >= 0:
        print(f"[Onion] ✓ PDF signature detected at offset {pdf_start}")
        
        # Check if this is a polyglot (PDF masquerading as MSG)
        is_polyglot = (file_extension == '.msg' and ole_header)
        
        if is_polyglot:
            print("[Onion] 🚨 POLYGLOT DETECTED: PDF hidden inside MSG container!")
            results["warnings"].append("🚨 CRITICAL: Extension spoofing detected - File has .msg extension but contains PDF data")
            results["score"] += 50
            results["score_reasons"].append(f"+50: PDF polyglot detected (file has .msg extension but contains PDF data)")
            results["verdict"] = "SUSPICIOUS (PDF Polyglot)"
            results["verdict_reasons"].append("PDF polyglot detected - file extension mismatch")
            
        results["file_type"] = "PDF"
        routing_strategy = "PDF"
        
        # Carve out PDF from the container (if embedded)
        if PDFMINER_AVAILABLE:
            try:
                print(f"[Onion] Carving PDF from offset {pdf_start}...")
                pdf_data_to_parse = raw_data[pdf_start:]
                
                # Create temporary file for PDFMiner
                fd, temp_pdf_path = tempfile.mkstemp(suffix='.pdf', prefix='onion_pdf_')
                try:
                    with os.fdopen(fd, 'wb') as tf:
                        tf.write(pdf_data_to_parse)
                        tf.flush()
                        os.fsync(tf.fileno())
                    
                    print("[Onion] Running PDFMiner on carved PDF data...")
                    pdf_text = extract_text(temp_pdf_path, laparams=LAParams())
                    print(f"[Onion] ✓ Extracted {len(pdf_text)} characters from PDF")
                    
                finally:
                    try:
                        os.unlink(temp_pdf_path)
                    except:
                        pass
                        
            except Exception as pdf_error:
                print(f"[Onion] ⚠️ PDF extraction error: {pdf_error}")
                traceback.print_exc()
        else:
            print("[Onion] ⚠️ PDFMiner not available, skipping PDF text extraction")
    
    # Strategy B: OLE / MSG Container
    elif ole_header or file_extension == '.msg' or file_extension == '.eml':
        print("[Onion] ✓ OLE/MSG/EML container detected")
        results["file_type"] = "MSG" if file_extension == '.msg' else "EML"
        routing_strategy = "MSG"
        
        # ENHANCED: Try Msg-Reader first for better attachment extraction
        msg_reader_attachments = []
        if MSG_READER_AVAILABLE and (file_extension == '.msg' or file_extension == '.eml'):
            print("[Onion] Attempting enhanced attachment extraction via Msg-Reader...")
            try:
                msg_reader_result = extract_email_attachments(file_path)
                if msg_reader_result.get('success'):
                    print(f"[Onion] ✓ Msg-Reader extracted {len(msg_reader_result.get('attachments', []))} attachment(s)")
                    
                    # Process Msg-Reader attachments
                    for att_data in msg_reader_result.get('attachments', []):
                        if att_data.get('content'):
                            # Analyze attachment
                            att_content = att_data['content']
                            att_filename = att_data.get('filename', 'unknown')
                            att_mime = att_data.get('mime_type', 'application/octet-stream')
                            # Calculate hash - normalize to lowercase for consistency
                            if att_data.get('hash'):
                                att_hash = str(att_data.get('hash')).lower().strip()
                            else:
                                att_hash = hashlib.sha256(att_content).hexdigest().lower()
                            
                            # Verify hash is valid SHA256 (64 hex chars)
                            if len(att_hash) != 64 or not all(c in '0123456789abcdef' for c in att_hash):
                                print(f"[Onion] WARNING: Invalid hash format, recalculating for {att_filename}")
                                att_hash = hashlib.sha256(att_content).hexdigest().lower()
                            
                            # Detect file type
                            detected_mime = None
                            detected_ext = None
                            if FILETYPE_AVAILABLE:
                                try:
                                    kind = filetype.guess(att_content)
                                    if kind:
                                        detected_mime = kind.mime
                                        detected_ext = f".{kind.extension}"
                                except:
                                    pass
                            
                            # VirusTotal lookup
                            vt_result = None
                            vt_error = None
                            try:
                                vt_result, vt_error = vt_fetch(att_hash)
                            except:
                                pass
                            
                            # Extract text from DOCX/PDF if possible
                            extracted_text = ""
                            if detected_mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                                extracted_text = extract_text_from_docx(att_content)
                            elif detected_mime == 'application/pdf' and PDFMINER_AVAILABLE:
                                try:
                                    fd, temp_pdf = tempfile.mkstemp(suffix='.pdf')
                                    try:
                                        with os.fdopen(fd, 'wb') as tf:
                                            tf.write(att_content)
                                        extracted_text = extract_text(temp_pdf, laparams=LAParams())
                                    finally:
                                        try:
                                            os.unlink(temp_pdf)
                                        except:
                                            pass
                                except:
                                    pass
                            
                            # Extract URLs from attachment text
                            extracted_urls = []
                            if extracted_text:
                                url_matches = _URL_REGEX_PATTERN.findall(extracted_text)
                                extracted_urls = list(set(url_matches))
                            
                            # Cache attachment for client-side viewing (BEFORE building attachment_info)
                            if att_hash and att_content:
                                cache_attachment(att_hash, att_content, att_filename, detected_mime or att_mime)
                            
                            attachment_info = {
                                "filename": att_filename,
                                "hash": att_hash,
                                "detected_mime": detected_mime or att_mime,
                                "detected_ext": detected_ext or os.path.splitext(att_filename)[1],
                                "detected_type": detected_ext.lstrip('.') if detected_ext else None,
                                                "extension_mismatch": False,
                                                "spoofing_warning": None,
                                                "has_dangerous_ext": False,
                                "vt_stats": {k: str(v) for k, v in vt_result.items()} if vt_result else {},
                                "error": str(vt_error) if vt_error else None,
                                            "extracted_text": extracted_text,
                                "extracted_urls": extracted_urls,
                                "size": len(att_content)
                            }
                            
                            # Check for spoofing
                            filename_ext = os.path.splitext(att_filename)[1].lower()
                            if detected_ext and filename_ext:
                                if filename_ext.lstrip('.') != detected_ext.lstrip('.'):
                                    attachment_info["extension_mismatch"] = True
                                    attachment_info["spoofing_warning"] = f"⚠️ TYPE MISMATCH! Filename shows {filename_ext} but actual type is {detected_ext}"
                            
                            # Check dangerous extensions
                            dangerous_extensions = ['.exe', '.bat', '.ps1', '.vbs', '.js', '.jar', '.scr', '.com', '.pif', '.cmd']
                            if filename_ext in dangerous_extensions:
                                attachment_info["has_dangerous_ext"] = True
                                results['score'] += 20
                                results["score_reasons"].append(f"+20: Attachment has dangerous extension ({filename_ext}): {att_filename}")
                            
                            # Add to score based on VT results and populate vt_stats
                            if vt_result:
                                # Populate vt_stats with all fields including community_score and detections
                                attachment_info["vt_stats"] = {
                                    "malicious": str(vt_result.get("malicious", 0) or 0),
                                    "suspicious": str(vt_result.get("suspicious", 0) or 0),
                                    "harmless": str(vt_result.get("harmless", 0) or 0),
                                    "undetected": str(vt_result.get("undetected", 0) or 0),
                                    "community_score": str(vt_result.get("community_score", 0) or 0),
                                    "detections": vt_result.get("detections", [])  # Keep as list, not string
                                }
                                
                                malicious = int(vt_result.get("malicious", 0) or 0)
                                suspicious = int(vt_result.get("suspicious", 0) or 0)
                                community_score = int(vt_result.get("community_score", 0) or 0)
                                
                                if malicious > 0:
                                    results['score'] += 100
                                    results["score_reasons"].append(f"+100: Attachment detected as malicious by {malicious} antivirus engines: {att_filename}")
                                elif suspicious > 0:
                                    results['score'] += 50
                                    results["score_reasons"].append(f"+50: Attachment detected as suspicious by {suspicious} antivirus engines: {att_filename}")
                                
                                # Add community score impact for attachments
                                if community_score < 0:
                                    # Negative community score = SET VERDICT TO MALICIOUS
                                    score_adjustment = max(-30, min(0, community_score // 2))
                                    results['score'] += abs(score_adjustment)
                                    results["score_reasons"].append(f"+{abs(score_adjustment)}: Attachment has negative community score ({community_score}): {att_filename}")
                                    results['verdict'] = "MALICIOUS"  # Negative community score = MALICIOUS
                                    results["verdict_reasons"].append(f"Attachment has negative community score ({community_score}): {att_filename}")
                                    print(f"[Onion] 🚨 Attachment has negative community score: {community_score} - SETTING VERDICT TO MALICIOUS")
                                
                                # Add VirusTotal community tab link
                                attachment_info["vt_community_link"] = get_vt_community_link(att_hash, "hash")
                                
                                if community_score > 10:
                                    score_reduction = min(10, community_score // 10)
                                    results['score'] = max(0, results['score'] - score_reduction)
                                    results["score_reasons"].append(f"-{score_reduction}: Attachment has positive community score ({community_score}): {att_filename}")
                                    print(f"[Onion] ✓ Attachment has positive community score: {community_score} (reducing score by {score_reduction})")
                            
                            # Cache attachment for client-side viewing
                            if att_hash and att_content:
                                cache_attachment(att_hash, att_content, att_filename, detected_mime or att_mime)
                            
                            msg_reader_attachments.append(attachment_info)
                    
                    # Merge headers from Msg-Reader if available
                    if msg_reader_result.get('headers'):
                        msg_headers = msg_reader_result['headers']
                        if msg_headers.get('subject'):
                            headers['Subject'] = msg_headers['subject']
                        if msg_headers.get('from'):
                            headers['From'] = msg_headers['from']
                        if msg_headers.get('to'):
                            headers['To'] = msg_headers['to']
                        if msg_headers.get('date'):
                            headers['Date'] = msg_headers['date']
                    
                    # Merge body text and HTML
                    if msg_reader_result.get('body'):
                        msg_body_text = msg_reader_result['body']
                    if msg_reader_result.get('body_html'):
                        msg_body_html = msg_reader_result['body_html']
                        # Use HTML as text source if no plain text body
                        if not msg_body_text:
                            msg_body_text = msg_body_html
                    
                    # DEEP HEADER EXTRACTION: Check for wrapper emails with embedded messages
                    print("[Onion] 🔍 Deep Header Extraction: Checking for wrapper/report emails...")
                    embedded_headers, is_wrapper, wrapper_info = extract_deep_headers_from_wrapper(file_path, msg_reader_result)
                    
                    if is_wrapper and embedded_headers:
                        print(f"[Onion] 🎯 WRAPPER EMAIL DETECTED: Found embedded email with {len(embedded_headers)} headers")
                        print(f"[Onion]   Wrapper Subject: {wrapper_info.get('wrapper_subject', 'N/A')}")
                        print(f"[Onion]   Embedded From: {embedded_headers.get('From', 'N/A')}")
                        print(f"[Onion]   Embedded Subject: {embedded_headers.get('Subject', 'N/A')}")
                        
                        # Store wrapper info in results
                        results["warnings"].append(f"⚠️ Wrapper/Report email detected: Original phishing email found as embedded attachment")
                        
                        # Use embedded headers for DKIM validation (will be used later)
                        # Store both wrapper and embedded headers
                        headers['_wrapper_detected'] = True
                        headers['_wrapper_info'] = wrapper_info
                        headers['_embedded_headers'] = embedded_headers
                        headers['_original_headers'] = headers.copy()  # Save original wrapper headers
                        
                        # Also update main headers to use embedded ones for display
                        if embedded_headers.get('From'):
                            headers['From'] = embedded_headers['From']
                        if embedded_headers.get('Subject'):
                            headers['Subject'] = embedded_headers['Subject']
                        if embedded_headers.get('To'):
                            headers['To'] = embedded_headers['To']
                        if embedded_headers.get('Date'):
                            headers['Date'] = embedded_headers['Date']
                    else:
                        print(f"[Onion] ✓ No wrapper detected - analyzing main email headers")
                
                else:
                    print(f"[Onion] ⚠️ Msg-Reader extraction failed: {msg_reader_result.get('error')}")
                    # Even if msg_reader failed, try deep header extraction directly on the file
                    print("[Onion] 🔍 Attempting deep header extraction directly on file (msg_reader failed)...")
                    embedded_headers, is_wrapper, wrapper_info = extract_deep_headers_from_wrapper(file_path, None)
                    if is_wrapper and embedded_headers:
                        print(f"[Onion] 🎯 WRAPPER EMAIL DETECTED (direct extraction): Found embedded email with {len(embedded_headers)} headers")
                        headers['_wrapper_detected'] = True
                        headers['_wrapper_info'] = wrapper_info
                        headers['_embedded_headers'] = embedded_headers
                        headers['_original_headers'] = headers.copy()
                        if embedded_headers.get('From'):
                            headers['From'] = embedded_headers['From']
                        if embedded_headers.get('Subject'):
                            headers['Subject'] = embedded_headers['Subject']
                        if embedded_headers.get('To'):
                            headers['To'] = embedded_headers['To']
                        if embedded_headers.get('Date'):
                            headers['Date'] = embedded_headers['Date']
            except Exception as msg_reader_error:
                print(f"[Onion] ⚠️ Msg-Reader error: {msg_reader_error}")
                traceback.print_exc()
                # Try deep header extraction as fallback
                try:
                    print("[Onion] 🔍 Attempting deep header extraction as fallback...")
                    embedded_headers, is_wrapper, wrapper_info = extract_deep_headers_from_wrapper(file_path, None)
                    if is_wrapper and embedded_headers:
                        print(f"[Onion] 🎯 WRAPPER EMAIL DETECTED (fallback): Found embedded email with {len(embedded_headers)} headers")
                        headers['_wrapper_detected'] = True
                        headers['_wrapper_info'] = wrapper_info
                        headers['_embedded_headers'] = embedded_headers
                        headers['_original_headers'] = headers.copy()
                        if embedded_headers.get('From'):
                            headers['From'] = embedded_headers['From']
                        if embedded_headers.get('Subject'):
                            headers['Subject'] = embedded_headers['Subject']
                        if embedded_headers.get('To'):
                            headers['To'] = embedded_headers['To']
                        if embedded_headers.get('Date'):
                            headers['Date'] = embedded_headers['Date']
                except Exception as fallback_error:
                    print(f"[Onion] ⚠️ Deep header extraction fallback also failed: {fallback_error}")
        
        # CRITICAL: Delegate to robust analyzer - DO NOT parse inline
        # For .eml files, skip analyze_msg_robust (already processed via extract_email_attachments)
        # For .msg files, use analyze_msg_robust for additional parsing
        msg_results = None
        if file_extension == '.msg':
            print("[Onion] Delegating to analyze_msg_robust (robust MSG analyzer)...")
            try:
                msg_results = analyze_msg_robust(file_path)
            except Exception as msg_robust_error:
                print(f"[Onion] ⚠️ analyze_msg_robust failed: {msg_robust_error}")
                msg_results = None
        else:
            print("[Onion] EML file detected - skipping analyze_msg_robust (already processed via extract_email_attachments)")
        
        if msg_results:
            if msg_results.get('error'):
                print(f"[Onion] ⚠️ Robust analyzer returned error: {msg_results['error']}")
                results["warnings"].append(f"MSG parsing warning: {msg_results['error']}")
            else:
                print("[Onion] ✓ Robust analysis completed successfully")
                
                # CRITICAL MERGE: Headers
                if msg_results.get('headers'):
                    for key, value in msg_results['headers'].items():
                        if value and value not in ["N/A", "Unknown", None]:
                            headers[key] = str(value)
                            print(f"[Onion] ✓ Merged header: {key} = {value}")
                
                # CRITICAL MERGE: Body text
                if msg_results.get('body_preview'):
                    msg_body_text = str(msg_results['body_preview'])
                    print(f"[Onion] ✓ Merged {len(msg_body_text)} characters of body text")
                
                # CRITICAL MERGE: URLs (handle both dict and string formats)
                for url_item in msg_results.get('urls', []):
                    if isinstance(url_item, dict):
                        url_str = url_item.get('url', '')
                    else:
                        url_str = str(url_item)
                    
                    if url_str and url_str.startswith('http'):
                        results["urls"].add(url_str)
                
                # CRITICAL MERGE: IPs (handle both dict and string formats)
                for ip_item in msg_results.get('ips', []):
                    if isinstance(ip_item, dict):
                        ip_str = ip_item.get('ip', '')
                    else:
                        ip_str = str(ip_item)
                    
                    if ip_str:
                        results["ips"].add(ip_str)
                
                # CRITICAL MERGE: Attachments
                # First add Msg-Reader attachments (if any)
                if msg_reader_attachments:
                    results['attachments'].extend(msg_reader_attachments)
                    print(f"[Onion] ✓ Added {len(msg_reader_attachments)} attachment(s) from Msg-Reader")
                
                # Then add attachments from robust analyzer (avoid duplicates)
                if msg_results.get('attachments'):
                    attachments_from_msg = msg_results['attachments']
                    # Check for duplicates by hash
                    existing_hashes = {att.get('hash') for att in results['attachments'] if att.get('hash')}
                    for att in attachments_from_msg:
                        att_hash = att.get('hash')
                        if att_hash and att_hash not in existing_hashes:
                            results['attachments'].append(att)
                            existing_hashes.add(att_hash)
                        elif not att_hash:  # Add even without hash to avoid missing attachments
                            results['attachments'].append(att)
                    print(f"[Onion] ✓ Merged {len(attachments_from_msg)} attachment(s) from robust analyzer")
                    
                    # ========== LAYER 4: RECURSIVE DEEP DIVE ==========
                    # Recursive Link Extraction: Extract text/URLs from attachments
                    print("[Onion] Layer 4: Recursive Deep Dive - Analyzing attachments recursively...")
                    for att in attachments_from_msg:
                        # Extract URLs from attachment metadata
                        if att.get('extracted_urls'):
                            for extracted_url in att['extracted_urls']:
                                if extracted_url:
                                    results["urls"].add(str(extracted_url))
                                    print(f"[Onion] ✓ Found URL in attachment: {extracted_url[:60]}...")
                        
                        # Extract text from attachments (e.g., DOCX content)
                        if att.get('extracted_text'):
                            attachment_text = str(att['extracted_text'])
                            msg_body_text += "\n" + attachment_text
                            print(f"[Onion] ✓ Extracted {len(attachment_text)} characters from attachment text")
                            
                            # Recursive: Extract URLs and IPs from attachment text
                            if attachment_text:
                                # Extract URLs from attachment text
                                att_urls = _URL_REGEX_PATTERN.findall(attachment_text)
                                for att_url in att_urls:
                                    if att_url and "http" in att_url.lower():
                                        results["urls"].add(att_url)
                                        print(f"[Onion] ✓ Found URL in attachment text: {att_url[:60]}...")
                                
                                # Extract IPs from attachment text
                                att_ip_matches = _IP_REGEX_PATTERN.finditer(attachment_text)
                                for ip_match in att_ip_matches:
                                    att_ip = ip_match.group(0)
                                    if att_ip and not is_private_ip(att_ip):
                                        results["ips"].add(att_ip)
                                        print(f"[Onion] ✓ Found IP in attachment text: {att_ip}")
                        
                        # Recursive: If attachment is a nested MSG/EML, analyze it recursively
                        att_filename = att.get('filename', '')
                        if att_filename and (att_filename.lower().endswith('.msg') or att_filename.lower().endswith('.eml')):
                            att_content = att.get('content') or att.get('data')
                            if att_content:
                                print(f"[Onion] 🔄 Recursive: Found nested email attachment: {att_filename}")
                                try:
                                    # Save nested attachment to temp file
                                    fd_nested, temp_nested_path = tempfile.mkstemp(suffix=os.path.splitext(att_filename)[1])
                                    try:
                                        with os.fdopen(fd_nested, 'wb') as tf_nested:
                                            if isinstance(att_content, bytes):
                                                tf_nested.write(att_content)
                                            else:
                                                tf_nested.write(att_content.encode('utf-8'))
                                        # Recursively analyze nested email
                                        nested_results = analyze_email_content(temp_nested_path)
                                        if nested_results and not nested_results.get('error'):
                                            # Merge nested results
                                            if nested_results.get('urls'):
                                                for nested_url in nested_results['urls']:
                                                    results["urls"].add(str(nested_url))
                                            if nested_results.get('ips'):
                                                for nested_ip in nested_results['ips']:
                                                    results["ips"].add(str(nested_ip))
                                            if nested_results.get('attachments'):
                                                results['attachments'].extend(nested_results['attachments'])
                                            print(f"[Onion] ✓ Recursively analyzed nested email: {len(nested_results.get('urls', []))} URLs, {len(nested_results.get('ips', []))} IPs")
                                    finally:
                                        try:
                                            os.unlink(temp_nested_path)
                                        except:
                                            pass
                                except Exception as nested_err:
                                    print(f"[Onion] ⚠️ Recursive analysis failed for nested email: {nested_err}")
                
                # CRITICAL MERGE: Score and verdict
                msg_score = msg_results.get('score', 0)
                if msg_score:
                    results['score'] += int(msg_score)
                    results["score_reasons"].append(f"+{int(msg_score)}: MSG analysis score from robust analyzer")
                
                if msg_results.get('verdict'):
                    msg_verdict = str(msg_results['verdict'])
                    if 'MALICIOUS' in msg_verdict:
                        results['verdict'] = "MALICIOUS"
                    elif 'SUSPICIOUS' in msg_verdict and results['verdict'] != "MALICIOUS":
                        results['verdict'] = "SUSPICIOUS"
                
                # Merge warnings if any
                if msg_results.get('warning'):
                    results["warnings"].append(str(msg_results['warning']))
                else:
                    # For .eml files, msg_results is None by design (not an error)
                    if file_extension != '.eml':
                        print("[Onion] ⚠️ Robust analyzer returned None (file may be corrupted)")
                        results["warnings"].append("MSG file structure appears corrupted")
        
        # CRITICAL: Add Msg-Reader attachments for both .msg and .eml files
        # (This needs to happen regardless of msg_results status)
        if msg_reader_attachments:
            existing_hashes = {att.get('hash') for att in results['attachments'] if att.get('hash')}
            new_attachments = []
            for msg_att in msg_reader_attachments:
                msg_hash = msg_att.get('hash')
                # Only add if hash doesn't exist or attachment has no hash
                if not msg_hash or msg_hash not in existing_hashes:
                    new_attachments.append(msg_att)
                    if msg_hash:
                        existing_hashes.add(msg_hash)
            if new_attachments:
                results['attachments'].extend(new_attachments)
                print(f"[Onion] ✓ Added {len(new_attachments)} attachment(s) from Msg-Reader ({file_extension} file)")
    
        else:
            print("[Onion] ⚠️ Unknown file format - no recognized signatures")
            results["file_type"] = "UNKNOWN"
            routing_strategy = "UNKNOWN"
    
    # ========== LAYER 3: DATA MERGING & ENRICHMENT ==========
    print("[Onion] Layer 3: Data Merging & Enrichment...")
    
    # Combine all text sources into one unified base_text
    # Include text from all attachments (recursive extraction results)
    text_parts = []
    if base_text:
        text_parts.append(base_text)
    if pdf_text:
        text_parts.append(pdf_text)
    if msg_body_text:
        text_parts.append(msg_body_text)
    
    # Add text from all attachments (from recursive deep dive)
    for att in results.get('attachments', []):
        if att.get('extracted_text'):
            att_text = str(att['extracted_text'])
            if att_text:
                text_parts.append(att_text)
                print(f"[Onion] ✓ Added {len(att_text)} characters from attachment: {att.get('filename', 'unknown')}")
    
    # Create unified base_text
    base_text = "\n".join(text_parts)
    print(f"[Onion] ✓ Unified text block: {len(base_text)} total characters (from all sources)")
    
    # ========== LAYER 5: INFORMATION ENRICHMENT ==========
    # Header Recovery: If headers are still missing, scrape from text
    print("[Onion] Layer 5: Information Enrichment - Recovering missing headers...")
    if headers.get("Subject") in [None, "N/A", "Unknown", ""]:
        subject_match = re.search(r'Subject:\s*([^\r\n<]+)', base_text, re.IGNORECASE)
        if subject_match:
            headers["Subject"] = subject_match.group(1).strip()
            print(f"[Onion] ✓ Recovered Subject from text: {headers['Subject'][:50]}...")
    
    if headers.get("From") in [None, "N/A", "Unknown", ""]:
        from_match = re.search(r'From:\s*([^\r\n<]+)', base_text, re.IGNORECASE)
        if from_match:
            from_value = from_match.group(1).strip()
            # Clean email addresses if present
            from_value = re.sub(r'<[^>]+>', '', from_value).strip()
            if from_value:
                headers["From"] = from_value
                print(f"[Onion] ✓ Recovered From from text: {headers['From'][:50]}...")
    
    if headers.get("To") in [None, "N/A", "Unknown", ""]:
        to_match = re.search(r'To:\s*([^\r\n<]+)', base_text, re.IGNORECASE)
        if to_match:
            to_value = to_match.group(1).strip()
            to_value = re.sub(r'<[^>]+>', '', to_value).strip()
            if to_value:
                headers["To"] = to_value
                print(f"[Onion] ✓ Recovered To from text: {headers['To'][:50]}...")
    
    if headers.get("Date") in [None, "N/A", "Unknown", ""]:
        date_match = re.search(r'(?:Date|Sent):\s*([^\r\n]+)', base_text, re.IGNORECASE)
        if date_match:
            headers["Date"] = date_match.group(1).strip()
            print(f"[Onion] ✓ Recovered Date from text: {headers['Date'][:50]}...")
    
    # Use scrape_visual_headers as additional fallback
    scraped_headers = scrape_visual_headers(base_text)
    for key in ["Subject", "From", "To", "Date"]:
        if headers.get(key) in [None, "N/A", "Unknown", ""] and scraped_headers.get(key):
            headers[key] = scraped_headers[key]
            print(f"[Onion] ✓ Recovered {key} via visual scraper: {headers[key][:50]}...")
    
    # ========== LAYER 6: THREAT INTELLIGENCE & VERDICT ==========
    print("[Onion] Layer 6: Threat Intelligence & Verdict...")
    
    # Final URL extraction from unified base_text
    url_list = _URL_REGEX_PATTERN.findall(base_text)
    IGNORED_DOMAINS = {'ns.adobe.com', 'www.w3.org', 'purl.org', 'schemas.xmlsoap.org', 
                       'adobe.com', 'w3.org', 'schemas.microsoft.com'}
    
    for url in url_list:
        if url and "http" in url.lower():
                # Filter out technical/namespace domains
            url_lower = url.lower()
            should_ignore = any(ignored in url_lower for ignored in IGNORED_DOMAINS)
            if not should_ignore:
                results["urls"].add(url)
    
    print(f"[Onion] ✓ Found {len(results['urls'])} unique URLs")
    
    # Final IP extraction from unified base_text
    ip_matches = _IP_REGEX_PATTERN.finditer(base_text)
    for match in ip_matches:
        ip = match.group(0)
        if ip and not is_private_ip(ip):
            results["ips"].add(ip)
    
    print(f"[Onion] ✓ Found {len(results['ips'])} unique public IPs")
    
    # VirusTotal Intelligence: Check top 5 URLs
    check_urls = list(results["urls"])[:5]
    final_url_objects = []
    
    for url in check_urls:
        try:
            vt_result, vt_error = vt_fetch(url)
            url_obj = {
                "url": url,
                        "malicious": 0,
                        "suspicious": 0,
                        "harmless": 0,
                        "undetected": 0,
                "total": 0
            }
            
            if vt_result and isinstance(vt_result, dict):
                url_obj.update({
                    "malicious": int(vt_result.get("malicious", 0) or 0),
                    "suspicious": int(vt_result.get("suspicious", 0) or 0),
                    "harmless": int(vt_result.get("harmless", 0) or 0),
                    "undetected": int(vt_result.get("undetected", 0) or 0),
                    "total": int(vt_result.get("malicious", 0) or 0) + 
                             int(vt_result.get("suspicious", 0) or 0) + 
                             int(vt_result.get("harmless", 0) or 0) + 
                             int(vt_result.get("undetected", 0) or 0),
                    "community_score": int(vt_result.get("community_score", 0) or 0),
                    "detections": vt_result.get("detections", [])
                })
                
                malicious = url_obj["malicious"]
                community_score = url_obj.get("community_score", 0)
                
                if malicious > 0:
                    results['score'] += 100
                    results["score_reasons"].append(f"+100: URL detected as malicious by {malicious} antivirus engines: {url[:80]}")
                    print(f"[Onion] 🚨 MALICIOUS URL detected: {url[:50]}... ({malicious} detections)")
                elif url_obj["suspicious"] > 0:
                    results['score'] += 50
                    results["score_reasons"].append(f"+50: URL detected as suspicious by {url_obj['suspicious']} antivirus engines: {url[:80]}")
                    print(f"[Onion] ⚠️ SUSPICIOUS URL detected: {url[:50]}...")
                
                # Add community score impact (negative = bad, positive = good)
                # Community score: (harmless votes - malicious votes)
                # Negative community score means more malicious votes
                if community_score < 0:
                    # Negative community score = bad reputation - SET VERDICT TO MALICIOUS
                    score_adjustment = max(-30, min(0, community_score // 2))  # Cap at -30 points
                    results['score'] += abs(score_adjustment)
                    results["score_reasons"].append(f"+{abs(score_adjustment)}: URL has negative community score ({community_score}): {url[:80]}")
                    results['verdict'] = "MALICIOUS"  # Negative community score = MALICIOUS
                    results["verdict_reasons"].append(f"URL has negative community score ({community_score}): {url[:80]}")
                    print(f"[Onion] 🚨 URL has negative community score: {community_score} - SETTING VERDICT TO MALICIOUS")
                
                # Add VirusTotal community tab link
                url_obj["vt_community_link"] = get_vt_community_link(url, "url")
                
                if community_score > 10:
                    # Very positive community score = good reputation (small reduction)
                    score_reduction = min(10, community_score // 10)
                    results['score'] = max(0, results['score'] - score_reduction)
                    results["score_reasons"].append(f"-{score_reduction}: URL has positive community score ({community_score}): {url[:80]}")
                    print(f"[Onion] ✓ URL has positive community score: {community_score} (reducing score by {score_reduction})")
            
            if vt_error:
                url_obj["error"] = str(vt_error)
            
            final_url_objects.append(url_obj)
            
        except Exception as vt_error:
            print(f"[Onion] ⚠️ VT lookup error for {url[:50]}...: {vt_error}")
            final_url_objects.append({
                "url": url,
                        "malicious": 0,
                        "suspicious": 0,
                        "harmless": 0,
                        "undetected": 0,
                        "total": 0,
                "error": str(vt_error)
            })
    
    # VirusTotal Intelligence: Check top 10 IPs
    check_ips = list(results["ips"])[:10]
    final_ip_objects = []
    
    for ip in check_ips:
        try:
            vt_result, vt_error = vt_fetch(ip)
            ip_obj = {
                "ip": ip,
                    "malicious": 0,
                    "suspicious": 0,
                "error": None
            }
            
            if vt_result and isinstance(vt_result, dict):
                ip_obj.update({
                    "malicious": int(vt_result.get("malicious", 0) or 0),
                    "suspicious": int(vt_result.get("suspicious", 0) or 0),
                    "community_score": int(vt_result.get("community_score", 0) or 0),
                    "detections": vt_result.get("detections", [])
                })
                
                malicious = ip_obj["malicious"]
                community_score = ip_obj.get("community_score", 0)
                
                if malicious > 0:
                    results['score'] += 100
                    results["score_reasons"].append(f"+100: IP address detected as malicious by {malicious} antivirus engines: {ip}")
                    print(f"[Onion] 🚨 MALICIOUS IP detected: {ip} ({malicious} detections)")
                elif ip_obj["suspicious"] > 0:
                    results['score'] += 50
                    results["score_reasons"].append(f"+50: IP address detected as suspicious by {ip_obj['suspicious']} antivirus engines: {ip}")
                    print(f"[Onion] ⚠️ SUSPICIOUS IP detected: {ip}")
                
                # Add community score impact (negative = bad, positive = good)
                if community_score < 0:
                    # Negative community score = SET VERDICT TO MALICIOUS
                    score_adjustment = max(-30, min(0, community_score // 2))  # Cap at -30 points
                    results['score'] += abs(score_adjustment)
                    results["score_reasons"].append(f"+{abs(score_adjustment)}: IP address has negative community score ({community_score}): {ip}")
                    results['verdict'] = "MALICIOUS"  # Negative community score = MALICIOUS
                    results["verdict_reasons"].append(f"IP address has negative community score ({community_score}): {ip}")
                    print(f"[Onion] 🚨 IP has negative community score: {community_score} - SETTING VERDICT TO MALICIOUS")
                
                # Add VirusTotal community tab link
                ip_obj["vt_community_link"] = get_vt_community_link(ip, "ip")
                
                if community_score > 10:
                    score_reduction = min(10, community_score // 10)
                    results['score'] = max(0, results['score'] - score_reduction)
                    results["score_reasons"].append(f"-{score_reduction}: IP address has positive community score ({community_score}): {ip}")
                    print(f"[Onion] ✓ IP has positive community score: {community_score} (reducing score by {score_reduction})")
            
            if vt_error:
                ip_obj["error"] = str(vt_error)
            
            final_ip_objects.append(ip_obj)
            
        except Exception as vt_error:
            print(f"[Onion] ⚠️ VT lookup error for {ip}: {vt_error}")
            final_ip_objects.append({
                "ip": ip,
                                                "malicious": 0,
                                                "suspicious": 0,
                "error": str(vt_error)
            })
    
    # VirusTotal Intelligence: Check Sender Domain
    sender_domain_info = None
    sender_domain = None
    if headers.get("From") and headers["From"] not in ["N/A", "Unknown", ""]:
        try:
            # Extract domain from email address (e.g., "user@example.com" -> "example.com")
            from_email = headers["From"]
            # Handle both "Name <email@domain.com>" and "email@domain.com" formats
            email_match = re.search(r'[\w\.-]+@([\w\.-]+\.[a-zA-Z]{2,})', from_email)
            if email_match:
                sender_domain = email_match.group(1).lower()
                print(f"[Onion] Checking sender domain: {sender_domain}")
                
                try:
                    vt_result, vt_error = vt_fetch(sender_domain)
                    sender_domain_info = {
                        "domain": sender_domain,
                                                    "malicious": 0,
                                                    "suspicious": 0,
                                                    "harmless": 0,
                                                    "undetected": 0,
                                                    "total": 0,
                        "error": None
                    }
                    
                    if vt_result and isinstance(vt_result, dict):
                        sender_domain_info.update({
                            "malicious": int(vt_result.get("malicious", 0) or 0),
                            "suspicious": int(vt_result.get("suspicious", 0) or 0),
                            "harmless": int(vt_result.get("harmless", 0) or 0),
                            "undetected": int(vt_result.get("undetected", 0) or 0),
                            "total": int(vt_result.get("malicious", 0) or 0) + 
                                     int(vt_result.get("suspicious", 0) or 0) + 
                                     int(vt_result.get("harmless", 0) or 0) + 
                                     int(vt_result.get("undetected", 0) or 0),
                            "community_score": int(vt_result.get("community_score", 0) or 0),
                            "detections": vt_result.get("detections", [])
                        })
                        
                        malicious = sender_domain_info["malicious"]
                        suspicious = sender_domain_info["suspicious"]
                        community_score = sender_domain_info.get("community_score", 0)
                        
                        # Add score based on domain reputation
                        if malicious > 0:
                            results['score'] += 75  # High score for malicious domains
                            results["score_reasons"].append(f"+75: Sender domain detected as malicious by {malicious} antivirus engines: {sender_domain}")
                            print(f"[Onion] 🚨 MALICIOUS sender domain detected: {sender_domain} ({malicious} detections)")
                        elif suspicious > 0:
                            results['score'] += 40  # Medium score for suspicious domains
                            results["score_reasons"].append(f"+40: Sender domain detected as suspicious by {suspicious} antivirus engines: {sender_domain}")
                            print(f"[Onion] ⚠️ SUSPICIOUS sender domain detected: {sender_domain} ({suspicious} detections)")
                        elif sender_domain_info["harmless"] > 0:
                            # Domain is clean (known good)
                            print(f"[Onion] ✓ Clean sender domain: {sender_domain}")
                    else:
                            # Domain not found or unknown
                            print(f"[Onion] ⚠️ Unknown sender domain: {sender_domain} (not in VirusTotal database)")
                            results['score'] += 5  # Small score for unknown domains
                            results["score_reasons"].append(f"+5: Sender domain unknown (not in VirusTotal database): {sender_domain}")
                            
                            # Add community score impact
                            if community_score < 0:
                                # Negative community score = SET VERDICT TO MALICIOUS
                                score_adjustment = max(-25, min(0, community_score // 2))  # Cap at -25 points
                                results['score'] += abs(score_adjustment)
                                results["score_reasons"].append(f"+{abs(score_adjustment)}: Sender domain has negative community score ({community_score}): {sender_domain}")
                                results['verdict'] = "MALICIOUS"  # Negative community score = MALICIOUS
                                results["verdict_reasons"].append(f"Sender domain has negative community score ({community_score}): {sender_domain}")
                                print(f"[Onion] 🚨 Domain has negative community score: {community_score} - SETTING VERDICT TO MALICIOUS")
                            
                            # Add VirusTotal community tab link
                            sender_domain_info["vt_community_link"] = get_vt_community_link(sender_domain, "domain")
                            
                            if community_score > 10:
                                score_reduction = min(8, community_score // 10)
                                results['score'] = max(0, results['score'] - score_reduction)
                                results["score_reasons"].append(f"-{score_reduction}: Sender domain has positive community score ({community_score}): {sender_domain}")
                            print(f"[Onion] ✓ Domain has positive community score: {community_score} (reducing score by {score_reduction})")
                    
                    if vt_error:
                        sender_domain_info["error"] = str(vt_error)
                        print(f"[Onion] ⚠️ VT lookup error for sender domain {sender_domain}: {vt_error}")
                    
                except Exception as domain_error:
                    print(f"[Onion] ⚠️ Error checking sender domain {sender_domain}: {domain_error}")
                    sender_domain_info = {
                        "domain": sender_domain,
                                                "malicious": 0,
                                                "suspicious": 0,
                        "error": str(domain_error)
                    }
        except Exception as extract_error:
            print(f"[Onion] ⚠️ Error extracting domain from From header: {extract_error}")
    
    # DKIM Check: Check DKIM configuration for sender domain
    # Use DEEP HEADER EXTRACTION: If wrapper email detected, use embedded headers for DKIM validation
    dkim_info = None
    deep_extracted_sender_domain = None
    is_wrapper_email = headers.get('_wrapper_detected', False)
    embedded_headers = headers.get('_embedded_headers', {})
    
    # Determine which sender domain to use for DKIM check
    if is_wrapper_email and embedded_headers:
        # Use embedded email's sender domain for DKIM validation
        embedded_from = embedded_headers.get('From', '')
        if embedded_from:
            email_match = re.search(r'[\w\.-]+@([\w\.-]+\.[a-zA-Z]{2,})', embedded_from)
            if email_match:
                deep_extracted_sender_domain = email_match.group(1).lower()
                print(f"[Onion] 🎯 DEEP HEADER EXTRACTION: Using embedded email sender domain for DKIM: {deep_extracted_sender_domain}")
                print(f"[Onion]   Embedded From: {embedded_from}")
                print(f"[Onion]   Embedded DKIM-Signature: {'Present' if embedded_headers.get('DKIM-Signature') else 'Missing'}")
        
        # Use the embedded sender domain for DKIM check
        sender_domain_for_dkim = deep_extracted_sender_domain or sender_domain
    else:
        # Use main email's sender domain
        sender_domain_for_dkim = sender_domain
    
    if sender_domain_for_dkim and DKIM_CHECK_AVAILABLE:
        try:
            print(f"[Onion] Checking DKIM configuration for domain: {sender_domain_for_dkim}")
            if is_wrapper_email:
                print(f"[Onion]   Note: This is a wrapper email - validating embedded phishing email's DKIM")
            
            # Setup colors for Magic-Spoofing (dummy colors for web app)
            colors = setup_colors() if DKIM_CHECK_AVAILABLE else {
                "white_bold": "", "info": "", "normal": "", "green": "", "red": ""
            }
            
            # Check DKIM for the domain
            dkim_result = check_dkim(sender_domain_for_dkim, colors)
            
            # Convert to simple dict format for JSON response
            dkim_info = {
                "domain": sender_domain_for_dkim,
                "selectors_found": dkim_result.get("selectors_found", []),
                "security_level": dkim_result.get("security_level", "None"),
                "issues": dkim_result.get("issues", []),
                "recommendations": dkim_result.get("recommendations", []),
                "records": {},
                "is_wrapper_email": is_wrapper_email,
                "deep_extraction_used": bool(is_wrapper_email and embedded_headers)
            }
            
            # Add wrapper information if available
            if is_wrapper_email:
                wrapper_info = headers.get('_wrapper_info', {})
                dkim_info["wrapper_info"] = {
                    "wrapper_subject": wrapper_info.get("wrapper_subject", ""),
                    "wrapper_from": wrapper_info.get("wrapper_from", ""),
                    "extraction_method": wrapper_info.get("method", "")
                }
                
                # Check if DKIM-Signature was found in embedded headers
                if embedded_headers.get('DKIM-Signature'):
                    dkim_info["dkim_signature_in_headers"] = True
                    dkim_info["dkim_signature_preview"] = embedded_headers['DKIM-Signature'][:100] + "..." if len(embedded_headers['DKIM-Signature']) > 100 else embedded_headers['DKIM-Signature']
                else:
                    dkim_info["dkim_signature_in_headers"] = False
                    dkim_info["dkim_signature_preview"] = None
            
            # Extract records info (simplified)
            for selector in dkim_result.get("selectors_found", []):
                record = dkim_result.get("records", {}).get(selector, "")
                if record:
                    dkim_info["records"][selector] = record[:200]  # Limit length
            
            print(f"[Onion] ✓ DKIM check complete: {len(dkim_info['selectors_found'])} selector(s) found, security level: {dkim_info['security_level']}")
            if is_wrapper_email:
                print(f"[Onion]   ✓ Deep header extraction: Validated embedded email's DKIM (not wrapper)")
            
            # Add score adjustment based on DKIM configuration
            if dkim_info["security_level"] == "None":
                results['score'] += 10
                domain_label = "embedded phishing email domain" if is_wrapper_email else "sender domain"
                results["score_reasons"].append(f"+10: No DKIM configuration found for {domain_label}: {sender_domain_for_dkim}")
                print(f"[Onion] ⚠️ No DKIM configuration found for {sender_domain_for_dkim}")
            elif dkim_info["security_level"] == "Low":
                results['score'] += 5
                domain_label = "embedded phishing email domain" if is_wrapper_email else "sender domain"
                results["score_reasons"].append(f"+5: Weak DKIM configuration for {domain_label}: {sender_domain_for_dkim}")
                print(f"[Onion] ⚠️ Weak DKIM configuration for {sender_domain_for_dkim}")
            elif dkim_info["security_level"] == "Medium":
                print(f"[Onion] ✓ Medium DKIM security for {sender_domain_for_dkim}")
            elif dkim_info["security_level"] == "High":
                results['score'] = max(0, results['score'] - 2)  # Slight reduction for good DKIM
                domain_label = "embedded phishing email domain" if is_wrapper_email else "sender domain"
                results["score_reasons"].append(f"-2: Strong DKIM configuration for {domain_label}: {sender_domain_for_dkim}")
                print(f"[Onion] ✓ Strong DKIM configuration for {sender_domain_for_dkim}")
                
        except Exception as dkim_error:
            print(f"[Onion] ⚠️ Error checking DKIM for {sender_domain_for_dkim}: {dkim_error}")
            traceback.print_exc()
            dkim_info = {
                "domain": sender_domain_for_dkim,
                "error": str(dkim_error),
                "selectors_found": [],
                "security_level": "Unknown",
                "is_wrapper_email": is_wrapper_email,
                "deep_extraction_used": bool(is_wrapper_email and embedded_headers)
            }
    elif sender_domain_for_dkim:
        # Domain found but DKIM check not available
        print(f"[Onion] ⚠️ Magic-Spoofing not available, skipping DKIM check for {sender_domain_for_dkim}")
        dkim_info = {
            "domain": sender_domain_for_dkim,
            "error": "Magic-Spoofing module not available. Please install required dependencies.",
            "selectors_found": [],
            "security_level": "Unknown",
            "is_wrapper_email": is_wrapper_email,
            "deep_extraction_used": bool(is_wrapper_email and embedded_headers)
        }
    
    # ========== MESSAGE DETAILS & SPOOFING ANALYSIS ==========
    print("[Onion] ====== Message Details & Spoofing Analysis ======")
    
    # Get headers to analyze (prefer embedded headers if wrapper detected)
    headers_to_analyze = embedded_headers if (is_wrapper_email and embedded_headers) else headers
    
    # Extract sender domain for spoofing check
    sender_domain_for_spoofing = None
    from_header = headers_to_analyze.get('From', '') or headers.get('From', '')
    if from_header and from_header not in ["N/A", "Unknown", ""]:
        try:
            email_match = re.search(r'[\w\.-]+@([\w\.-]+\.[a-zA-Z]{2,})', from_header)
            if email_match:
                sender_domain_for_spoofing = email_match.group(1).lower()
                print(f"[Onion] Extracted sender domain for spoofing check: {sender_domain_for_spoofing}")
        except Exception as e:
            print(f"[Onion] ⚠️ Error extracting sender domain: {e}")
    
    # Extract all Received headers
    received_headers = headers_to_analyze.get('Received', [])
    if not received_headers:
        # Try to get from main headers if not in embedded
        received_headers = headers.get('Received', [])
    
    # Extract domains from Received headers
    domains_from_received = extract_domains_from_received_headers(received_headers)
    print(f"[Onion] ✓ Extracted {len(domains_from_received)} unique domain(s) from Received headers: {', '.join(domains_from_received[:5])}")
    
    # Parse Authentication-Results header
    auth_results = None
    auth_results_header = headers_to_analyze.get('Authentication-Results', '') or headers_to_analyze.get('Authentication-Results-Original', '')
    if auth_results_header:
        auth_results = parse_authentication_results(auth_results_header)
        print(f"[Onion] ✓ Parsed Authentication-Results: SPF={auth_results.get('spf_result', 'N/A')}, DKIM={auth_results.get('dkim_result', 'N/A')}")
    
    # ========== COMPREHENSIVE SPOOFING CHECK ==========
    # Perform comprehensive spoofing check (SPF, DKIM, DMARC) on sender domain
    spoofing_analysis = {
        "sender_domain": sender_domain_for_spoofing,
        "spf": None,
        "dkim": None,
        "dmarc": None,
        "overall_risk": "Unknown",
        "issues": [],
        "recommendations": []
    }
    
    if sender_domain_for_spoofing:
        colors = setup_colors() if (SPF_CHECK_AVAILABLE or DKIM_CHECK_AVAILABLE or DMARC_CHECK_AVAILABLE) else {
            "white_bold": "", "info": "", "normal": "", "green": "", "red": ""
        }
        
        # SPF Check
        if SPF_CHECK_AVAILABLE:
            try:
                print(f"[Onion] 🔍 Checking SPF for spoofing analysis: {sender_domain_for_spoofing}")
                spf_result = check_spf(sender_domain_for_spoofing, colors)
                spoofing_analysis["spf"] = {
                    "exists": spf_result.get("exists", False),
                    "record": spf_result.get("record", ""),
                    "security_level": spf_result.get("security_level", "None"),
                    "all_mechanism": spf_result.get("all_mechanism", ""),
                    "issues": spf_result.get("issues", []),
                    "recommendations": spf_result.get("recommendations", [])
                }
                
                # Add to overall issues if SPF missing or weak
                if not spf_result.get("exists"):
                    spoofing_analysis["issues"].append(f"SPF: No SPF record found for {sender_domain_for_spoofing}")
                    spoofing_analysis["recommendations"].append(f"SPF: Implement SPF record for {sender_domain_for_spoofing}")
                    results['score'] += 15
                    results["score_reasons"].append(f"+15: No SPF record - domain can be easily spoofed: {sender_domain_for_spoofing}")
                elif spf_result.get("security_level") in ["None", "Low"]:
                    spoofing_analysis["issues"].append(f"SPF: Weak SPF configuration for {sender_domain_for_spoofing}")
                    results['score'] += 10
                    results["score_reasons"].append(f"+10: Weak SPF configuration allows spoofing: {sender_domain_for_spoofing}")
                elif spf_result.get("all_mechanism") == "~all":
                    spoofing_analysis["issues"].append(f"SPF: Using soft fail (~all) instead of hard fail (-all) for {sender_domain_for_spoofing}")
                    results['score'] += 5
                    results["score_reasons"].append(f"+5: SPF soft fail allows potential spoofing: {sender_domain_for_spoofing}")
                    
            except Exception as spf_spoof_error:
                print(f"[Onion] ⚠️ Error checking SPF for spoofing: {spf_spoof_error}")
                spoofing_analysis["spf"] = {"error": str(spf_spoof_error)}
        
        # DKIM Check (already done above, reuse if available)
        if DKIM_CHECK_AVAILABLE and dkim_info:
            spoofing_analysis["dkim"] = {
                "exists": len(dkim_info.get("selectors_found", [])) > 0,
                "security_level": dkim_info.get("security_level", "None"),
                "selectors_found": dkim_info.get("selectors_found", []),
                "issues": dkim_info.get("issues", []),
                "recommendations": dkim_info.get("recommendations", [])
            }
            
            # Add to overall issues if DKIM missing or weak
            if not spoofing_analysis["dkim"]["exists"]:
                spoofing_analysis["issues"].append(f"DKIM: No DKIM configuration found for {sender_domain_for_spoofing}")
                spoofing_analysis["recommendations"].append(f"DKIM: Implement DKIM signing for {sender_domain_for_spoofing}")
            elif spoofing_analysis["dkim"]["security_level"] in ["None", "Low"]:
                spoofing_analysis["issues"].append(f"DKIM: Weak DKIM configuration for {sender_domain_for_spoofing}")
        
        # DMARC Check - Parse DMARC directly for detailed info
        if DMARC_CHECK_AVAILABLE:
            try:
                print(f"[Onion] 🔍 Checking DMARC for spoofing analysis: {sender_domain_for_spoofing}")
                
                # Parse DMARC directly for more details
                try:
                    import pydig
                    dmarc_records = pydig.query('_dmarc.' + sender_domain_for_spoofing, 'TXT')
                    dmarc_found = False
                    dmarc_record = None
                    for record in dmarc_records:
                        if "v=DMARC1" in record:
                            dmarc_found = True
                            dmarc_record = record
                            break
                    
                    dmarc_info = {
                        "exists": dmarc_found,
                        "record": dmarc_record,
                        "policy": None,
                        "subdomain_policy": None,
                        "pct": None,
                        "rua": [],
                        "ruf": [],
                        "security_level": "None",
                        "issues": [],
                        "recommendations": []
                    }
                    
                    if dmarc_found and dmarc_record:
                        # Parse DMARC record using same logic as analyze_dmarc_record
                        clean_record = dmarc_record.strip('"\'')
                        
                        # Parse fields
                        fields = {}
                        for field in clean_record.split(';'):
                            field = field.strip()
                            if '=' in field:
                                key, value = field.split('=', 1)
                                key = key.strip().lower()
                                value = value.strip()
                                fields[key] = value
                        
                        # Extract policy
                        if 'p' in fields:
                            dmarc_info["policy"] = fields['p'].lower()
                        if 'sp' in fields:
                            dmarc_info["subdomain_policy"] = fields['sp'].lower()
                        if 'pct' in fields:
                            try:
                                dmarc_info["pct"] = int(fields['pct'].strip('"\''))
                            except:
                                dmarc_info["pct"] = 100  # Default
                        else:
                            dmarc_info["pct"] = 100  # Default
                        if 'rua' in fields:
                            dmarc_info["rua"] = [addr.strip() for addr in fields['rua'].split(',')]
                        if 'ruf' in fields:
                            dmarc_info["ruf"] = [addr.strip() for addr in fields['ruf'].split(',')]
                        
                        # Determine security level
                        if dmarc_info["policy"] == "reject":
                            dmarc_info["security_level"] = "High"
                        elif dmarc_info["policy"] == "quarantine":
                            dmarc_info["security_level"] = "Medium"
                            dmarc_info["issues"].append("DMARC policy is 'quarantine' - should be 'reject' for maximum protection")
                            dmarc_info["recommendations"].append("Change DMARC policy from 'quarantine' to 'reject'")
                        elif dmarc_info["policy"] == "none":
                            dmarc_info["security_level"] = "Low"
                            dmarc_info["issues"].append("DMARC policy is 'none' - provides no protection against spoofing")
                            dmarc_info["recommendations"].append("Change DMARC policy from 'none' to 'quarantine' or 'reject'")
                        
                        # Check pct (percentage)
                        if dmarc_info.get("pct", 100) < 100:
                            dmarc_info["issues"].append(f"DMARC only applies to {dmarc_info['pct']}% of emails (pct={dmarc_info['pct']})")
                            dmarc_info["recommendations"].append("Set pct=100 or remove pct to apply DMARC to all emails")
                        
                        spoofing_analysis["dmarc"] = dmarc_info
                    else:
                        spoofing_analysis["dmarc"] = {
                            "exists": False, 
                            "security_level": "None",
                            "issues": ["No DMARC record found"],
                            "recommendations": ["Implement DMARC record to prevent email spoofing"]
                        }
                    
                except Exception as dmarc_parse_error:
                    print(f"[Onion] ⚠️ Error parsing DMARC: {dmarc_parse_error}")
                    traceback.print_exc()
                    spoofing_analysis["dmarc"] = {"error": str(dmarc_parse_error)}
                
                # Add to overall issues if DMARC missing or weak
                if not spoofing_analysis["dmarc"].get("exists"):
                    spoofing_analysis["issues"].append(f"DMARC: No DMARC record found for {sender_domain_for_spoofing}")
                    spoofing_analysis["recommendations"].append(f"DMARC: Implement DMARC policy for {sender_domain_for_spoofing}")
                    results['score'] += 20
                    results["score_reasons"].append(f"+20: No DMARC record - high spoofing risk: {sender_domain_for_spoofing}")
                elif spoofing_analysis["dmarc"].get("policy") == "none":
                    spoofing_analysis["issues"].append(f"DMARC: Policy set to 'none' (no protection) for {sender_domain_for_spoofing}")
                    results['score'] += 15
                    results["score_reasons"].append(f"+15: DMARC policy 'none' provides no protection: {sender_domain_for_spoofing}")
                elif spoofing_analysis["dmarc"].get("policy") == "quarantine":
                    results['score'] += 5
                    results["score_reasons"].append(f"+5: DMARC policy 'quarantine' (should be 'reject'): {sender_domain_for_spoofing}")
                    
            except Exception as dmarc_spoof_error:
                print(f"[Onion] ⚠️ Error checking DMARC for spoofing: {dmarc_spoof_error}")
                spoofing_analysis["dmarc"] = {"error": str(dmarc_spoof_error)}
        
        # Calculate overall spoofing risk
        spf_level = spoofing_analysis["spf"].get("security_level", "None") if spoofing_analysis["spf"] else "Unknown"
        dkim_level = spoofing_analysis["dkim"].get("security_level", "None") if spoofing_analysis["dkim"] else "Unknown"
        dmarc_level = spoofing_analysis["dmarc"].get("security_level", "None") if spoofing_analysis["dmarc"] else "Unknown"
        
        # Risk assessment
        if (spf_level in ["None", "Unknown"] and dkim_level in ["None", "Unknown"] and dmarc_level in ["None", "Unknown"]):
            spoofing_analysis["overall_risk"] = "CRITICAL - No email authentication configured"
        elif (spf_level in ["None", "Low"] or dkim_level in ["None", "Low"] or dmarc_level in ["None", "Low"]):
            spoofing_analysis["overall_risk"] = "HIGH - Weak email authentication"
        elif (spf_level == "Medium" or dkim_level == "Medium" or dmarc_level == "Medium"):
            spoofing_analysis["overall_risk"] = "MEDIUM - Moderate protection"
        else:
            spoofing_analysis["overall_risk"] = "LOW - Strong email authentication"
        
        print(f"[Onion] ✓ Spoofing analysis complete: {spoofing_analysis['overall_risk']}")
    
    # Perform SPF checks on all domains found in Received headers
    spf_results = []
    if SPF_CHECK_AVAILABLE and domains_from_received:
        colors = setup_colors() if SPF_CHECK_AVAILABLE else {
            "white_bold": "", "info": "", "normal": "", "green": "", "red": ""
        }
        
        for domain in domains_from_received:
            try:
                print(f"[Onion] Checking SPF for domain: {domain}")
                spf_result = check_spf(domain, colors)
                
                spf_info = {
                    "domain": domain,
                    "exists": spf_result.get("exists", False),
                    "record": spf_result.get("record", ""),
                    "security_level": spf_result.get("security_level", "None"),
                    "all_mechanism": spf_result.get("all_mechanism", ""),
                    "issues": spf_result.get("issues", []),
                    "recommendations": spf_result.get("recommendations", []),
                    "mechanisms": spf_result.get("mechanisms", []),
                    "has_hard_fail": spf_result.get("all_mechanism") == "-all",
                    "has_soft_fail": spf_result.get("all_mechanism") == "~all"
                }
                
                # Check Authentication-Results for SPF result
                # Match by domain if available, otherwise use first/only SPF result
                if auth_results:
                    auth_spf_domain = auth_results.get("spf_domain")
                    auth_spf_result = auth_results.get("spf_result", "")
                    
                    # Match domain if specified, or if only one domain in Received headers
                    if (auth_spf_domain and auth_spf_domain == domain) or (not auth_spf_domain and len(domains_from_received) == 1):
                        spf_info["auth_result"] = auth_spf_result
                        auth_spf_result_lower = auth_spf_result.lower() if auth_spf_result else ""
                        if auth_spf_result_lower == "fail":
                            spf_info["hard_fail_detected"] = True
                        elif auth_spf_result_lower in ["softfail", "soft fail"]:
                            spf_info["soft_fail_detected"] = True
                
                spf_results.append(spf_info)
                
                # Add score adjustments based on SPF results
                if not spf_info["exists"]:
                    results['score'] += 10
                    results["score_reasons"].append(f"+10: No SPF record found for domain in Received headers: {domain}")
                    print(f"[Onion] ⚠️ No SPF configuration found for {domain}")
                elif spf_info["has_soft_fail"]:
                    results['score'] += 5
                    results["score_reasons"].append(f"+5: Soft SPF fail (~all) configured for domain: {domain}")
                    print(f"[Onion] ⚠️ Soft SPF fail (~all) for {domain}")
                elif spf_info["has_hard_fail"]:
                    print(f"[Onion] ✓ Hard SPF fail (-all) configured for {domain} - good security")
                
                # Check Authentication-Results for hard/soft fail
                if spf_info.get("hard_fail_detected"):
                    results['score'] += 20
                    results["score_reasons"].append(f"+20: SPF hard fail detected in Authentication-Results for {domain}")
                    print(f"[Onion] 🚨 SPF HARD FAIL detected for {domain}")
                elif spf_info.get("soft_fail_detected"):
                    results['score'] += 10
                    results["score_reasons"].append(f"+10: SPF soft fail detected in Authentication-Results for {domain}")
                    print(f"[Onion] ⚠️ SPF SOFT FAIL detected for {domain}")
                    
            except Exception as spf_error:
                print(f"[Onion] ⚠️ Error checking SPF for {domain}: {spf_error}")
                spf_results.append({
                    "domain": domain,
                    "error": str(spf_error),
                    "exists": False,
                    "security_level": "Unknown"
                })
    
    # Prepare full Message Details (all headers)
    message_details = {}
    for key, value in headers_to_analyze.items():
        # Skip internal metadata headers
        if not key.startswith('_'):
            if isinstance(value, list):
                message_details[key] = value
            else:
                message_details[key] = str(value) if value else ""
    
    # Final Verdict Calculation
    # Check if verdict was already set to MALICIOUS due to negative community_score
    # If not, calculate based on score
    if results.get('verdict') != "MALICIOUS":
        if results['score'] >= 100:
            results['verdict'] = "MALICIOUS"
            results["verdict_reasons"].append(f"Threat score ({results['score']}) exceeds threshold (>= 100)")
        elif results['score'] >= 50:
            results['verdict'] = "SUSPICIOUS"
            results["verdict_reasons"].append(f"Threat score ({results['score']}) is in suspicious range (50-99)")
        elif len(results['urls']) > 0 or len(results['ips']) > 0 or len(results['attachments']) > 0:
            results['verdict'] = "CLEAN"
            results["verdict_reasons"].append(f"Email contains IOCs but threat score ({results['score']}) is below suspicious threshold")
    else:
            results['verdict'] = "UNKNOWN"
            results["verdict_reasons"].append("No IOCs found and threat score is low")
    
    # Sort URLs by threat level (malicious first)
    final_url_objects.sort(key=lambda x: (x.get("malicious", 0) * 1000 + x.get("suspicious", 0)), reverse=True)
    
    # ========== FINAL RESULT ==========
    print(f"[Onion] ====== Analysis Complete ======")
    print(f"[Onion] Verdict: {results['verdict']}, Score: {results['score']}")
    print(f"[Onion] URLs: {len(results['urls'])}, IPs: {len(results['ips'])}, Attachments: {len(results['attachments'])}")
    
    # Prepare body preview - prefer HTML if available, otherwise use text
    body_preview_text = ""
    body_preview_html = None
    
    # Use msgreader body_html if available (from msg-analyzer via msg_reader_wrapper)
    # This ensures we use the full email body extracted by msgreader
    if msg_body_html:
        # Use full HTML body from msgreader (not truncated)
        body_preview_html = msg_body_html
        # Extract plain text from HTML for preview
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(msg_body_html, 'html.parser')
            body_preview_text = soup.get_text(separator=' ', strip=True)
        except:
            body_preview_text = msg_body_text if msg_body_text else ""
    elif msg_body_text:
        # Use full text body from msgreader
        body_preview_text = msg_body_text
    else:
        body_preview_text = base_text if base_text else ""
    
    return {
        "verdict": results["verdict"],
        "score": results["score"],
        "file_type": results["file_type"],
        "file_hash": file_hash,
        "body_preview": body_preview_text,
        "body_html": body_preview_html,  # Include HTML version if available
        "urls": final_url_objects,
        "ips": final_ip_objects,
        "attachments": results["attachments"],
        "headers": headers,
        "warnings": results["warnings"],
        "keywords_found": [],
        "total_urls_found": len(results["urls"]),
        "urls_checked": len(final_url_objects),
        "spoofing_warning": "\n".join([w for w in results["warnings"] if "spoofing" in w.lower() or "polyglot" in w.lower()]) if results["warnings"] else None,
        "sender_domain": sender_domain_info,  # Include sender domain threat intelligence
        "dkim_info": dkim_info,  # Include DKIM check results
        "score_reasons": results.get("score_reasons", []),  # Include threat score reasons
        "verdict_reasons": results.get("verdict_reasons", []),  # Include verdict reasons
        "message_details": message_details,  # Full Message Details (all headers)
        "spf_results": spf_results,  # SPF check results for all domains
        "auth_results": auth_results,  # Parsed Authentication-Results
        "domains_from_received": domains_from_received,  # Domains extracted from Received headers
        "spoofing_analysis": spoofing_analysis  # Comprehensive spoofing analysis (SPF/DKIM/DMARC)
    }

@app.route("/analyze_email", methods=["POST"])
def analyze_email():
    """
    Route to handle email file upload and analysis.
    
    Smart File Handling Strategy:
    1. Reads file content into memory first (no permanent disk write)
    2. Creates temporary file ONLY when needed (analysis libraries require physical file paths)
    3. Analyzes the temporary file
    4. AUTOMATICALLY deletes temporary file after analysis completes
    
    IMPORTANT: The file is NEVER saved permanently - it's created temporarily, analyzed, and deleted automatically.
    This ensures files are only used during analysis and never stored on disk permanently.
    
    Uses low-level file descriptor (mkstemp) to prevent file locking issues.
    """
    print("[Email Analysis] ====== Starting email analysis request ======")
    
    if not EXTRACT_MSG_AVAILABLE:
        print("[Email Analysis] ERROR: extract-msg library not installed")
        return jsonify({"error": "extract-msg library not installed"}), 400
    
    # Step 1: Input Validation - Check if file exists and filename ends with .msg
    if 'file' not in request.files:
        print("[Email Analysis] ERROR: No file in request.files")
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    print(f"[Email Analysis] File received: filename='{file.filename}', content_type='{file.content_type}'")
    
    # Handle files without filename (Outlook drag & drop)
    if not file.filename or file.filename == '':
        # Try to determine file type from content_type
        content_type = file.content_type or ''
        if 'ms-outlook' in content_type or 'x-msmsg' in content_type:
            file.filename = 'email.msg'
            print("[Email Analysis] Assigned default filename 'email.msg' based on content_type")
        elif 'message/rfc822' in content_type:
            file.filename = 'email.eml'
            print("[Email Analysis] Assigned default filename 'email.eml' based on content_type")
        else:
            # Default to .msg for Outlook drags
            file.filename = 'email.msg'
            print("[Email Analysis] Assigned default filename 'email.msg' (no content_type detected)")
    
    # Check file extension (case insensitive) - support both .msg and .eml
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ['.msg', '.eml']:
        # If no extension, try to determine from content_type
        content_type = file.content_type or ''
        if 'ms-outlook' in content_type or 'x-msmsg' in content_type:
            file.filename = file.filename if file.filename.endswith('.msg') else file.filename + '.msg'
            file_ext = '.msg'
            print(f"[Email Analysis] Added .msg extension based on content_type. New filename: {file.filename}")
        elif 'message/rfc822' in content_type:
            file.filename = file.filename if file.filename.endswith('.eml') else file.filename + '.eml'
            file_ext = '.eml'
            print(f"[Email Analysis] Added .eml extension based on content_type. New filename: {file.filename}")
        else:
            # Default to .msg if no content_type
            if not file_ext:
                file.filename = file.filename + '.msg'
                file_ext = '.msg'
                print(f"[Email Analysis] Added .msg extension (default). New filename: {file.filename}")
            else:
                print(f"[Email Analysis] ERROR: Invalid file extension. Filename: {file.filename}")
                return jsonify({"error": "Only .msg and .eml files are supported"}), 400
    
    # ========== STEP 1: RECEPTION & SAVING ==========
    # Step 1.1: Read file content from request (keep in memory first)
    file.seek(0)  # CRITICAL: Ensure file stream is at the beginning (Flask may have read it)
    file_content = file.read()
    file_size = len(file_content)
    print(f"[Email Analysis] Read {file_size} bytes from uploaded file (stored in memory)")
    
    # Step 1.2: Validate file size (handle zero-size files)
    if file_size == 0:
        print("[Email Analysis] ERROR: File size is 0 bytes - empty file detected")
        return jsonify({
            "error": "Empty file (size 0). Please ensure the file is valid and try again.",
            "verdict": "CORRUPTED",
            "score": 0,
            "file_type": "UNKNOWN"
        }), 400
    
    # Step 1.3: Smart File Handling Strategy
    # Note: analyze_email_content requires a physical file path (msg-parser, extract-msg, msg-reader all need files)
    # Therefore, we will create a temporary file, analyze it, and delete it automatically
    # The file is NEVER saved permanently - it's created, analyzed, and deleted in one flow
    
    fd = None
    temp_path = None
    file_saved = False  # Track if we successfully saved the file
    
    # Wrap all operations (saving + analysis + cleanup) in try-finally for proper cleanup
    try:
        # Step 1.4: Create temporary physical file on disk (required for analysis libraries)
        print("[Email Analysis] Creating temporary file using mkstemp (automatic cleanup after analysis)...")
        # Use mkstemp to get a file descriptor and path (preserve original extension)
        fd, temp_path = tempfile.mkstemp(suffix=file_ext)
        print(f"[Email Analysis] Temp file path: {temp_path} (will be auto-deleted after analysis)")
        
        # Step 1.5: Write data from memory to temp file
        # os.fdopen() automatically closes the file descriptor when the with block exits
        try:
            # Open file descriptor as writable binary file
            with os.fdopen(fd, 'wb') as f:
                # Write file content from memory to temp file
                f.write(file_content)
                # CRITICAL: Flush and sync to force OS write to disk (must be inside with block)
                f.flush()
                os.fsync(f.fileno())
            # File descriptor is automatically closed by os.fdopen() when exiting the with block
            fd = None  # Mark as closed to prevent double-close
            print("[Email Analysis] File content written from memory to disk")
        except Exception as write_err:
            # If write fails, the file descriptor is already closed by os.fdopen()
            # Only need to mark it as None to prevent double-close attempts
            fd = None
            raise write_err
        
        # Step 1.7: Verify file was written correctly
        saved_file_size = os.path.getsize(temp_path)
        print(f"[Email Analysis] File saved to disk successfully. Size: {saved_file_size} bytes")
        
        if saved_file_size != file_size:
            print(f"[Email Analysis] WARNING: Size mismatch! Memory: {file_size}, Disk: {saved_file_size}")
        
        if saved_file_size == 0:
            print("[Email Analysis] ERROR: Saved file is empty (0 bytes)")
            file_saved = False
            # Cleanup empty file immediately
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                    temp_path = None
                except:
                    pass
            return jsonify({
                "error": "Error saving file - file was saved as empty. Please try again.",
                "verdict": "CORRUPTED",
                "score": 0,
                "file_type": "UNKNOWN"
            }), 500
        
        file_saved = True
        print("[Email Analysis] ✓ File ready for analysis (temporary, will be auto-deleted)")
        
        # ========== STEP 2: ANALYSIS ==========
        # Step 2.1: Analyze email content using temporary file path
        # Note: The file will be automatically deleted after analysis completes (see finally block)
        if not file_saved or not temp_path:
            raise Exception("File was not saved successfully - cannot proceed with analysis")
        
        print(f"[Email Analysis] Starting analysis of temporary file: {temp_path}")
        print("[Email Analysis] Note: File is temporary and will be auto-deleted after analysis")
        result = analyze_email_content(temp_path)
        
        # Step 2.2: Validate result
        if result is None:
            print("[Email Analysis] WARNING: analyze_email_content returned None")
            result = {"error": "Analysis failed: No results returned"}
        
        if not isinstance(result, dict):
            print(f"[Email Analysis] WARNING: analyze_email_content returned {type(result)}")
            result = {"error": f"Analysis failed: Invalid result type ({type(result).__name__})"}
        
        print("[Email Analysis] Analysis completed successfully")
        # Log attachments count for debugging
        attachments_count = len(result.get('attachments', [])) if result else 0
        print(f"[Email Analysis] Result contains {attachments_count} attachment(s)")
        if attachments_count > 0:
            print(f"[Email Analysis] Attachment filenames: {[att.get('filename', 'Unknown') for att in result.get('attachments', [])[:5]]}")
        
        # Check if result contains an error
        if isinstance(result, dict) and 'error' in result:
            print(f"[Email Analysis] Analysis returned error: {result['error']}")
        
        # Generate graph_data for Threat Constellation Map
        graph_data = []
        if isinstance(result, dict) and 'error' not in result:
            try:
                # Extract IPs and URLs from result
                ips = result.get('ips', [])
                urls = result.get('urls', [])
                
                # Build graph data using the same function as log analyzer
                # We need to create a simplified analysis dict for build_graph_data
                analysis_dict = {
                    'ips': [ip.get('ip', ip) if isinstance(ip, dict) else ip for ip in ips],
                    'urls': [url.get('url', url) if isinstance(url, dict) else url for url in urls],
                    'relationships': {}
                }
                
                # Create empty VT and AbuseIPDB results (we could enhance this later with actual lookups)
                vt_results = {
                    'ips': {},
                    'urls': {},
                    'domains': {},
                    'hashes': {}
                }
                abuseipdb_results = {}
                
                # Build graph data
                graph_data = build_graph_data(analysis_dict, vt_results, abuseipdb_results, {})
                
            except Exception as graph_error:
                print(f"[Email Analysis] Warning: Failed to generate graph_data: {graph_error}")
                traceback.print_exc()
                graph_data = []
        
        # Add graph_data to result
        if isinstance(result, dict):
            result['graph_data'] = graph_data
        
        return jsonify(result)
            
    except Exception as processing_error:
        print(f"[Email Analysis] ERROR during processing: {processing_error}")
        traceback.print_exc()
        return jsonify({
            "error": f"Error processing file: {str(processing_error)}",
            "verdict": "CORRUPTED",
            "score": 0,
            "file_type": "UNKNOWN"
        }), 500
        
    # ========== STEP 3: CLEANUP (AUTOMATIC) ==========
    finally:
        # Step 3.1: CRITICAL - Delete temporary file immediately after analysis completes
        # This ensures the file is NEVER permanently saved - it's only used during analysis
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
                print(f"[Email Analysis] ✓ Temporary file automatically deleted: {temp_path}")
                print("[Email Analysis] ✓ File was never permanently saved - cleaned up successfully")
            except Exception as cleanup_err:
                print(f"[Email Analysis] ⚠️ WARNING: Could not delete temp file: {cleanup_err}")
                print(f"[Email Analysis] ⚠️ Manual cleanup may be needed for: {temp_path}")
        
        # Step 3.2: Ensure file descriptor is closed (safety check)
        # Note: os.fdopen() automatically closes the file descriptor, so this is just a safety check
        # for edge cases where the with block might not have executed
        if fd is not None:
            try:
                os.close(fd)
                fd = None
                print("[Email Analysis] ✓ File descriptor closed in finally block (safety check)")
            except (OSError, ValueError) as close_err:
                # File descriptor was already closed by os.fdopen() - this is expected and safe to ignore
                fd = None
                # Only log if it's not the expected "Bad file descriptor" error
                if "Bad file descriptor" not in str(close_err):
                    print(f"[Email Analysis] ⚠️ Warning: Could not close file descriptor: {close_err}")

# ================================

# ======= PDF TO HTML CONVERSION =======
def convert_pdf_to_html(pdf_path: str, output_dir: str = None) -> dict:
    """
    Convert PDF file to HTML using pdf2htmlEX from Pdf-Viewer directory.
    
    Args:
        pdf_path: Path to the PDF file
        output_dir: Directory to save HTML output (default: same as PDF)
    
    Returns:
        dict with 'success', 'html_path', 'error' keys
    """
    if not os.path.exists(pdf_path):
        return {'success': False, 'error': f'PDF file not found: {pdf_path}'}
    
    # Determine output directory
    if not output_dir:
        output_dir = os.path.dirname(pdf_path)
    
    # Output HTML filename (same name as PDF but .html extension)
    pdf_basename = os.path.splitext(os.path.basename(pdf_path))[0]
    output_html = os.path.join(output_dir, f"{pdf_basename}.html")
    
    # Find pdf2htmlEX executable
    # Check common locations and current directory
    pdf2html_exe = None
    
    # Check environment variable first (for custom path)
    custom_path = os.environ.get('PDF2HTMLEX_PATH')
    if custom_path and os.path.exists(custom_path):
        pdf2html_exe = custom_path
        print(f"[PDF2HTML] Using custom path from PDF2HTMLEX_PATH: {custom_path}")
    
    # Check if pdf2htmlEX is in PATH
    if not pdf2html_exe:
        exe_name = 'pdf2htmlEX.exe' if os.name == 'nt' else 'pdf2htmlEX'
        try:
            result = subprocess.run([exe_name, '--version'], 
                                  capture_output=True, timeout=5)
            if result.returncode == 0:
                pdf2html_exe = exe_name
                print(f"[PDF2HTML] Found pdf2htmlEX in PATH")
        except:
            pass
        
    # Check in Pdf-Viewer directory (for local builds)
    if not pdf2html_exe:
        pdf_viewer_paths = [
            os.path.join('Pdf-Viewer', 'pdf2htmlEX', 'build', 'pdf2htmlEX', 'pdf2htmlEX.exe'),
            os.path.join('Pdf-Viewer', 'build', 'pdf2htmlEX', 'pdf2htmlEX.exe'),
            os.path.join('Pdf-Viewer', 'pdf2htmlEX', 'pdf2htmlEX.exe'),
            os.path.join('Pdf-Viewer', 'pdf2htmlEX.exe'),
            # Linux/Unix paths
            os.path.join('Pdf-Viewer', 'pdf2htmlEX', 'build', 'pdf2htmlEX', 'pdf2htmlEX'),
            os.path.join('Pdf-Viewer', 'build', 'pdf2htmlEX', 'pdf2htmlEX'),
            os.path.join('Pdf-Viewer', 'pdf2htmlEX', 'pdf2htmlEX'),
            # Common Windows installation paths
            r'C:\Program Files\pdf2htmlEX\pdf2htmlEX.exe',
            r'C:\Program Files (x86)\pdf2htmlEX\pdf2htmlEX.exe',
            r'C:\tools\pdf2htmlEX\pdf2htmlEX.exe',
            os.path.join(os.path.expanduser('~'), 'tools', 'pdf2htmlEX', 'pdf2htmlEX.exe'),
        ]
        for path in pdf_viewer_paths:
            if os.path.exists(path):
                pdf2html_exe = path
                print(f"[PDF2HTML] Found pdf2htmlEX at: {path}")
                break
    
    if not pdf2html_exe:
        error_msg = '''pdf2htmlEX not found. Please install pdf2htmlEX:
        
For Windows:
  1. Download from: https://github.com/pdf2htmlEX/pdf2htmlEX/releases
  2. Extract pdf2htmlEX.exe to a folder (e.g., C:\\tools\\pdf2htmlEX\\)
  3. Add folder to PATH or set environment variable: PDF2HTMLEX_PATH=C:\\tools\\pdf2htmlEX\\pdf2htmlEX.exe

For Linux/WSL:
  1. Run: sudo apt-get install pdf2htmlEX
  2. Or build from source in Pdf-Viewer directory: ./buildScripts/buildInstallLocallyApt

See INSTALL_PDF2HTMLEX.md for detailed instructions.'''
        return {
            'success': False, 
            'error': error_msg
        }
    
    try:
        # Run pdf2htmlEX with options optimized for web viewing
        cmd = [
            pdf2html_exe,
            '--zoom', '1.5',  # Zoom for better readability
            '--font-size-multiplier', '1',
            '--fit-width', '1024',  # Fit to 1024px width
            '--page-filename', 'page%03d.html',  # Page filenames
            '--external-hint',  # Use external resources
            '--embed-css', '1',  # Embed CSS
            '--embed-font', '1',  # Embed fonts
            '--embed-image', '1',  # Embed images
            pdf_path,
            output_html
        ]
        
        print(f"[PDF2HTML] Converting {pdf_path} to HTML...")
        result = subprocess.run(cmd, capture_output=True, timeout=60, text=True)
        
        if result.returncode == 0 and os.path.exists(output_html):
            # Get all generated files (HTML + resources)
            generated_files = [output_html]
            
            # Find page files if multi-page
            base_dir = os.path.dirname(output_html)
            base_name = os.path.splitext(os.path.basename(output_html))[0]
            page_pattern = os.path.join(base_dir, f"{base_name}.page*.html")
            import glob
            generated_files.extend(glob.glob(page_pattern))
            
            return {
                'success': True,
                'html_path': output_html,
                'generated_files': generated_files,
                'error': None
            }
        else:
            error_msg = result.stderr or result.stdout or 'Unknown error'
            return {
                'success': False,
                'error': f'pdf2htmlEX conversion failed: {error_msg}'
            }
            
    except subprocess.TimeoutExpired:
        return {'success': False, 'error': 'PDF conversion timeout (60 seconds)'}
    except Exception as e:
        return {'success': False, 'error': f'Error converting PDF: {str(e)}'}

@app.route("/view_pdf/<pdf_hash>")
def view_pdf_as_html(pdf_hash: str):
    """
    Route to view PDF attachment as HTML using pdf2htmlEX.
    Uses cached PDF content from email analysis.
    """
    print(f"[PDF Viewer] Request to view PDF with hash: {pdf_hash[:16]}...")
    
    # Get cached PDF
    cached_pdf = get_cached_pdf(pdf_hash)
    if not cached_pdf:
        return jsonify({'error': 'PDF not found or expired from cache'}), 404
    
    pdf_content = cached_pdf['content']
    pdf_filename = cached_pdf['filename']
    
    # Create temporary directory for conversion
    temp_dir = tempfile.mkdtemp(prefix='pdf_viewer_')
    temp_pdf_path = os.path.join(temp_dir, f"{pdf_hash[:16]}.pdf")
    temp_html_path = os.path.join(temp_dir, f"{pdf_hash[:16]}.html")
    
    try:
        # Write PDF to temp file
        with open(temp_pdf_path, 'wb') as f:
            f.write(pdf_content)
        
        print(f"[PDF Viewer] Converting PDF to HTML: {temp_pdf_path}")
        conversion_result = convert_pdf_to_html(temp_pdf_path, temp_dir)
        
        if not conversion_result['success']:
            # Cleanup
            shutil.rmtree(temp_dir, ignore_errors=True)
            return jsonify({'error': conversion_result['error']}), 500
        
        html_path = conversion_result['html_path']
        
        # Read HTML content
        with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        # Serve HTML directly (Flask will handle it)
        # Note: Resources (fonts, images) need to be served separately
        # For now, we'll embed everything or serve as static
        return html_content, 200, {'Content-Type': 'text/html; charset=utf-8'}
        
    except Exception as e:
        print(f"[PDF Viewer] Error serving PDF HTML: {e}")
        traceback.print_exc()
        # Cleanup on error
        shutil.rmtree(temp_dir, ignore_errors=True)
        return jsonify({'error': f'Error viewing PDF: {str(e)}'}), 500

@app.route("/pdf_resource/<pdf_hash>/<path:resource_path>")
def serve_pdf_resource(pdf_hash: str, resource_path: str):
    """Serve resources (fonts, images, CSS) for converted PDF HTML."""
    # Find the temp directory for this PDF hash
    temp_base = tempfile.gettempdir()
    temp_dirs = [d for d in os.listdir(temp_base) if d.startswith('pdf_viewer_')]
    
    for temp_dir_name in temp_dirs:
        temp_dir = os.path.join(temp_base, temp_dir_name)
        if not os.path.isdir(temp_dir):
            continue
        
        # Check if PDF hash matches
        pdf_file = os.path.join(temp_dir, f"{pdf_hash[:16]}.pdf")
        if os.path.exists(pdf_file):
            resource_file = os.path.join(temp_dir, resource_path)
            if os.path.exists(resource_file) and os.path.abspath(resource_file).startswith(os.path.abspath(temp_dir)):
                # Serve the file
                from flask import send_from_directory
                return send_from_directory(temp_dir, resource_path)
    
    return jsonify({'error': 'Resource not found'}), 404

@app.route("/analyze_file/<file_hash>")
def analyze_file(file_hash: str):
    """
    Analyze a file by hash - extract links and check them in VirusTotal.
    Returns JSON with links and their VirusTotal scores.
    """
    print(f"[File Analyzer] Analyzing file with hash: {file_hash[:16]}...")
    
    # Get cached attachment
    cached_attachment = get_cached_attachment(file_hash)
    if not cached_attachment:
        return jsonify({'error': 'File not found in cache'}), 404
    
    file_content = cached_attachment['content']
    filename = cached_attachment.get('filename', 'unknown')
    mime_type = cached_attachment.get('mime_type', '')
    
    # Extract URLs based on file type
    extracted_urls = []
    
    try:
        # Determine file type
        file_ext = os.path.splitext(filename)[1].lower() if filename else ''
        detected_type = mime_type.lower() if mime_type else ''
        
        # DOCX files
        if file_ext == '.docx' or 'wordprocessingml' in detected_type or 'docx' in detected_type:
            print("[File Analyzer] Detected DOCX file, extracting URLs...")
            extracted_urls = extract_urls_from_docx(file_content)
        
        # PDF files
        elif file_ext == '.pdf' or 'pdf' in detected_type:
            print("[File Analyzer] Detected PDF file, extracting URLs...")
            
            # Try PyPDF2 approach first (like LinksReader) - extracts hyperlinks from annotations
            pdf_urls_from_annotations = []
            if PYPDF2_AVAILABLE:
                try:
                    pdf_file = PyPDF2.PdfReader(io.BytesIO(file_content))
                    for page_num in range(len(pdf_file.pages)):
                        page = pdf_file.pages[page_num]
                        if '/Annots' in page:
                            annotations = page['/Annots']
                            for annotation in annotations:
                                a_entry = annotation.get_object().get('/A')
                                if isinstance(a_entry, PyPDF2.generic.DictionaryObject):
                                    uri = a_entry.get('/URI')
                                    if uri:
                                        pdf_urls_from_annotations.append(uri)
                    if pdf_urls_from_annotations:
                        extracted_urls.extend(pdf_urls_from_annotations)
                        print(f"[File Analyzer] Extracted {len(pdf_urls_from_annotations)} URLs from PDF annotations using PyPDF2")
                except Exception as e:
                    print(f"[File Analyzer] Error extracting URLs from PDF annotations (PyPDF2): {e}")
            
            # Also extract URLs from text using pdfminer (complements PyPDF2)
            if PDFMINER_AVAILABLE:
                try:
                    fd, tmp_path = tempfile.mkstemp(suffix='.pdf')
                    try:
                        with os.fdopen(fd, 'wb') as f:
                            f.write(file_content)
                        extracted_text = extract_text(tmp_path, laparams=LAParams())
                        if extracted_text:
                            url_matches = _URL_REGEX_PATTERN.findall(extracted_text)
                            extracted_urls.extend(url_matches)
                            print(f"[File Analyzer] Extracted {len(url_matches)} URLs from PDF text using pdfminer")
                    finally:
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
                except Exception as e:
                    print(f"[File Analyzer] Error extracting from PDF text: {e}")
            
            # Remove duplicates from all extracted URLs
            if extracted_urls:
                extracted_urls = list(set(extracted_urls))
                print(f"[File Analyzer] Total unique URLs from PDF: {len(extracted_urls)}")
        
        # XLSX files (Excel)
        elif file_ext in ['.xlsx', '.xls'] or 'spreadsheetml' in detected_type or 'excel' in detected_type:
            print("[File Analyzer] Detected XLSX file, extracting URLs...")
            try:
                import openpyxl
                from io import BytesIO
                workbook = openpyxl.load_workbook(BytesIO(file_content))
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.hyperlink and cell.hyperlink.target:
                                extracted_urls.append(cell.hyperlink.target)
                            elif cell.value and isinstance(cell.value, str):
                                url_matches = _URL_REGEX_PATTERN.findall(cell.value)
                                extracted_urls.extend(url_matches)
                extracted_urls = list(set(extracted_urls))
            except Exception as e:
                print(f"[File Analyzer] Error extracting from XLSX: {e}")
                # Fallback: try regex on raw content
                try:
                    content_str = file_content.decode('utf-8', errors='ignore')
                    url_matches = _URL_REGEX_PATTERN.findall(content_str)
                    extracted_urls = list(set(url_matches))
                except:
                    pass
        
        # Text files
        else:
            print("[File Analyzer] Detected text file, extracting URLs...")
            try:
                # Try UTF-8 first
                content_str = file_content.decode('utf-8', errors='ignore')
            except:
                try:
                    # Try other encodings
                    content_str = file_content.decode('latin-1', errors='ignore')
                except:
                    content_str = ''
            
            if content_str:
                url_matches = _URL_REGEX_PATTERN.findall(content_str)
                extracted_urls = list(set(url_matches))
        
        print(f"[File Analyzer] Extracted {len(extracted_urls)} URLs")
        
        # Remove duplicates and filter valid URLs
        unique_urls = list(set(extracted_urls))
        valid_urls = []
        for url in unique_urls:
            try:
                parsed = urlparse(url)
                if parsed.scheme and parsed.netloc:
                    valid_urls.append(url)
            except:
                pass
        
        # Limit to 30 URLs for performance
        valid_urls = valid_urls[:30]
        
        # Lookup URLs in VirusTotal
        vt_results_urls = {}
        
        def lookup_url(url):
            try:
                parsed = urlparse(url)
                if not parsed.scheme or not parsed.netloc:
                    return url, None, "Invalid URL format"
                
                # Try full URL lookup first
                try:
                    result, error = vt_fetch(url)
                    if result and not error:
                        return url, result, error
                except:
                    pass
                
                # Fall back to domain lookup
                domain = parsed.netloc
                if domain:
                    result, error = vt_fetch(domain)
                    return url, result, error
                return url, None, "Invalid URL"
            except Exception as e:
                return url, None, str(e)
        
        # Lookup URLs (limit to 30)
        if valid_urls:
            with ThreadPoolExecutor(max_workers=3) as executor:
                url_futures = {executor.submit(lookup_url, url): url for url in valid_urls}
                for future in as_completed(url_futures):
                    url, result, error = future.result()
                    if result:
                        vt_results_urls[url] = result
                    elif error:
                        vt_results_urls[url] = {"error": str(error)}
        
        # Format results
        urls_list = []
        for url in valid_urls:
            result = vt_results_urls.get(url, {})
            if isinstance(result, dict) and "error" not in result:
                community_score = result.get("community_score", 0) or 0
                malicious = result.get("malicious", 0) or 0
                suspicious = result.get("suspicious", 0) or 0
                harmless = result.get("harmless", 0) or 0
                undetected = result.get("undetected", 0) or 0
                total = malicious + suspicious + harmless + undetected
                
                urls_list.append({
                    "url": str(url),
                    "malicious": int(malicious),
                    "suspicious": int(suspicious),
                    "harmless": int(harmless),
                    "undetected": int(undetected),
                    "total": int(total),
                    "community_score": int(community_score),
                    "vt_community_link": get_vt_community_link(url, "url"),
                    "error": None
                })
            else:
                urls_list.append({
                    "url": str(url),
                    "malicious": 0,
                    "suspicious": 0,
                    "harmless": 0,
                    "undetected": 0,
                    "total": 0,
                    "community_score": 0,
                    "vt_community_link": None,
                    "error": str(result.get("error")) if isinstance(result, dict) and "error" in result else "No result"
                })
        
        # Sort by threat level (malicious first, then suspicious)
        urls_list.sort(key=lambda x: (x.get("malicious", 0) * 1000 + x.get("suspicious", 0)), reverse=True)
        
        # Calculate overall file score
        file_score = 0
        for url_data in urls_list:
            if url_data.get("malicious", 0) > 0:
                file_score += 100
            elif url_data.get("suspicious", 0) > 0:
                file_score += 50
        
        return jsonify({
            "success": True,
            "filename": filename,
            "file_hash": file_hash,
            "urls_count": len(urls_list),
            "file_score": file_score,
            "urls": urls_list
        })
        
    except Exception as e:
        print(f"[File Analyzer] Error analyzing file: {e}")
        traceback.print_exc()
        return jsonify({'error': f'Error analyzing file: {str(e)}'}), 500

@app.route("/attachment/<file_hash>")
def get_attachment(file_hash: str):
    """
    Route to retrieve attachment file content by hash.
    Used for client-side file viewing.
    """
    print(f"[Attachment] Request to retrieve attachment with hash: {file_hash[:16]}...")
    print(f"[Attachment] Full hash: {file_hash}")
    print(f"[Attachment] Cache contains {len(_attachment_cache)} attachment(s)")
    
    # Debug: List all cached hashes
    if _attachment_cache:
        cached_hashes = list(_attachment_cache.keys())
        print(f"[Attachment] Cached hashes (first 50 chars each): {[h[:50] + '...' for h in cached_hashes[:5]]}")
    
    # Get cached attachment
    cached_attachment = get_cached_attachment(file_hash)
    if not cached_attachment:
        # Try case-insensitive match
        file_hash_lower = file_hash.lower()
        for cached_hash in list(_attachment_cache.keys()):
            if cached_hash.lower() == file_hash_lower:
                print(f"[Attachment] Found case-insensitive match: {cached_hash[:16]}...")
                cached_attachment = _attachment_cache[cached_hash]
                break
        
        if not cached_attachment:
            print(f"[Attachment] ERROR: Attachment not found in cache")
            print(f"[Attachment] Searched for: {file_hash}")
            print(f"[Attachment] Available hashes: {list(_attachment_cache.keys())[:3]}")
            return jsonify({'error': 'Attachment not found or expired from cache'}), 404
    
    file_content = cached_attachment['content']
    filename = cached_attachment['filename']
    mime_type = cached_attachment.get('mime_type', 'application/octet-stream')
    
    # Determine MIME type from filename if not set
    if not mime_type or mime_type == 'application/octet-stream':
        ext = os.path.splitext(filename)[1].lower()
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.txt': 'text/plain',
            '.log': 'text/plain'
        }
        mime_type = mime_types.get(ext, 'application/octet-stream')
    
    from flask import Response
    response = Response(file_content, mimetype=mime_type)
    response.headers['Content-Disposition'] = f'inline; filename="{filename}"'
    return response

# ================================

if __name__ == "__main__":
    app.config.update(
        SESSION_COOKIE_SECURE=True,
        SESSION_COOKIE_HTTPONLY=True,
        SESSION_COOKIE_SAMESITE="Lax",
    )
    print(">>> Flask starting on http://127.0.0.1:5000")
    print(">>> NOTE: Make sure you have created a .env file with your API keys.")
    print(">>> DEBUG MODE IS ON. Server will auto-reload on code changes.")
    app.run(host="127.0.0.1", port=5000, debug=True)